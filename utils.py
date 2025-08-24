import os
import json
import re
from datetime import datetime
from decimal import Decimal
import locale
from calendar import monthrange
import math
import shutil
from flask import current_app
from collections import defaultdict
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor
import zipfile
import tempfile
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

# Konfigurasi upload file
UPLOAD_FOLDER = 'static/uploads'
ALLOWED_EXTENSIONS = {'pdf'}

def format_tanggal_indonesia(dt):
    try:
        locale.setlocale(locale.LC_TIME, 'id_ID.UTF-8')
    except:
        try:
            locale.setlocale(locale.LC_TIME, 'indonesian')
        except:
            return dt.strftime('%A, %d %B %Y')
    return dt.strftime('%A, %d %B %Y')

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Fungsi helper untuk mengkonversi Decimal ke float untuk JSON serialization
def convert_decimal_for_json(obj):
    """Mengkonversi objek Decimal ke float untuk JSON serialization"""
    if isinstance(obj, dict):
        return {key: convert_decimal_for_json(value) for key, value in obj.items()}
    elif isinstance(obj, list):
        return [convert_decimal_for_json(item) for item in obj]
    elif hasattr(obj, '__class__') and obj.__class__.__name__ == 'Decimal':
        return float(obj)
    else:
        return obj

def create_upload_path(perguruan_tinggi, nama_ketua):
    """Membuat path upload berdasarkan perguruan tinggi dan nama ketua"""
    # Bersihkan nama perguruan tinggi dan ketua untuk path yang aman
    perguruan_tinggi_clean = re.sub(r'[^\w\s-]', '', perguruan_tinggi).strip()
    nama_ketua_clean = re.sub(r'[^\w\s-]', '', nama_ketua).strip()
    
    # Ganti spasi dengan underscore
    perguruan_tinggi_clean = re.sub(r'\s+', '_', perguruan_tinggi_clean)
    nama_ketua_clean = re.sub(r'\s+', '_', nama_ketua_clean)
    
    # Buat path
    upload_path = os.path.join(UPLOAD_FOLDER, 'Proposal', perguruan_tinggi_clean, nama_ketua_clean)
    
    # Buat direktori jika belum ada
    os.makedirs(upload_path, exist_ok=True)
    
    return upload_path

def create_standardized_file_path(file_type, perguruan_tinggi, nama_ketua, judul_usaha, original_filename, subfolder=None):
    """
    Membuat path file yang standar untuk upload
    Returns: (physical_path, db_path)
    """
    # Bersihkan nama untuk path yang aman
    perguruan_tinggi_clean = re.sub(r'[^\w\s-]', '', perguruan_tinggi).strip()
    nama_ketua_clean = re.sub(r'[^\w\s-]', '', nama_ketua).strip()
    judul_clean = re.sub(r'[^\w\s-]', '', judul_usaha).strip()
    
    # Ganti spasi dengan underscore
    perguruan_tinggi_clean = re.sub(r'\s+', '_', perguruan_tinggi_clean)
    nama_ketua_clean = re.sub(r'\s+', '_', nama_ketua_clean)
    judul_clean = re.sub(r'\s+', '_', judul_clean)
    
    # Buat base path
    base_path = os.path.join(UPLOAD_FOLDER, 'Proposal', perguruan_tinggi_clean, nama_ketua_clean)
    
    # Tambahkan subfolder jika ada
    if subfolder:
        base_path = os.path.join(base_path, subfolder)
    
    # Buat direktori jika belum ada
    os.makedirs(base_path, exist_ok=True)
    
    # Generate nama file yang unik
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    file_ext = os.path.splitext(original_filename)[1]
    safe_filename = f"{file_type}_{judul_clean}_{timestamp}{file_ext}"
    
    # Path fisik untuk penyimpanan
    physical_path = os.path.join(base_path, safe_filename)
    
    # Path untuk database (relative dari static/uploads)
    db_path = os.path.join('Proposal', perguruan_tinggi_clean, nama_ketua_clean)
    if subfolder:
        db_path = os.path.join(db_path, subfolder)
    db_path = os.path.join(db_path, safe_filename)
    
    return physical_path, db_path

def group_anggaran_data(anggaran_data):
    """Mengelompokkan data anggaran berdasarkan kegiatan_utama, kegiatan, penanggung_jawab"""
    grouped = {}
    for item in anggaran_data:
        key = (item['kegiatan_utama'], item['kegiatan'], item['penanggung_jawab'])
        if key not in grouped:
            grouped[key] = []
        grouped[key].append(item)
    return grouped

def flatten_anggaran_data(anggaran_data):
    """Flatten data anggaran dengan informasi rowspan untuk tabel"""
    
    # Group by kegiatan_utama, kegiatan, penanggung_jawab
    flat_rows = []
    # Grouping
    grouped = defaultdict(lambda: defaultdict(lambda: defaultdict(list)))
    for row in anggaran_data:
        grouped[row['kegiatan_utama']][row['kegiatan']][row['penanggung_jawab']].append(row)
    
    # Flatten with rowspan info
    for kegiatan_utama, keg_dict in grouped.items():
        kegiatan_utama_rowspan = sum(
            sum(len(pj_rows) for pj_rows in kegiatan_dict.values())
            for kegiatan_dict in keg_dict.values()
        )
        first_utama = True
        for kegiatan, kegiatan_dict in keg_dict.items():
            kegiatan_rowspan = sum(len(pj_rows) for pj_rows in kegiatan_dict.values())
            first_kegiatan = True
            
            # Process each penanggung_jawab group
            for penanggung_jawab, pj_rows in kegiatan_dict.items():
                pj_rowspan = len(pj_rows)
                first_pj = True
                
                # Group by target_capaian untuk penanggung_jawab ini
                target_capaian_groups = {}
                for row in pj_rows:
                    target_capaian = row['target_capaian']
                    if target_capaian not in target_capaian_groups:
                        target_capaian_groups[target_capaian] = []
                    target_capaian_groups[target_capaian].append(row)
                
                # Process each target_capaian group
                for target_capaian, target_rows in target_capaian_groups.items():
                    target_rowspan = len(target_rows)
                    first_target = True
                    
                    for row in target_rows:
                        flat_row = dict(row)
                        flat_row['show_kegiatan_utama'] = first_utama
                        flat_row['kegiatan_utama_rowspan'] = kegiatan_utama_rowspan if first_utama else 0
                        flat_row['show_kegiatan'] = first_kegiatan
                        flat_row['kegiatan_rowspan'] = kegiatan_rowspan if first_kegiatan else 0
                        flat_row['show_penanggung_jawab'] = first_pj
                        flat_row['penanggung_jawab_rowspan'] = pj_rowspan if first_pj else 0
                        # Target capaian dengan rowspan berdasarkan grup yang sama
                        flat_row['show_target_capaian'] = first_target
                        flat_row['target_capaian_rowspan'] = target_rowspan if first_target else 0
                        flat_rows.append(flat_row)
                        first_utama = False
                        first_kegiatan = False
                        first_pj = False
                        first_target = False
    
    return flat_rows 

def copy_word_header(src_docx_path, dst_doc):
    """
    Menyalin header dari src_docx_path (misal: Kop Surat.docx) ke dokumen Word baru (dst_doc).
    dst_doc: objek Document dari python-docx
    """
    try:
        # Buka file sumber untuk mengekstrak teks header
        src_doc = Document(src_docx_path)
        
        # Ekstrak teks dari header (tanpa gambar)
        header_texts = []
        for paragraph in src_doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                header_texts.append(text)
        
        # Tambahkan teks header ke dokumen tujuan (tanpa logo)
        for i, text in enumerate(header_texts[:6]):  # Ambil 6 baris pertama
            if text:
                header_para = dst_doc.add_paragraph()
                
                # Cek apakah ada link dalam teks
                if 'www.unjaya.ac.id' in text or 'info@unjaya.ac.id' in text:
                    # Pisahkan teks menjadi bagian-bagian
                    parts = []
                    if 'www.unjaya.ac.id' in text:
                        parts = text.split('www.unjaya.ac.id')
                        # Tambahkan teks sebelum link
                        if parts[0]:
                            run1 = header_para.add_run(parts[0])
                            run1.font.name = 'Times New Roman'
                            run1.font.size = Pt(10)
                            run1.bold = False
                        
                        # Tambahkan link dengan warna biru dan garis bawah
                        link_run = header_para.add_run('www.unjaya.ac.id')
                        link_run.font.name = 'Times New Roman'
                        link_run.font.size = Pt(10)
                        link_run.bold = False
                        link_run.font.color.rgb = RGBColor(0, 0, 255)  # Warna biru
                        link_run.font.underline = True
                        
                        # Tambahkan teks setelah link
                        if len(parts) > 1 and parts[1]:
                            run2 = header_para.add_run(parts[1])
                            run2.font.name = 'Times New Roman'
                            run2.font.size = Pt(10)
                            run2.bold = False
                    elif 'info@unjaya.ac.id' in text:
                        parts = text.split('info@unjaya.ac.id')
                        # Tambahkan teks sebelum link
                        if parts[0]:
                            run1 = header_para.add_run(parts[0])
                            run1.font.name = 'Times New Roman'
                            run1.font.size = Pt(10)
                            run1.bold = False
                        
                        # Tambahkan link dengan warna biru dan garis bawah
                        link_run = header_para.add_run('info@unjaya.ac.id')
                        link_run.font.name = 'Times New Roman'
                        link_run.font.size = Pt(10)
                        link_run.bold = False
                        link_run.font.color.rgb = RGBColor(0, 0, 255)  # Warna biru
                        link_run.font.underline = True
                        
                        # Tambahkan teks setelah link
                        if len(parts) > 1 and parts[1]:
                            run2 = header_para.add_run(parts[1])
                            run2.font.name = 'Times New Roman'
                            run2.font.size = Pt(10)
                            run2.bold = False
                else:
                    # Teks biasa tanpa link
                    header_run = header_para.add_run(text)
                    header_run.font.name = 'Times New Roman'
                    
                    # Hanya baris 1-2 (Yayasan dan Universitas) yang bold dan ukuran 12pt
                    if i < 2:  # Baris 0 dan 1 (Yayasan dan Universitas)
                        header_run.bold = True
                        header_run.font.size = Pt(12)
                    else:  # Baris 2-5 (alamat, telp, email) tidak bold dan ukuran 10pt
                        header_run.bold = False
                        header_run.font.size = Pt(10)
                
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        print("✅ Header Word disalin tanpa logo (Yayasan & Universitas bold, lainnya normal)")
        return dst_doc
        
    except Exception as e:
        print(f"Error copying Word header: {e}")
        # Jika gagal, return dokumen asli tanpa header
        return dst_doc

def copy_pdf_header(src_pdf_path, dst_canvas):
    """
    Menyalin header dari src_pdf_path (misal: Kop Surat.pdf) ke canvas PDF baru.
    dst_canvas: objek canvas dari reportlab
    """
    try:
        # Baca halaman pertama dari PDF sumber
        PdfReader = None
        
        # Coba import PyPDF2 terlebih dahulu
        try:
            import PyPDF2  # type: ignore
            PdfReader = PyPDF2.PdfReader
        except ImportError:
            pass
        
        # Jika PyPDF2 tidak tersedia, coba pypdf
        if PdfReader is None:
            try:
                import pypdf  # type: ignore
                PdfReader = pypdf.PdfReader
            except ImportError:
                pass
        
        # Jika kedua library tidak tersedia, return False
        if PdfReader is None:
            # PyPDF2 atau pypdf tidak tersedia, menggunakan fallback
            return False
        
        reader = PdfReader(src_pdf_path)
        if len(reader.pages) > 0:
            # Ambil halaman pertama
            page = reader.pages[0]
            
            # Salin ukuran halaman
            dst_canvas.setPageSize(page.mediabox)
            
            # Coba salin konten halaman (gambar, teks, dll)
            try:
                # Salin gambar dari halaman pertama
                if '/XObject' in page['/Resources']:
                    xObject = page['/Resources']['/XObject'].get_object()
                    for obj in xObject:
                        if xObject[obj]['/Subtype'] == '/Image':
                            # Salin gambar ke canvas
                            img_data = xObject[obj].get_data()
                            # Simpan gambar sementara
                            temp_img_path = f"temp_img_{obj}.png"
                            with open(temp_img_path, 'wb') as img_file:
                                img_file.write(img_data)
                            
                            # Tambahkan gambar ke canvas
                            from reportlab.platypus import Image
                            img = Image(temp_img_path)
                            img.drawOn(dst_canvas, 0, 0)
                            
                            # Hapus file sementara
                            os.remove(temp_img_path)
                
                # Salin teks dari halaman pertama
                if '/Font' in page['/Resources']:
                    # Extract text content
                    text_content = page.extract_text()
                    if text_content:
                        # Tambahkan teks ke canvas
                        dst_canvas.setFont("Helvetica", 12)
                        dst_canvas.drawString(50, 750, text_content)
                
                return True
                
            except Exception as e:
                # Tidak dapat menyalin konten PDF, menggunakan fallback
                # Tetap return True karena ukuran halaman sudah disalin
                return True
                
    except Exception as e:
        print(f"Error copying PDF header: {e}")
        return False

def extract_logo_from_kop_surat():
    """Mengekstrak gambar logo dari file Kop Surat.docx"""
    try:
        from docx import Document
        import tempfile
        
        kop_surat_path = "Kop Surat.docx"
        if not os.path.exists(kop_surat_path):
            # File Kop Surat.docx tidak ditemukan
            return None, None
        
        doc = Document(kop_surat_path)
        
        # Buat direktori temp jika belum ada
        temp_dir = os.path.join('static', 'temp')
        os.makedirs(temp_dir, exist_ok=True)
        
        logo_paths = []
        
        # Cari semua gambar dalam dokumen
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                # Ekstrak gambar
                image_data = rel.target_part.blob
                
                # Simpan gambar sementara
                logo_filename = f'logo_{len(logo_paths)}.png'
                logo_path = os.path.join(temp_dir, logo_filename)
                
                with open(logo_path, 'wb') as f:
                    f.write(image_data)
                
                logo_paths.append(logo_path)
                # Logo berhasil diekstrak
        
        # Jika ada 2 logo, return keduanya
        if len(logo_paths) >= 2:
            return logo_paths[0], logo_paths[1]  # Logo Kartika, Logo Unjaya
        elif len(logo_paths) == 1:
            return logo_paths[0], logo_paths[0]  # Gunakan logo yang sama untuk keduanya
        else:
            # Tidak ada logo yang ditemukan dalam Kop Surat.docx
            return None, None
        
        return None
        
    except Exception as e:
        print(f"Error extracting logo from kop surat: {e}")
        return None

def create_pdf_header_from_kop_surat_pdf():
    """
    Membuat header PDF yang identik dengan Kop Surat.pdf
    """
    # Path ke Kop Surat.pdf
    kop_surat_pdf_path = os.path.join(os.getcwd(), 'Kop Surat.pdf')
    
    if not os.path.exists(kop_surat_pdf_path):
        # Silent fallback tanpa warning yang mengganggu
        return create_pdf_header_from_template()
    
    try:
        # Baca Kop Surat.pdf untuk mendapatkan ukuran dan layout
        PdfReader = None
        try:
            import PyPDF2  # type: ignore
            PdfReader = PyPDF2.PdfReader
        except ImportError:
            pass
        
        if PdfReader is None:
            try:
                import pypdf  # type: ignore
                PdfReader = pypdf.PdfReader
            except ImportError:
                pass
        
        if PdfReader is None:
            # Silent fallback tanpa warning yang mengganggu
            return create_pdf_header_from_template()
        
        reader = PdfReader(kop_surat_pdf_path)
        if len(reader.pages) > 0:
            page = reader.pages[0]
            # Gunakan ukuran halaman dari Kop Surat.pdf
            page_width = float(page.mediabox.width)
            page_height = float(page.mediabox.height)
            
            # Coba ekstrak teks dari Kop Surat.pdf untuk mendapatkan konten yang tepat
            try:
                text_content = page.extract_text()
                # Konten dari Kop Surat.pdf berhasil diekstrak
            except:
                text_content = ""
            
            # HAPUS LOGO - Set logo_path dan shield_logo_path ke None
            logo_path = None
            shield_logo_path = None
            # Logo Kartika dan Unjaya dihapus dari header PDF
            
            # Buat header data berdasarkan konten Kop Surat.pdf
            header_data = []
            
            # Parse konten teks untuk mendapatkan header yang tepat
            lines = text_content.split('\n') if text_content else []
            
            # Jika berhasil mengekstrak teks, gunakan itu
            if lines and len(lines) >= 3:
                # Ambil baris-baris header dari PDF
                for i, line in enumerate(lines[:6]):  # Ambil 6 baris pertama
                    line = line.strip()
                    if line:
                        header_data.append(['', line, ''])  # Hapus logo, hanya teks
                    else:
                        header_data.append(['', '', ''])  # Baris kosong tanpa logo
            else:
                # Fallback ke header default tanpa logo
                # Baris 1: YAYASAN KARTIKA EKA PAKSI (tanpa logo)
                header_data.append(['', 'YAYASAN KARTIKA EKA PAKSI', ''])
                
                # Baris 2: UNIVERSITAS JENDERAL ACHMAD YANI YOGYAKARTA (tanpa logo)
                header_data.append(['', 'UNIVERSITAS JENDERAL ACHMAD YANI YOGYAKARTA', ''])
                
                # Baris 3: (spasi) (tanpa logo)
                header_data.append(['', '', ''])
                
                # Baris 4: Alamat lengkap (tanpa logo)
                alamat_lengkap = "Jl Siliwangi Ringroad Barat, Banyuraden, Gamping, Sleman, Yogyakarta 55293"
                header_data.append(['', alamat_lengkap, ''])
                
                # Baris 5: Telp, Fax, Website (tanpa logo)
                kontak_lengkap = "Telp. (0274) 552489, 552851 Fax. (0274) 557228 Website: www.unjaya.ac.id"
                header_data.append(['', kontak_lengkap, ''])
                
                # Baris 6: Email (tanpa logo)
                email_lengkap = "Email: info@unjaya.ac.id"
                header_data.append(['', email_lengkap, ''])
            
            return {
                'header_data': header_data,
                'logo_path': logo_path,
                'shield_logo_path': shield_logo_path,
                'page_width': page_width,
                'page_height': page_height,
                'source_pdf': kop_surat_pdf_path  # Tambahkan referensi ke file sumber
            }
            
    except Exception as e:
        print(f"Error reading Kop Surat.pdf: {e}")
        return create_pdf_header_from_template()

def copy_pdf_header_to_document(src_pdf_path, dst_doc):
    """
    Menyalin header dari Kop Surat.pdf ke dokumen PDF yang baru dibuat
    dst_doc: objek SimpleDocTemplate dari reportlab
    """
    try:
        # Baca halaman pertama dari Kop Surat.pdf
        PdfReader = None
        try:
            import PyPDF2  # type: ignore
            PdfReader = PyPDF2.PdfReader
        except ImportError:
            pass
        
        if PdfReader is None:
            try:
                import pypdf  # type: ignore
                PdfReader = pypdf.PdfReader
            except ImportError:
                pass
        
        if PdfReader is None:
            # PyPDF2 atau pypdf tidak tersedia, menggunakan fallback
            return False
        
        reader = PdfReader(src_pdf_path)
        if len(reader.pages) > 0:
            page = reader.pages[0]
            
            # Salin ukuran halaman
            dst_doc.pagesize = page.mediabox
            
            # Coba salin gambar dari header
            try:
                if '/XObject' in page['/Resources']:
                    xObject = page['/Resources']['/XObject'].get_object()
                    for obj in xObject:
                        if xObject[obj]['/Subtype'] == '/Image':
                            # Simpan gambar sementara
                            img_data = xObject[obj].get_data()
                            temp_img_path = f"temp_header_img_{obj}.png"
                            with open(temp_img_path, 'wb') as img_file:
                                img_file.write(img_data)
                            
                            # Tambahkan gambar ke dokumen
                            from reportlab.platypus import Image
                            img = Image(temp_img_path, width=1*inch, height=1*inch)
                            dst_doc.story.append(img)
                            
                            # Hapus file sementara
                            os.remove(temp_img_path)
                
                return True
                
            except Exception as e:
                print(f"Warning: Tidak dapat menyalin gambar header: {e}")
                return True
                
    except Exception as e:
        print(f"Error copying PDF header to document: {e}")
        return False

def merge_pdf_with_header(header_pdf_path, content_pdf_path, output_path):
    """
    Menggabungkan header dari Kop Surat.pdf dengan konten PDF yang baru dibuat
    """
    try:
        PdfReader = None
        PdfWriter = None
        
        # Coba import PyPDF2 terlebih dahulu
        try:
            import PyPDF2  # type: ignore
            PdfReader = PyPDF2.PdfReader
            PdfWriter = PyPDF2.PdfWriter
        except ImportError:
            pass
        
        # Jika PyPDF2 tidak tersedia, coba pypdf
        if PdfReader is None:
            try:
                import pypdf  # type: ignore
                PdfReader = pypdf.PdfReader
                PdfWriter = pypdf.PdfWriter
            except ImportError:
                pass
        
        if PdfReader is None or PdfWriter is None:
            # PyPDF2 atau pypdf tidak tersedia, menggunakan fallback
            return False
        
        # Baca header PDF
        header_reader = PdfReader(header_pdf_path)
        if len(header_reader.pages) == 0:
            print("Header PDF kosong")
            return False
        
        # Baca konten PDF
        content_reader = PdfReader(content_pdf_path)
        if len(content_reader.pages) == 0:
            print("Konten PDF kosong")
            return False
        
        # Buat PDF writer
        writer = PdfWriter()
        
        # Ambil halaman pertama dari header (hanya header)
        header_page = header_reader.pages[0]
        
        # Untuk setiap halaman konten, gabungkan dengan header
        for i, content_page in enumerate(content_reader.pages):
            # Buat halaman baru dengan ukuran header
            merged_page = writer.add_blank_page(
                width=header_page.mediabox.width,
                height=header_page.mediabox.height
            )
            
            # Salin konten header ke halaman baru
            merged_page.merge_page(header_page)
            
            # Salin konten ke halaman baru (dengan offset untuk header)
            # Offset header sekitar 100 point dari atas
            content_page.mediabox.upper_left = (
                content_page.mediabox.left,
                header_page.mediabox.height - 100
            )
            merged_page.merge_page(content_page)
        
        # Tulis PDF hasil
        with open(output_path, 'wb') as output_file:
            writer.write(output_file)
        
        return True
        
    except Exception as e:
        print(f"Error merging PDF with header: {e}")
        return False

def create_pdf_header_from_template():
    """
    Membuat header PDF berdasarkan template default
    """
    # Gunakan ukuran A4 standar
    page_width = 595.276
    page_height = 841.890
    
    # HAPUS LOGO - Set logo_path dan shield_logo_path ke None
    logo_path = None
    shield_logo_path = None
    # Logo Kartika dan Unjaya dihapus dari template header PDF
    
    # Buat header data default tanpa logo
    header_data = []
    
    # Baris 1: YAYASAN KARTIKA EKA PAKSI (tanpa logo)
    header_data.append(['', 'YAYASAN KARTIKA EKA PAKSI', ''])
    
    # Baris 2: UNIVERSITAS JENDERAL ACHMAD YANI YOGYAKARTA (tanpa logo)
    header_data.append(['', 'UNIVERSITAS JENDERAL ACHMAD YANI YOGYAKARTA', ''])
    
    # Baris 3: (spasi) (tanpa logo)
    header_data.append(['', '', ''])
    
    # Baris 4: Alamat lengkap (tanpa logo)
    alamat_lengkap = "Jl Siliwangi Ringroad Barat, Banyuraden, Gamping, Sleman, Yogyakarta 55293"
    header_data.append(['', alamat_lengkap, ''])
    
    # Baris 5: Telp, Fax, Website (tanpa logo) - dengan link biru
    kontak_lengkap = "Telp. (0274) 552489, 552851 Fax. (0274) 557228 Website: www.unjaya.ac.id"
    header_data.append(['', kontak_lengkap, ''])
    
    # Baris 6: Email (tanpa logo) - dengan link biru
    email_lengkap = "Email: info@unjaya.ac.id"
    header_data.append(['', email_lengkap, ''])
    
    return {
        'header_data': header_data,
        'logo_path': logo_path,
        'shield_logo_path': shield_logo_path,
        'page_width': page_width,
        'page_height': page_height
    }

def create_header_with_styled_links(header_data):
    """
    Membuat header dengan styling link yang benar menggunakan ReportLab
    """
    from reportlab.platypus import Paragraph
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib import colors
    
    styles = getSampleStyleSheet()
    
    # Buat style untuk teks biasa
    normal_style = ParagraphStyle(
        'NormalStyle',
        parent=styles['Normal'],
        fontName='Times-Roman',
        fontSize=10,
        textColor=colors.black,
        alignment=1  # Center
    )
    
    # Buat style untuk teks bold
    bold_style = ParagraphStyle(
        'BoldStyle',
        parent=styles['Normal'],
        fontName='Times-Bold',
        fontSize=12,
        textColor=colors.black,
        alignment=1  # Center
    )
    
    styled_headers = []
    
    for i, row in enumerate(header_data):
        if row[1]:  # Jika ada teks di kolom tengah
            text = row[1]
            
            # Teks biasa tanpa link (ReportLab tidak mendukung HTML tags)
            if i < 2:  # Baris 1-2 (Yayasan dan Universitas)
                styled_headers.append(Paragraph(text, bold_style))
            else:  # Baris 3-6 (alamat, telp, email)
                styled_headers.append(Paragraph(text, normal_style))
        else:
            # Baris kosong
            styled_headers.append(Paragraph('', normal_style))
    
    return styled_headers 