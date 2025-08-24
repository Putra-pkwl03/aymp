from flask import Flask, render_template, request, redirect, url_for, flash, session, jsonify, send_file, abort
from datetime import datetime, timedelta
import json
from flask_mysqldb import MySQL
import MySQLdb.cursors
import re
from flask import send_from_directory
import os
import shutil
import mimetypes
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
import uuid
from collections import defaultdict
from datetime import datetime
from decimal import Decimal
import locale
from calendar import monthrange
import math
from utils import copy_word_header
import sys
import threading
import time
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Inches
import io

# Import konfigurasi
try:
    from config import AUTO_FILL_CONFIG, INTERVAL_CONFIG
except ImportError:
    # Fallback konfigurasi jika file config.py tidak ada
    AUTO_FILL_CONFIG = {
        'CHECK_INTERVAL': 60,  # 60 detik = 1 menit
        'PRIORITY_CHECK_INTERVAL': 30,  # 30 detik untuk sesi urgent
        'DB_HOST': 'localhost',
        'DB_USER': 'root',
        'DB_PASSWD': '',
        'DB_NAME': 'aymp',
        'AUTO_FILL_ENABLED': True,
        'ENABLE_SMART_CHECKING': True,  # ✅ Enable smart checking
        'MAX_BATCH_SIZE': 10,  # Batch size untuk processing
        'DEFAULT_SCORE': 0,
        'DEFAULT_STATUS': 'terkunci',
        'DEFAULT_KETERANGAN': 'Auto-filled: Sesi terlewat',
        'SKIP_OLD_DATA': True,  # ✅ Skip auto-fill untuk data lama
        'OLD_DATA_THRESHOLD_DAYS': 7,  # Data yang dibuat lebih dari 7 hari dianggap lama
        'STRICT_OLD_DATA_CHECK': True,  # ✅ Periksa ketat untuk data lama
        'DISABLE_AUTO_FILL_FOR_OLD_DATA': True,  # ✅ Nonaktifkan auto-fill untuk data lama
        'STRICT_DATE_VALIDATION': True  # ✅ Validasi tanggal yang ketat
    }
    
    INTERVAL_CONFIG = {
        'jam': {'multiplier': 3600, 'description': 'per jam'},
        'hari': {'multiplier': 86400, 'description': 'per hari'},
        'minggu': {'multiplier': 604800, 'description': 'per minggu'},
        'bulan': {'multiplier': 2592000, 'description': 'per bulan'}
    }

app = Flask(__name__)

app.secret_key = 'your_secret_key'

# Custom filter untuk nl2br
@app.template_filter('nl2br')
def nl2br_filter(text):
    if text is None:
        return ''
    return text.replace('\n', '<br>')

app.config['MYSQL_HOST'] = 'localhost'
app.config['MYSQL_USER'] = 'root'
app.config['MYSQL_PASSWORD'] = ''  # Ganti jika ada password
app.config['MYSQL_DB'] = 'AYMP'

# Konfigurasi upload file
# Konfigurasi upload yang lebih lengkap
UPLOAD_FOLDER = 'static/uploads'
ALLOWED_EXTENSIONS = {'pdf'}
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Tambahkan konfigurasi keamanan
app.config['SESSION_COOKIE_SECURE'] = False  # Set False untuk HTTP development
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=24)
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'

mysql = MySQL(app)

# ✅ PERBAIKAN: Tambahkan definisi yang hilang
def get_app_functions():
    """Return app functions untuk kompatibilitas dengan file lain"""
    return {
        'mysql': mysql,
        'app': app
    }

# ✅ PERBAIKAN: Setup logger sederhana
import logging
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

# Jika belum ada handler, tambahkan
if not logger.handlers:
    handler = logging.StreamHandler()
    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    handler.setFormatter(formatter)
    logger.addHandler(handler)


def format_tanggal_indonesia(dt):
    try:
        locale.setlocale(locale.LC_TIME, 'id_ID.UTF-8')
    except:
        try:
            locale.setlocale(locale.LC_TIME, 'indonesian')
        except:
            return dt.strftime('%A, %d %B %Y')
    return dt.strftime('%A, %d %B %Y')

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Fungsi helper untuk mengkonversi Decimal ke float untuk JSON serialization
def convert_decimal_for_json(obj):
    """Mengkonversi objek Decimal ke float untuk JSON serialization"""
    if isinstance(obj, dict):
        return {key: convert_decimal_for_json(value) for key, value in obj.items()}
    elif isinstance(obj, list):
        return [convert_decimal_for_json(item) for item in obj]
    elif hasattr(obj, '__class__') and obj.__class__.__name__ == 'Decimal':
        return float(obj)
    else:
        return obj

# Fungsi untuk membaca kop surat dari file
def get_kop_surat_content():
    """Membaca konten kop surat dari file Kop Surat.docx"""
    try:
        from docx import Document
        import os
        
        kop_surat_path = "Kop Surat.docx"
        if not os.path.exists(kop_surat_path):
            # Jika file tidak ada, gunakan kop surat default
            return {
                'yayasan': 'YAYASAN KARTIKA EKA PAKSI',
                'universitas': 'UNIVERSITAS JENDERAL ACHMAD YANI YOGYAKARTA',
                'alamat': 'Jl Siliwangi Ringroad Barat, Banyuraden, Gamping, Sleman, Yogyakarta 55293',
                'kontak': 'Telp. (0274) 552489, 552851 Fax. (0274) 557228 Website: www.unjaya.ac.id',
                'email': 'Email: info@unjaya.ac.id'
            }
        
        doc = Document(kop_surat_path)
        content = []
        
        # Ambil semua paragraph dari dokumen kop surat
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text.strip())
        
        # Parse konten kop surat
        kop_surat = {
            'yayasan': content[0] if len(content) > 0 else 'YAYASAN KARTIKA EKA PAKSI',
            'universitas': content[1] if len(content) > 1 else 'UNIVERSITAS JENDERAL ACHMAD YANI YOGYAKARTA',
            'alamat': content[2] if len(content) > 2 else 'Jl Siliwangi Ringroad Barat, Banyuraden, Gamping, Sleman, Yogyakarta 55293',
            'kontak': content[3] if len(content) > 3 else 'Telp. (0274) 552489, 552851 Fax. (0274) 557228 Website: www.unjaya.ac.id',
            'email': content[4] if len(content) > 4 else 'Email: info@unjaya.ac.id'
        }
        
        return kop_surat
        
    except Exception as e:
        print(f"Error reading kop surat: {e}")
        # Return default kop surat jika ada error
        return {
            'yayasan': 'YAYASAN KARTIKA EKA PAKSI',
            'universitas': 'UNIVERSITAS JENDERAL ACHMAD YANI YOGYAKARTA',
            'alamat': 'Jl Siliwangi Ringroad Barat, Banyuraden, Gamping, Sleman, Yogyakarta 55293',
            'kontak': 'Telp. (0274) 552489, 552851 Fax. (0274) 557228 Website: www.unjaya.ac.id',
            'email': 'Email: info@unjaya.ac.id'
        }

def extract_logo_from_kop_surat():
    """Mengekstrak gambar logo dari file Kop Surat.docx"""
    try:
        from docx import Document
        import os
        import tempfile
        
        kop_surat_path = "Kop Surat.docx"
        if not os.path.exists(kop_surat_path):
            return None
        
        doc = Document(kop_surat_path)
        
        # Cari gambar dalam dokumen
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                # Ekstrak gambar
                image_data = rel.target_part.blob
                
                # Simpan gambar sementara
                temp_dir = os.path.join('static', 'temp')
                os.makedirs(temp_dir, exist_ok=True)
                
                logo_path = os.path.join(temp_dir, 'logo_kop_surat.png')
                with open(logo_path, 'wb') as f:
                    f.write(image_data)
                
                return logo_path
        
        return None
        
    except Exception as e:
        print(f"Error extracting logo from kop surat: {e}")
        return None

def create_kop_surat_header():
    """Membuat header kop surat untuk laporan dengan tata letak horizontal"""
    kop_surat = get_kop_surat_content()
    
    # HAPUS LOGO - Set logo_path dan shield_logo_path ke None
    logo_path = None
    shield_logo_path = None
    print("✅ Logo Kartika dan Unjaya dihapus dari header kop surat")
    
    return {
        'kop_surat': kop_surat,
        'logo_path': logo_path,
        'shield_logo_path': shield_logo_path,
        'layout': 'horizontal'  # Menandakan tata letak horizontal
        }

def safe_remove_empty_folder(folder_path, folder_name="folder"):
    """
    Menghapus folder kosong dengan aman dan menghapus parent folder jika kosong
    
    Args:
        folder_path (str): Path folder yang akan dihapus
        folder_name (str): Nama folder untuk pesan log (opsional)
    
    Returns:
        bool: True jika berhasil menghapus folder, False jika gagal
    """
    try:
        if os.path.exists(folder_path) and not os.listdir(folder_path):
            os.rmdir(folder_path)
            print(f"✅ {folder_name} kosong berhasil dihapus: {folder_path}")
            
            # Hapus parent folder jika kosong
            parent_folder = os.path.dirname(folder_path)
            if os.path.exists(parent_folder) and not os.listdir(parent_folder):
                os.rmdir(parent_folder)
                print(f"✅ Parent folder kosong berhasil dihapus: {parent_folder}")
            return True
    except Exception as e:
        print(f"❌ Error saat menghapus {folder_name} {folder_path}: {e}")
    return False

# Fungsi helper untuk logging aktivitas mahasiswa
def log_mahasiswa_activity(mahasiswa_id, nim, nama_mahasiswa, jenis_aktivitas, modul, detail_modul=None, deskripsi="", data_lama=None, data_baru=None):
    """
    Mencatat aktivitas mahasiswa ke database
    Hanya mencatat untuk mahasiswa dengan status 'aktif' atau 'proses'
    """
    try:
        if not hasattr(mysql, 'connection') or mysql.connection is None:
            return False
            
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        
        # Cek status mahasiswa sebelum mencatat log
        cursor.execute('SELECT status FROM mahasiswa WHERE id = %s', (mahasiswa_id,))
        mahasiswa_status = cursor.fetchone()
        
        if not mahasiswa_status or mahasiswa_status['status'] not in ['aktif', 'proses']:
            # Tidak mencatat log untuk mahasiswa dengan status selain 'aktif' atau 'proses'
            cursor.close()
            return False
        
        # Ambil informasi IP dan User Agent
        ip_address = request.environ.get('HTTP_X_FORWARDED_FOR', request.environ.get('REMOTE_ADDR'))
        user_agent = request.environ.get('HTTP_USER_AGENT')
        
        # Insert log aktivitas
        cursor.execute('''
            INSERT INTO log_aktivitas_mahasiswa 
            (mahasiswa_id, nim, nama_mahasiswa, jenis_aktivitas, modul, detail_modul, deskripsi, data_lama, data_baru, ip_address, user_agent)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        ''', (
            mahasiswa_id, nim, nama_mahasiswa, jenis_aktivitas, modul, detail_modul, deskripsi,
            json.dumps(convert_decimal_for_json(data_lama)) if data_lama else None,
            json.dumps(convert_decimal_for_json(data_baru)) if data_baru else None,
            ip_address, user_agent
        ))
        
        mysql.connection.commit()
        cursor.close()
        return True
        
    except Exception as e:
        print(f"Error logging activity: {str(e)}")
        return False

# Fungsi untuk tracking session mahasiswa
def start_mahasiswa_session(mahasiswa_id, nim, nama_mahasiswa):
    """
    Memulai session tracking untuk mahasiswa
    Hanya untuk mahasiswa dengan status 'aktif' atau 'proses'
    """
    try:
        print(f"DEBUG: Starting session for mahasiswa_id={mahasiswa_id}, nim={nim}, nama={nama_mahasiswa}")
        
        if not hasattr(mysql, 'connection') or mysql.connection is None:
            print("DEBUG: MySQL connection not available")
            return False
            
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        
        # Cek status mahasiswa sebelum membuat session
        cursor.execute('SELECT status FROM mahasiswa WHERE id = %s', (mahasiswa_id,))
        mahasiswa_status = cursor.fetchone()
        
        print(f"DEBUG: Mahasiswa status: {mahasiswa_status}")
        
        if not mahasiswa_status or mahasiswa_status['status'] not in ['aktif', 'proses']:
            # Tidak membuat session untuk mahasiswa dengan status selain 'aktif' atau 'proses'
            print(f"DEBUG: Mahasiswa status not allowed: {mahasiswa_status['status'] if mahasiswa_status else 'None'}")
            cursor.close()
            return False
        
        # Tutup session aktif sebelumnya jika ada
        cursor.execute('''
            UPDATE session_mahasiswa 
            SET status = 'ended', logout_time = NOW(), 
                durasi_detik = TIMESTAMPDIFF(SECOND, login_time, NOW())
            WHERE mahasiswa_id = %s AND status = 'active'
        ''', (mahasiswa_id,))
        
        # Buat session baru
        ip_address = request.environ.get('HTTP_X_FORWARDED_FOR', request.environ.get('REMOTE_ADDR'))
        user_agent = request.environ.get('HTTP_USER_AGENT')
        
        # Generate session ID unik jika belum ada
        import uuid
        if 'session_id' not in session or not session['session_id']:
            session['session_id'] = str(uuid.uuid4())
        session_id = session['session_id']
        
        print(f"DEBUG: Session ID: {session_id}")
        print(f"DEBUG: IP Address: {ip_address}")
        print(f"DEBUG: User Agent: {user_agent}")
        
        cursor.execute('''
            INSERT INTO session_mahasiswa 
            (mahasiswa_id, nim, nama_mahasiswa, session_id, ip_address, user_agent)
            VALUES (%s, %s, %s, %s, %s, %s)
        ''', (mahasiswa_id, nim, nama_mahasiswa, session_id, ip_address, user_agent))
        
        mysql.connection.commit()
        print(f"DEBUG: Session inserted successfully")
        cursor.close()
        
        # Log aktivitas login (hanya untuk status 'aktif' atau 'proses')
        log_mahasiswa_activity(mahasiswa_id, nim, nama_mahasiswa, 'login', 'sistem', 'authentication', 'Login ke sistem AYMP')
        
        return True
        
    except Exception as e:
        print(f"Error starting session: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

# Fungsi untuk mengakhiri session mahasiswa
def end_mahasiswa_session(mahasiswa_id):
    """
    Mengakhiri session tracking untuk mahasiswa
    """
    try:
        if not hasattr(mysql, 'connection') or mysql.connection is None:
            return False
            
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        
        # Update session aktif
        cursor.execute('''
            UPDATE session_mahasiswa 
            SET status = 'ended', logout_time = NOW(), 
                durasi_detik = TIMESTAMPDIFF(SECOND, login_time, NOW())
            WHERE mahasiswa_id = %s AND status = 'active'
        ''', (mahasiswa_id,))
        
        mysql.connection.commit()
        cursor.close()
        return True
        
    except Exception as e:
        print(f"Error ending session: {str(e)}")
        return False

# Fungsi untuk mendapatkan informasi mahasiswa dari session
def get_mahasiswa_info_from_session():
    """
    Mengambil informasi mahasiswa dari session aktif
    """
    if 'user_type' not in session or session['user_type'] != 'mahasiswa':
        return None
        
    if not hasattr(mysql, 'connection') or mysql.connection is None:
        return None
        
    try:
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        cursor.execute('SELECT id, nim, nama_ketua FROM mahasiswa WHERE nim = %s', (session['nim'],))
        mahasiswa = cursor.fetchone()
        cursor.close()
        return mahasiswa
    except:
        return None

# Fungsi untuk menghapus file laporan kemajuan dari folder kemajuan
def hapus_file_laporan_kemajuan(proposal_id, nim):
    """
    Menghapus file laporan kemajuan dari folder static/uploads/kemajuan/
    berdasarkan proposal_id dan nim
    """
    try:
        # Ambil data proposal untuk mendapatkan perguruan tinggi, nama_ketua, judul_usaha, dan tanggal_buat
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        cursor.execute('''
            SELECT p.tanggal_buat, p.judul_usaha, m.perguruan_tinggi, m.nama_ketua
            FROM proposal p 
            INNER JOIN mahasiswa m ON p.nim = m.nim 
            WHERE p.id = %s
        ''', (proposal_id,))
        proposal_data = cursor.fetchone()
        cursor.close()
        
        if not proposal_data:
            print(f"⚠️ Data proposal tidak ditemukan untuk ID: {proposal_id}")
            return
        
        perguruan_tinggi = proposal_data['perguruan_tinggi']
        nama_ketua = proposal_data['nama_ketua']
        judul_usaha = proposal_data['judul_usaha']
        tanggal_buat = proposal_data['tanggal_buat']
        
        # Path folder kemajuan (menggunakan nama_ketua, bukan NIM)
        kemajuan_folder = f"static/uploads/kemajuan/{perguruan_tinggi}/{nama_ketua}"
        
        if not os.path.exists(kemajuan_folder):
            print(f"ℹ️ Folder kemajuan tidak ditemukan: {kemajuan_folder}")
            return
        
        # Cari dan hapus file laporan kemajuan yang terkait dengan proposal ini
        files_deleted = 0
        for filename in os.listdir(kemajuan_folder):
            file_path = os.path.join(kemajuan_folder, filename)
            
            # Cek apakah file terkait dengan proposal ini berdasarkan timestamp
            # File laporan kemajuan bisa memiliki format: nota_laporan_kemajuan_[kategori]_[nomor]_[timestamp].pdf
            if filename.startswith('nota_laporan_kemajuan_') and filename.endswith('.pdf'):
                # Ekstrak timestamp dari filename untuk mencocokkan dengan tanggal_buat proposal
                try:
                    # Ambil timestamp dari filename (format: nota_laporan_kemajuan_Isp_6_20250720_133806.pdf)
                    parts = filename.split('_')
                    if len(parts) >= 6:
                        file_date = parts[-2]  # 20250720
                        
                        # Konversi tanggal_buat proposal ke format yang sama
                        proposal_date = tanggal_buat.strftime('%Y%m%d')
                        
                        # Jika tanggal cocok, hapus file
                        if file_date == proposal_date:
                            os.remove(file_path)
                            print(f"✅ File laporan kemajuan berhasil dihapus: {file_path}")
                            files_deleted += 1
                except Exception as e:
                    print(f"❌ Error saat memproses file {filename}: {e}")
        
        # Cek juga folder kemajuan dengan judul_usaha sebagai perguruan_tinggi (fallback)
        kemajuan_folder_fallback = f"static/uploads/kemajuan/{judul_usaha}/{nama_ketua}"
        if os.path.exists(kemajuan_folder_fallback) and kemajuan_folder_fallback != kemajuan_folder:
            print(f"🔍 Mencari file di folder fallback: {kemajuan_folder_fallback}")
            for filename in os.listdir(kemajuan_folder_fallback):
                file_path = os.path.join(kemajuan_folder_fallback, filename)
                
                if filename.startswith('nota_laporan_kemajuan_') and filename.endswith('.pdf'):
                    try:
                        parts = filename.split('_')
                        if len(parts) >= 6:
                            file_date = parts[-2]
                            proposal_date = tanggal_buat.strftime('%Y%m%d')
                            
                            if file_date == proposal_date:
                                os.remove(file_path)
                                print(f"✅ File laporan kemajuan (fallback) berhasil dihapus: {file_path}")
                                files_deleted += 1
                    except Exception as e:
                        print(f"❌ Error saat memproses file fallback {filename}: {e}")
            
            # Hapus folder fallback jika kosong
            safe_remove_empty_folder(kemajuan_folder_fallback, "folder kemajuan fallback")
        
        # Hapus folder kemajuan jika kosong
        safe_remove_empty_folder(kemajuan_folder, "folder kemajuan")
        
        if files_deleted == 0:
            print(f"ℹ️ Tidak ada file laporan kemajuan yang ditemukan untuk proposal {proposal_id}")
        else:
            print(f"✅ Berhasil menghapus {files_deleted} file laporan kemajuan untuk proposal {proposal_id}")
            
    except Exception as e:
        print(f"❌ Error dalam fungsi hapus_file_laporan_kemajuan: {str(e)}")

# Fungsi untuk menghapus file laporan akhir yang terkait dengan laporan kemajuan tertentu
def hapus_file_laporan_akhir_terkait(proposal_id, nim, kegiatan_utama, kegiatan, nama_barang):
    """
    Menghapus file laporan akhir yang terkait dengan laporan kemajuan tertentu
    berdasarkan kegiatan_utama, kegiatan, dan nama_barang
    """
    try:
        # Ambil data proposal untuk mendapatkan perguruan tinggi, nama_ketua, judul_usaha
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        cursor.execute('''
            SELECT p.judul_usaha, m.perguruan_tinggi, m.nama_ketua
            FROM proposal p 
            INNER JOIN mahasiswa m ON p.nim = m.nim 
            WHERE p.id = %s
        ''', (proposal_id,))
        proposal_data = cursor.fetchone()
        cursor.close()
        
        if not proposal_data:
            print(f"⚠️ Data proposal tidak ditemukan untuk ID: {proposal_id}")
            return
        
        perguruan_tinggi = proposal_data['perguruan_tinggi']
        nama_ketua = proposal_data['nama_ketua']
        judul_usaha = proposal_data['judul_usaha']
        
        # Bersihkan judul usaha untuk nama folder
        judul_usaha_clean = judul_usaha.replace(' ', '_').replace('/', '_').replace('\\', '_').replace(':', '_')
        
        files_deleted = 0
        
        # 1. HAPUS FILE DARI STRUKTUR FOLDER BARU: static/uploads/Laporan Akhir/[judul_usaha]/
        laporan_akhir_folder_baru = f"static/uploads/Laporan Akhir/{judul_usaha_clean}"
        
        if os.path.exists(laporan_akhir_folder_baru):
            print(f"🔍 Mencari file laporan akhir terkait di folder baru: {laporan_akhir_folder_baru}")
            
            for filename in os.listdir(laporan_akhir_folder_baru):
                file_path = os.path.join(laporan_akhir_folder_baru, filename)
                
                # Cek file dengan format baru: Bukti_[judul_usaha]_[id]_laporan_akhir.pdf
                if filename.startswith('Bukti_') and filename.endswith('_laporan_akhir.pdf'):
                    try:
                        # Ekstrak ID dari filename: Bukti_Bakso_Enak_5_laporan_akhir.pdf
                        parts = filename.split('_')
                        if len(parts) >= 4:
                            # Cari ID laporan (biasanya setelah judul usaha)
                            for i, part in enumerate(parts):
                                if part.isdigit() and i < len(parts) - 2:  # Pastikan bukan bagian akhir
                                    file_id = int(part)
                                    if file_id == proposal_id:
                                        os.remove(file_path)
                                        print(f"✅ File laporan akhir terkait (struktur baru) berhasil dihapus: {file_path}")
                                        files_deleted += 1
                                        break
                    except Exception as e:
                        print(f"❌ Error saat memproses file baru {filename}: {e}")
        
        # 2. HAPUS FILE DARI STRUKTUR FOLDER LAMA: static/uploads/proposal/[perguruan_tinggi]/[nama_ketua]/laporan_akhir/
        laporan_akhir_folder_lama = f"static/uploads/proposal/{perguruan_tinggi}/{nama_ketua}/laporan_akhir"
        
        if os.path.exists(laporan_akhir_folder_lama):
            print(f"🔍 Mencari file laporan akhir terkait di folder lama: {laporan_akhir_folder_lama}")
            
            for filename in os.listdir(laporan_akhir_folder_lama):
                file_path = os.path.join(laporan_akhir_folder_lama, filename)
                
                # Cek file dengan format lama: nota_laporan_akhir_[kategori]_[nomor]_[timestamp].pdf atau bukti_laporan_akhir_[kategori]_[nomor]_[timestamp].pdf
                if (filename.startswith('nota_laporan_akhir_') or filename.startswith('bukti_laporan_akhir_')) and filename.endswith('.pdf'):
                    try:
                        # Ekstrak informasi dari filename untuk mencocokkan dengan kegiatan
                        parts = filename.split('_')
                        if len(parts) >= 6:
                            # Cek apakah file terkait dengan kegiatan yang sama
                            # Format: nota_laporan_akhir_[kegiatan]_[nomor]_[timestamp].pdf
                            kegiatan_from_filename = parts[3] if len(parts) > 3 else ''
                            
                            if kegiatan_from_filename.lower() in kegiatan.lower() or kegiatan.lower() in kegiatan_from_filename.lower():
                                os.remove(file_path)
                                print(f"✅ File laporan akhir terkait (struktur lama) berhasil dihapus: {file_path}")
                                files_deleted += 1
                    except Exception as e:
                        print(f"❌ Error saat memproses file lama {filename}: {e}")
        
        print(f"✅ Total {files_deleted} file laporan akhir terkait berhasil dihapus")
        
    except Exception as e:
        print(f"❌ Error dalam fungsi hapus_file_laporan_akhir_terkait: {str(e)}")

# Fungsi untuk menghapus file laporan akhir dari folder proposal dan Laporan Akhir
def hapus_file_laporan_akhir(proposal_id, nim):
    """
    Menghapus file laporan akhir dari folder:
    1. static/uploads/Laporan Akhir/[judul_usaha]/ (struktur baru)
    2. static/uploads/proposal/[perguruan_tinggi]/[nama_ketua]/laporan_akhir/ (struktur lama)
    berdasarkan proposal_id dan nim
    """
    try:
        # Ambil data proposal untuk mendapatkan perguruan tinggi, nama_ketua, judul_usaha, dan tanggal_buat
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        cursor.execute('''
            SELECT p.tanggal_buat, p.judul_usaha, m.perguruan_tinggi, m.nama_ketua
            FROM proposal p 
            INNER JOIN mahasiswa m ON p.nim = m.nim 
            WHERE p.id = %s
        ''', (proposal_id,))
        proposal_data = cursor.fetchone()
        cursor.close()
        
        if not proposal_data:
            print(f"⚠️ Data proposal tidak ditemukan untuk ID: {proposal_id}")
            return
        
        perguruan_tinggi = proposal_data['perguruan_tinggi']
        nama_ketua = proposal_data['nama_ketua']
        judul_usaha = proposal_data['judul_usaha']
        tanggal_buat = proposal_data['tanggal_buat']
        
        # Bersihkan judul usaha untuk nama folder
        judul_usaha_clean = judul_usaha.replace(' ', '_').replace('/', '_').replace('\\', '_').replace(':', '_')
        
        files_deleted = 0
        
        # 1. HAPUS FILE DARI STRUKTUR FOLDER BARU: static/uploads/Laporan Akhir/[judul_usaha]/
        laporan_akhir_folder_baru = f"static/uploads/Laporan Akhir/{judul_usaha_clean}"
        
        if os.path.exists(laporan_akhir_folder_baru):
            print(f"🔍 Mencari file laporan akhir di folder baru: {laporan_akhir_folder_baru}")
            folder_files_deleted = 0
            
            for filename in os.listdir(laporan_akhir_folder_baru):
                file_path = os.path.join(laporan_akhir_folder_baru, filename)
                
                # Cek file dengan format baru: Bukti_[judul_usaha]_[id]_laporan_akhir.pdf
                if filename.startswith('Bukti_') and filename.endswith('_laporan_akhir.pdf'):
                    try:
                        # Ekstrak ID dari filename: Bukti_Bakso_Enak_5_laporan_akhir.pdf
                        parts = filename.split('_')
                        if len(parts) >= 4:
                            # Cari ID laporan (biasanya setelah judul usaha)
                            for i, part in enumerate(parts):
                                if part.isdigit() and i < len(parts) - 2:  # Pastikan bukan bagian akhir
                                    file_id = int(part)
                                    if file_id == proposal_id:
                                        os.remove(file_path)
                                        print(f"✅ File laporan akhir (struktur baru) berhasil dihapus: {file_path}")
                                        files_deleted += 1
                                        folder_files_deleted += 1
                                        break
                    except Exception as e:
                        print(f"❌ Error saat memproses file baru {filename}: {e}")
            
            # Hapus folder baru jika kosong
            if folder_files_deleted > 0 and os.path.exists(laporan_akhir_folder_baru) and not os.listdir(laporan_akhir_folder_baru):
                try:
                    os.rmdir(laporan_akhir_folder_baru)
                    print(f"✅ Folder laporan akhir baru kosong berhasil dihapus: {laporan_akhir_folder_baru}")
                    
                    # Cek parent folder (Laporan Akhir) jika kosong
                    parent_folder = os.path.dirname(laporan_akhir_folder_baru)
                    if os.path.exists(parent_folder) and not os.listdir(parent_folder):
                        os.rmdir(parent_folder)
                        print(f"✅ Parent folder Laporan Akhir kosong berhasil dihapus: {parent_folder}")
                except Exception as e:
                    print(f"❌ Error saat menghapus folder baru {laporan_akhir_folder_baru}: {e}")
        
        # 2. HAPUS FILE DARI STRUKTUR FOLDER LAMA: static/uploads/proposal/[perguruan_tinggi]/[nama_ketua]/laporan_akhir/
        laporan_akhir_folder_lama = f"static/uploads/proposal/{perguruan_tinggi}/{nama_ketua}/laporan_akhir"
        
        if os.path.exists(laporan_akhir_folder_lama):
            print(f"🔍 Mencari file laporan akhir di folder lama: {laporan_akhir_folder_lama}")
            folder_files_deleted = 0
            
            for filename in os.listdir(laporan_akhir_folder_lama):
                file_path = os.path.join(laporan_akhir_folder_lama, filename)
                
                # Cek file dengan format lama: nota_laporan_akhir_[kategori]_[nomor]_[timestamp].pdf atau bukti_laporan_akhir_[kategori]_[nomor]_[timestamp].pdf
                if (filename.startswith('nota_laporan_akhir_') or filename.startswith('bukti_laporan_akhir_')) and filename.endswith('.pdf'):
                    try:
                        # Ekstrak timestamp dari filename untuk mencocokkan dengan tanggal_buat proposal
                        parts = filename.split('_')
                        if len(parts) >= 6:
                            file_date = parts[-2]  # 20250720
                            proposal_date = tanggal_buat.strftime('%Y%m%d')
                            
                            if file_date == proposal_date:
                                os.remove(file_path)
                                print(f"✅ File laporan akhir (struktur lama) berhasil dihapus: {file_path}")
                                files_deleted += 1
                                folder_files_deleted += 1
                    except Exception as e:
                        print(f"❌ Error saat memproses file lama {filename}: {e}")
            
            # Hapus folder lama jika kosong
            if folder_files_deleted > 0 and os.path.exists(laporan_akhir_folder_lama) and not os.listdir(laporan_akhir_folder_lama):
                try:
                    os.rmdir(laporan_akhir_folder_lama)
                    print(f"✅ Folder laporan akhir lama kosong berhasil dihapus: {laporan_akhir_folder_lama}")
                except Exception as e:
                    print(f"❌ Error saat menghapus folder lama {laporan_akhir_folder_lama}: {e}")
        
        # 3. HAPUS FILE DARI FOLDER FALLBACK: static/uploads/proposal/[judul_usaha]/[nama_ketua]/laporan_akhir/
        laporan_akhir_folder_fallback = f"static/uploads/proposal/{judul_usaha}/{nama_ketua}/laporan_akhir"
        
        if os.path.exists(laporan_akhir_folder_fallback) and laporan_akhir_folder_fallback != laporan_akhir_folder_lama:
            print(f"🔍 Mencari file di folder fallback: {laporan_akhir_folder_fallback}")
            folder_files_deleted = 0
            
            for filename in os.listdir(laporan_akhir_folder_fallback):
                file_path = os.path.join(laporan_akhir_folder_fallback, filename)
                
                if (filename.startswith('nota_laporan_akhir_') or filename.startswith('bukti_laporan_akhir_')) and filename.endswith('.pdf'):
                    try:
                        parts = filename.split('_')
                        if len(parts) >= 6:
                            file_date = parts[-2]
                            proposal_date = tanggal_buat.strftime('%Y%m%d')
                            
                            if file_date == proposal_date:
                                os.remove(file_path)
                                print(f"✅ File laporan akhir (fallback) berhasil dihapus: {file_path}")
                                files_deleted += 1
                                folder_files_deleted += 1
                    except Exception as e:
                        print(f"❌ Error saat memproses file fallback {filename}: {e}")
            
            # Hapus folder fallback jika kosong
            safe_remove_empty_folder(laporan_akhir_folder_fallback, "folder laporan akhir fallback")
        
        if files_deleted == 0:
            print(f"ℹ️ Tidak ada file laporan akhir yang ditemukan untuk proposal {proposal_id}")
        else:
            print(f"✅ Berhasil menghapus {files_deleted} file laporan akhir untuk proposal {proposal_id}")
        
    except Exception as e:
        print(f"❌ Error dalam fungsi hapus_file_laporan_akhir: {str(e)}")

# Fungsi komprehensif untuk menghapus semua file terkait proposal
def hapus_semua_file_proposal(proposal_id, nim):
    """
    Menghapus semua file terkait proposal dari berbagai folder
    berdasarkan proposal_id dan nim
    """
    try:
        # Ambil data proposal
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        cursor.execute('''
            SELECT p.tanggal_buat, p.judul_usaha, m.perguruan_tinggi, m.nama_ketua
            FROM proposal p 
            INNER JOIN mahasiswa m ON p.nim = m.nim 
            WHERE p.id = %s
        ''', (proposal_id,))
        proposal_data = cursor.fetchone()
        cursor.close()
        
        if not proposal_data:
            print(f"⚠️ Data proposal tidak ditemukan untuk ID: {proposal_id}")
            return
        
        perguruan_tinggi = proposal_data['perguruan_tinggi']
        nama_ketua = proposal_data['nama_ketua']
        judul_usaha = proposal_data['judul_usaha']
        tanggal_buat = proposal_data['tanggal_buat']
        proposal_date = tanggal_buat.strftime('%Y%m%d')
        
        total_files_deleted = 0
        
        # 1. Hapus file laporan kemajuan dari folder kemajuan
        kemajuan_folders = [
            f"static/uploads/kemajuan/{perguruan_tinggi}/{nama_ketua}",
            f"static/uploads/kemajuan/{judul_usaha}/{nama_ketua}"
        ]
        
        for kemajuan_folder in kemajuan_folders:
            if os.path.exists(kemajuan_folder):
                print(f"🔍 Mencari file laporan kemajuan di: {kemajuan_folder}")
                files_deleted = 0
                
                for filename in os.listdir(kemajuan_folder):
                    file_path = os.path.join(kemajuan_folder, filename)
                    
                    # Cek file laporan kemajuan
                    if filename.startswith('nota_laporan_kemajuan_') and filename.endswith('.pdf'):
                        try:
                            parts = filename.split('_')
                            if len(parts) >= 6:
                                file_date = parts[-2]
                                if file_date == proposal_date:
                                    os.remove(file_path)
                                    print(f"✅ File laporan kemajuan berhasil dihapus: {file_path}")
                                    files_deleted += 1
                                    total_files_deleted += 1
                        except Exception as e:
                            print(f"❌ Error saat memproses file {filename}: {e}")
                
                                        # Hapus folder jika kosong
                safe_remove_empty_folder(kemajuan_folder, "folder kemajuan")
        
        # 2. Hapus file laporan akhir dari folder proposal (struktur lama) dan Laporan Akhir (struktur baru)
        # Bersihkan judul usaha untuk nama folder
        judul_usaha_clean = judul_usaha.replace(' ', '_').replace('/', '_').replace('\\', '_').replace(':', '_')
        
        # Struktur folder lama
        laporan_akhir_folders_lama = [
            f"static/uploads/proposal/{perguruan_tinggi}/{nama_ketua}/laporan_akhir",
            f"static/uploads/proposal/{judul_usaha}/{nama_ketua}/laporan_akhir"
        ]
        
        # Struktur folder baru
        laporan_akhir_folder_baru = f"static/uploads/Laporan Akhir/{judul_usaha_clean}"
        
        # Hapus dari folder lama
        for laporan_akhir_folder in laporan_akhir_folders_lama:
            if os.path.exists(laporan_akhir_folder):
                print(f"🔍 Mencari file laporan akhir di folder lama: {laporan_akhir_folder}")
                files_deleted = 0
                
                for filename in os.listdir(laporan_akhir_folder):
                    file_path = os.path.join(laporan_akhir_folder, filename)
                    
                    # Cek file laporan akhir (nota dan bukti) - format lama
                    if (filename.startswith('nota_laporan_akhir_') or filename.startswith('bukti_laporan_akhir_')) and filename.endswith('.pdf'):
                        try:
                            parts = filename.split('_')
                            if len(parts) >= 6:
                                file_date = parts[-2]
                                if file_date == proposal_date:
                                    os.remove(file_path)
                                    print(f"✅ File laporan akhir (struktur lama) berhasil dihapus: {file_path}")
                                    files_deleted += 1
                                    total_files_deleted += 1
                        except Exception as e:
                            print(f"❌ Error saat memproses file lama {filename}: {e}")
                
                # Hapus folder jika kosong
                safe_remove_empty_folder(laporan_akhir_folder, "folder laporan akhir lama")
        
        # Hapus dari folder baru
        if os.path.exists(laporan_akhir_folder_baru):
            print(f"🔍 Mencari file laporan akhir di folder baru: {laporan_akhir_folder_baru}")
            files_deleted = 0
            
            for filename in os.listdir(laporan_akhir_folder_baru):
                file_path = os.path.join(laporan_akhir_folder_baru, filename)
                
                # Cek file dengan format baru: Bukti_[judul_usaha]_[id]_laporan_akhir.pdf
                if filename.startswith('Bukti_') and filename.endswith('_laporan_akhir.pdf'):
                    try:
                        # Ekstrak ID dari filename: Bukti_Bakso_Enak_5_laporan_akhir.pdf
                        parts = filename.split('_')
                        if len(parts) >= 4:
                            # Cari ID laporan (biasanya setelah judul usaha)
                            for i, part in enumerate(parts):
                                if part.isdigit() and i < len(parts) - 2:  # Pastikan bukan bagian akhir
                                    file_id = int(part)
                                    if file_id == proposal_id:
                                        os.remove(file_path)
                                        print(f"✅ File laporan akhir (struktur baru) berhasil dihapus: {file_path}")
                                        files_deleted += 1
                                        total_files_deleted += 1
                                        break
                    except Exception as e:
                        print(f"❌ Error saat memproses file baru {filename}: {e}")
            
            # Hapus folder baru jika kosong
            safe_remove_empty_folder(laporan_akhir_folder_baru, "folder laporan akhir baru")
        
        
        print(f"🎯 Total file yang berhasil dihapus: {total_files_deleted}")
        return total_files_deleted
        
    except Exception as e:
        print(f"❌ Error dalam fungsi hapus_semua_file_proposal: {str(e)}")
        return 0

# Fungsi untuk logging aktivitas pembimbing
def log_pembimbing_activity(pembimbing_id, nip, nama_pembimbing, jenis_aktivitas, modul, detail_modul=None, deskripsi="", data_lama=None, data_baru=None, target_mahasiswa_id=None, target_proposal_id=None):
    """
    Mencatat aktivitas pembimbing ke database
    
    Args:
        pembimbing_id: ID pembimbing
        nip: NIP pembimbing
        nama_pembimbing: Nama lengkap pembimbing
        jenis_aktivitas: login, logout, tambah, edit, hapus, view, konfirmasi, setuju, tolak, revisi
        modul: proposal, anggaran_awal, anggaran_bertumbuh, laporan_kemajuan, laporan_akhir, monitoring, sistem
        detail_modul: Detail spesifik (proposal_id_6, anggaran_id_12, dll)
        deskripsi: Deskripsi aktivitas
        data_lama: Data sebelum perubahan (JSON)
        data_baru: Data setelah perubahan (JSON)
        target_mahasiswa_id: ID mahasiswa yang terkait (opsional)
        target_proposal_id: ID proposal yang terkait (opsional)
    """
    try:
        if not hasattr(mysql, 'connection') or mysql.connection is None:
            return False
            
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        
        # Get IP address and user agent
        ip_address = request.environ.get('HTTP_X_FORWARDED_FOR', request.environ.get('REMOTE_ADDR'))
        user_agent = request.environ.get('HTTP_USER_AGENT')
        
        # Convert data to JSON if needed
        data_lama_json = json.dumps(convert_decimal_for_json(data_lama)) if data_lama else None
        data_baru_json = json.dumps(convert_decimal_for_json(data_baru)) if data_baru else None
        
        cursor.execute('''
            INSERT INTO log_aktivitas_pembimbing 
            (pembimbing_id, nip, nama_pembimbing, jenis_aktivitas, modul, detail_modul, 
             deskripsi, data_lama, data_baru, target_mahasiswa_id, target_proposal_id, 
             ip_address, user_agent)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        ''', (pembimbing_id, nip, nama_pembimbing, jenis_aktivitas, modul, detail_modul,
              deskripsi, data_lama_json, data_baru_json, target_mahasiswa_id, target_proposal_id,
              ip_address, user_agent))
        
        # Update last_activity di session_pembimbing yang aktif
        cursor.execute('''
            UPDATE session_pembimbing 
            SET last_activity = NOW()
            WHERE pembimbing_id = %s AND status = 'active'
        ''', (pembimbing_id,))
        
        mysql.connection.commit()
        cursor.close()
        return True
        
    except Exception as e:
        print(f"Error logging pembimbing activity: {str(e)}")
        # Pastikan error logging tidak menghalangi proses utama
        try:
            if hasattr(mysql, 'connection') and mysql.connection:
                mysql.connection.rollback()
        except:
            pass
        return False

# Fungsi untuk tracking session pembimbing
def start_pembimbing_session(pembimbing_id, nip, nama_pembimbing):
    """
    Memulai session tracking untuk pembimbing
    """
    try:
        print(f"DEBUG: Starting session for pembimbing_id={pembimbing_id}, nip={nip}, nama={nama_pembimbing}")
        
        if not hasattr(mysql, 'connection') or mysql.connection is None:
            print("DEBUG: MySQL connection not available")
            return False
            
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        
        # Tutup session aktif sebelumnya jika ada
        cursor.execute('''
            UPDATE session_pembimbing 
            SET status = 'ended', logout_time = NOW(), 
                durasi_detik = TIMESTAMPDIFF(SECOND, login_time, NOW())
            WHERE pembimbing_id = %s AND status = 'active'
        ''', (pembimbing_id,))
        
        # Buat session baru
        ip_address = request.environ.get('HTTP_X_FORWARDED_FOR', request.environ.get('REMOTE_ADDR'))
        user_agent = request.environ.get('HTTP_USER_AGENT')
        
        # Generate session ID unik jika belum ada
        import uuid
        if 'session_id' not in session or not session['session_id']:
            session['session_id'] = str(uuid.uuid4())
        session_id = session['session_id']
        
        print(f"DEBUG: Session ID: {session_id}")
        print(f"DEBUG: IP Address: {ip_address}")
        print(f"DEBUG: User Agent: {user_agent}")
        
        cursor.execute('''
            INSERT INTO session_pembimbing 
            (pembimbing_id, nip, nama_pembimbing, session_id, ip_address, user_agent)
            VALUES (%s, %s, %s, %s, %s, %s)
        ''', (pembimbing_id, nip, nama_pembimbing, session_id, ip_address, user_agent))
        
        mysql.connection.commit()
        print(f"DEBUG: Pembimbing session inserted successfully")
        cursor.close()
        
        # Log aktivitas login
        log_pembimbing_activity(pembimbing_id, nip, nama_pembimbing, 'login', 'sistem', 'authentication', 'Login ke sistem AYMP')
        
        return True
        
    except Exception as e:
        print(f"Error starting pembimbing session: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

# Fungsi untuk mengakhiri session pembimbing
def end_pembimbing_session(pembimbing_id):
    """
    Mengakhiri session tracking untuk pembimbing
    """
    try:
        if not hasattr(mysql, 'connection') or mysql.connection is None:
            return False
            
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        
        # Update session aktif
        cursor.execute('''
            UPDATE session_pembimbing 
            SET status = 'ended', logout_time = NOW(), 
                durasi_detik = TIMESTAMPDIFF(SECOND, login_time, NOW())
            WHERE pembimbing_id = %s AND status = 'active'
        ''', (pembimbing_id,))
        
        mysql.connection.commit()
        cursor.close()
        return True
        
    except Exception as e:
        print(f"Error ending pembimbing session: {str(e)}")
        return False

# Fungsi untuk mendapatkan informasi pembimbing dari session
def get_pembimbing_info_from_session():
    """
    Mengambil informasi pembimbing dari session aktif
    """
    if 'user_type' not in session or session['user_type'] != 'pembimbing':
        return None
        
    if not hasattr(mysql, 'connection') or mysql.connection is None:
        return None
        
    try:
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        cursor.execute('SELECT id, nip, nama FROM pembimbing WHERE nip = %s', (session['nip'],))
        pembimbing = cursor.fetchone()
        cursor.close()
        return pembimbing
    except:
        return None

def create_upload_path(perguruan_tinggi, nama_ketua):
    """Membuat path folder upload berdasarkan perguruan tinggi dan nama ketua"""
    # Bersihkan nama folder dari karakter yang tidak aman
    safe_perguruan_tinggi = re.sub(r'[^\w\s-]', '', perguruan_tinggi).strip()
    safe_nama_ketua = re.sub(r'[^\w\s-]', '', nama_ketua).strip()
    
    # Ganti spasi dengan underscore
    safe_perguruan_tinggi = safe_perguruan_tinggi.replace(' ', '_')
    safe_nama_ketua = safe_nama_ketua.replace(' ', '_')
    
    # Buat path lengkap
    upload_path = os.path.join(app.config['UPLOAD_FOLDER'], safe_perguruan_tinggi, safe_nama_ketua)
    
    # Buat folder jika belum ada
    os.makedirs(upload_path, exist_ok=True)
    
    return upload_path

def create_standardized_file_path(file_type, perguruan_tinggi, nama_ketua, judul_usaha, original_filename, subfolder=None):
    """
    Membuat path file yang standar dan konsisten
    
    Args:
        file_type: 'proposal', 'laporan_kemajuan', 'laporan_akhir', 
        perguruan_tinggi: Nama perguruan tinggi
        nama_ketua: Nama ketua tim
        judul_usaha: Judul usaha
        original_filename: Nama file asli
        subfolder: Subfolder tambahan (opsional)
    
    Returns:
        tuple: (physical_path, database_path)
    """
    # Bersihkan nama-nama untuk keamanan
    safe_perguruan_tinggi = re.sub(r'[^\w\s-]', '', perguruan_tinggi).strip().replace(' ', '_')
    safe_nama_ketua = re.sub(r'[^\w\s-]', '', nama_ketua).strip().replace(' ', '_')
    safe_judul = re.sub(r'[^\w\s-]', '', judul_usaha).strip().replace(' ', '_')
    
    # Dapatkan ekstensi file
    file_extension = original_filename.rsplit('.', 1)[1].lower() if '.' in original_filename else 'pdf'
    
    # Buat nama file yang standar
    if file_type == 'proposal':
        filename = f"{safe_judul}_{safe_nama_ketua}_proposal.{file_extension}"
        base_folder = 'Proposal'
    elif file_type == 'laporan_kemajuan':
        filename = f"Laporan_Kemajuan_{safe_judul}_{safe_nama_ketua}.{file_extension}"
        base_folder = 'Laporan Kemajuan'
        subfolder = 'Laporan' if not subfolder else subfolder
    elif file_type == 'laporan_akhir':
        filename = f"Laporan_Akhir_{safe_judul}_{safe_nama_ketua}.{file_extension}"
        base_folder = 'Laporan Kemajuan'
        subfolder = 'Laporan' if not subfolder else subfolder

    else:
        filename = f"{file_type}_{safe_judul}_{safe_nama_ketua}.{file_extension}"
        base_folder = 'Misc'
    
    # Buat path fisik untuk penyimpanan
    if subfolder:
        physical_path = os.path.join('static', 'uploads', base_folder, subfolder, filename)
    else:
        physical_path = os.path.join('static', 'uploads', base_folder, filename)
    
    # Buat path untuk database (relatif terhadap static/uploads)
    if subfolder:
        db_path = f"uploads/{base_folder}/{subfolder}/{filename}"
    else:
        db_path = f"uploads/{base_folder}/{filename}"
    
    # Buat direktori jika belum ada
    os.makedirs(os.path.dirname(physical_path), exist_ok=True)
    
    return physical_path, db_path

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        # Validasi form
        username = request.form.get('nim')
        password = request.form.get('password')
        if not hasattr(mysql, 'connection') or mysql.connection is None:
            return jsonify({'success': False, 'message': 'Koneksi ke database gagal. Cek konfigurasi database!'})
        try:
            cursor = mysql.connection.cursor()
        except Exception as e:
            return jsonify({'success': False, 'message': f'Koneksi database gagal: {e}'})
        username_found = False
        password_correct = False
        # Cek admin - bisa login dengan nama atau NIP
        cursor.execute('SELECT * FROM admin WHERE nama = %s OR nip = %s', (username, username))
        admin = cursor.fetchone()
        if admin:
            username_found = True
            if admin[4] and admin[4].strip():
                try:
                    if check_password_hash(admin[4], password):
                        password_correct = True
                        session['user_type'] = 'admin'
                        session['nama'] = admin[1]
                        session['nip'] = admin[2]
                        cursor.close()
                        return jsonify({'success': True, 'redirect': url_for('admin.das_admin'), 'message': 'Login admin berhasil!'})
                except ValueError:
                    cursor.close()
                    return jsonify({'success': False, 'message': 'Password admin tidak valid. Silakan hubungi administrator.'})
        
        # Cek pembimbing - bisa login dengan nama atau NIP
        if not username_found:
            cursor.execute('SELECT * FROM pembimbing WHERE nama = %s OR nip = %s', (username, username))
            pembimbing = cursor.fetchone()
            if pembimbing:
                username_found = True
                if pembimbing[4] and pembimbing[4].strip():
                    try:
                        if check_password_hash(pembimbing[4], password):
                            password_correct = True
                            session['user_type'] = 'pembimbing'
                            session['nama'] = pembimbing[1]
                            session['nip'] = pembimbing[2]
                            
                            # Start session tracking untuk pembimbing
                            try:
                                start_pembimbing_session(pembimbing[0], pembimbing[2], pembimbing[1])
                            except Exception as e:
                                print(f"Warning: Error starting pembimbing session: {str(e)}")
                            
                            cursor.close()
                            return jsonify({'success': True, 'redirect': url_for('pembimbing.das_pembimbing'), 'message': 'Login pembimbing berhasil!'})
                    except ValueError:
                        cursor.close()
                        return jsonify({'success': False, 'message': 'Password pembimbing tidak valid. Silakan hubungi administrator.'})
        # Cek reviewer
        if not username_found:
            cursor.execute('SELECT * FROM reviewer WHERE username = %s OR nama = %s', (username, username))
            reviewer = cursor.fetchone()
            if reviewer:
                username_found = True
                if reviewer[3] and reviewer[3].strip():
                    try:
                        if check_password_hash(reviewer[3], password):
                            password_correct = True
                            session['user_type'] = 'reviewer'
                            session['user_id'] = reviewer[0]
                            session['nama'] = reviewer[1]
                            session['username'] = reviewer[2]
                            
                            cursor.close()
                            return jsonify({'success': True, 'redirect': url_for('reviewer.dashboard'), 'message': 'Login reviewer berhasil!'})
                    except ValueError:
                        cursor.close()
                        return jsonify({'success': False, 'message': 'Password reviewer tidak valid. Silakan hubungi administrator.'})
        
        # Cek mahasiswa
        if not username_found:
            cursor.execute('SELECT * FROM mahasiswa WHERE nim = %s OR nama_ketua = %s', (username, username))
            mahasiswa = cursor.fetchone()
            if mahasiswa:
                username_found = True
                if mahasiswa[7] and mahasiswa[7].strip():
                    try:
                        if check_password_hash(mahasiswa[7], password):
                            password_correct = True
                            session['user_type'] = 'mahasiswa'
                            session['nim'] = mahasiswa[3]
                            session['nama'] = mahasiswa[4]
                            
                            # Start session tracking untuk mahasiswa
                            try:
                                start_mahasiswa_session(mahasiswa[0], mahasiswa[3], mahasiswa[4])
                            except Exception as e:
                                print(f"Warning: Error starting mahasiswa session: {str(e)}")
                            
                            cursor.close()
                            return jsonify({'success': True, 'redirect': url_for('mahasiswa.dashboard_mahasiswa'), 'message': 'Login mahasiswa berhasil!'})
                    except ValueError:
                        cursor.close()
                        return jsonify({'success': False, 'message': 'Password mahasiswa tidak valid. Silakan hubungi administrator.'})
        cursor.close()
        if not username_found:
            return jsonify({'success': False, 'message': 'Username/Nama/NIM/NIP tidak ditemukan!'})
        elif username_found and not password_correct:
            return jsonify({'success': False, 'message': 'Password yang Anda masukkan salah!'})
        else:
            return jsonify({'success': False, 'message': 'Username/Nama/NIM/NIP dan password salah!'})
    return render_template('logdaf/login.html')


@app.route('/daftar', methods=['GET', 'POST'])
def daftar():
    # Ambil data program studi untuk dropdown
    program_studi_list = []
    try:
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        cursor.execute('SELECT id, nama_program_studi FROM program_studi ORDER BY nama_program_studi')
        program_studi_list = cursor.fetchall()
        cursor.close()
    except Exception as e:
        flash(f'Error mengambil data program studi: {e}', 'danger')
        program_studi_list = []
    
    if request.method == 'POST':
        # Validasi form
        perguruan_tinggi = "Universitas Jenderal Achmad Yani Yogyakarta"
        program_studi = request.form.get('program_studi', '').strip()
        nim = request.form.get('nim', '').strip()
        nama_ketua = request.form.get('nama_ketua', '').strip().title()
        no_telp = request.form.get('no_telp', '').strip()
        email = request.form.get('email', '').strip().lower()
        password = request.form.get('password')
        ulangi_password = request.form.get('ulangi_password')
        status = 'proses'
        # Validasi sederhana
        if not all([perguruan_tinggi, program_studi, nim, nama_ketua, no_telp, email, password, ulangi_password]):
            flash('Semua field harus diisi!', 'danger')
            return render_template('logdaf/daftar.html', program_studi_list=program_studi_list)
        if password != ulangi_password:
            flash('Password dan Ulangi Password harus sama!', 'danger')
            return render_template('logdaf/daftar.html', program_studi_list=program_studi_list)
        if not hasattr(mysql, 'connection') or mysql.connection is None:
            flash('Koneksi ke database gagal. Cek konfigurasi database!', 'danger')
            return render_template('logdaf/daftar.html', program_studi_list=program_studi_list)
        try:
            cursor = mysql.connection.cursor()
        except Exception as e:
            flash(f'Koneksi database gagal: {e}', 'danger')
            return render_template('logdaf/daftar.html', program_studi_list=program_studi_list)
        # Cek duplikasi komprehensif
        from admin import check_comprehensive_duplicates
        is_duplicate, message, table_name = check_comprehensive_duplicates(
            cursor, 
            nama=nama_ketua, 
            nim=nim, 
            email=email
        )
        if is_duplicate:
            flash(message, 'danger')
            cursor.close()
            return render_template('logdaf/daftar.html', program_studi_list=program_studi_list)
        try:
            password_hash = generate_password_hash(password)
            cursor.execute('''
                INSERT INTO mahasiswa (perguruan_tinggi, program_studi, nim, nama_ketua, no_telp, email, password, status)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
            ''', (perguruan_tinggi, program_studi, nim, nama_ketua, no_telp, email, password_hash, status))
            mysql.connection.commit()
            flash('Pendaftaran berhasil!', 'success')
            return redirect(url_for('daftar', sukses=1))
        except MySQLdb.IntegrityError:
            flash('NIM atau Email sudah terdaftar!', 'danger')
            return render_template('logdaf/daftar.html', program_studi_list=program_studi_list)
        except Exception as e:
            flash(f'Error saat menyimpan data: {e}', 'danger')
            return render_template('logdaf/daftar.html', program_studi_list=program_studi_list)
        finally:
            cursor.close()
    if request.args.get('sukses') == '1':
        pass
    else:
        session.pop('_flashes', None)
    return render_template('logdaf/daftar.html', program_studi_list=program_studi_list)

# Route das_mahasiswa dihapus karena tidak digunakan
# Semua template menggunakan dashboard_mahasiswa

def group_anggaran_data(anggaran_data):
    grouped = {}
    for row in anggaran_data:
        keg_utama = row['kegiatan_utama']
        kegiatan = row['kegiatan']
        if keg_utama not in grouped:
            grouped[keg_utama] = {}
        if kegiatan not in grouped[keg_utama]:
            grouped[keg_utama][kegiatan] = []
        grouped[keg_utama][kegiatan].append(row)
    return grouped

def flatten_anggaran_data(anggaran_data):
    """Flatten data anggaran dengan informasi rowspan untuk tabel"""
    
    # Group by kegiatan_utama, kegiatan, penanggung_jawab
    flat_rows = []
    # Grouping
    grouped = defaultdict(lambda: defaultdict(lambda: defaultdict(list)))
    for row in anggaran_data:
        grouped[row['kegiatan_utama']][row['kegiatan']][row['penanggung_jawab']].append(row)
    
    # Flatten with rowspan info
    for kegiatan_utama, keg_dict in grouped.items():
        kegiatan_utama_rowspan = sum(
            sum(len(pj_rows) for pj_rows in kegiatan_dict.values())
            for kegiatan_dict in keg_dict.values()
        )
        first_utama = True
        for kegiatan, kegiatan_dict in keg_dict.items():
            kegiatan_rowspan = sum(len(pj_rows) for pj_rows in kegiatan_dict.values())
            first_kegiatan = True
            
            # Process each penanggung_jawab group
            for penanggung_jawab, pj_rows in kegiatan_dict.items():
                pj_rowspan = len(pj_rows)
                first_pj = True
                
                # Group by target_capaian untuk penanggung_jawab ini
                target_capaian_groups = {}
                for row in pj_rows:
                    target_capaian = row['target_capaian']
                    if target_capaian not in target_capaian_groups:
                        target_capaian_groups[target_capaian] = []
                    target_capaian_groups[target_capaian].append(row)
                
                # Process each target_capaian group
                for target_capaian, target_rows in target_capaian_groups.items():
                    target_rowspan = len(target_rows)
                    first_target = True
                    
                    for row in target_rows:
                        flat_row = dict(row)
                        flat_row['show_kegiatan_utama'] = first_utama
                        flat_row['kegiatan_utama_rowspan'] = kegiatan_utama_rowspan if first_utama else 0
                        flat_row['show_kegiatan'] = first_kegiatan
                        flat_row['kegiatan_rowspan'] = kegiatan_rowspan if first_kegiatan else 0
                        flat_row['show_penanggung_jawab'] = first_pj
                        flat_row['penanggung_jawab_rowspan'] = pj_rowspan if first_pj else 0
                        # Target capaian dengan rowspan berdasarkan grup yang sama
                        flat_row['show_target_capaian'] = first_target
                        flat_row['target_capaian_rowspan'] = target_rowspan if first_target else 0
                        flat_rows.append(flat_row)
                        first_utama = False
                        first_kegiatan = False
                        first_pj = False
                        first_target = False
    
    return flat_rows


def calculate_and_save_laba_rugi(proposal_id):
    """Fungsi untuk menghitung dan menyimpan data laba rugi"""
    if not hasattr(mysql, 'connection') or mysql.connection is None:
        print("Koneksi database tidak tersedia")
        return
        
    try:
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        
        # Ambil data produksi
        cursor.execute('''
            SELECT tanggal_produksi, nama_produk, total_biaya
            FROM produksi
            WHERE proposal_id = %s
            ORDER BY tanggal_produksi DESC
        ''', (proposal_id,))
        
        produksi_data = cursor.fetchall()
        
        # Ambil data biaya operasional dengan estimasi hari
        cursor.execute('''
            SELECT total_harga, estimasi_hari_habis, estimasi_aktif_digunakan
            FROM biaya_operasional
            WHERE proposal_id = %s
        ''', (proposal_id,))
        
        biaya_operasional_data = cursor.fetchall()
        
        # Hapus data laba rugi lama untuk proposal ini
        cursor.execute('DELETE FROM laba_rugi WHERE proposal_id = %s', (proposal_id,))
        
        # Hitung dan simpan data laba rugi untuk setiap produksi
        for produksi in produksi_data:
            # Hitung pendapatan untuk produk ini
            cursor.execute('''
                SELECT SUM(total) as pendapatan_produk
                FROM penjualan
                WHERE proposal_id = %s AND nama_produk = %s
            ''', (proposal_id, produksi['nama_produk']))
            
            penjualan_produk = cursor.fetchone()
            pendapatan_produk = penjualan_produk['pendapatan_produk'] or 0
            
            # Hitung biaya operasional berdasarkan rumus baru
            # biaya_operasional = (total_harga / estimasi_hari_habis) * estimasi_aktif_digunakan
            total_biaya_operasional_periode = 0
            for biaya in biaya_operasional_data:
                if biaya['estimasi_hari_habis'] > 0:  # Hindari division by zero
                    biaya_per_hari = biaya['total_harga'] / biaya['estimasi_hari_habis']
                    biaya_operasional_item = biaya_per_hari * biaya['estimasi_aktif_digunakan']
                    total_biaya_operasional_periode += biaya_operasional_item
            
            # Hitung laba rugi
            laba_kotor = pendapatan_produk - produksi['total_biaya']
            laba_bersih = laba_kotor - total_biaya_operasional_periode
            
            # Simpan ke tabel laba_rugi
            cursor.execute('''
                INSERT INTO laba_rugi (
                    proposal_id, tanggal_produksi, nama_produk, pendapatan, 
                    total_biaya_produksi, laba_rugi_kotor, biaya_operasional, laba_rugi_bersih
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
            ''', (
                proposal_id, produksi['tanggal_produksi'], produksi['nama_produk'],
                pendapatan_produk, produksi['total_biaya'], laba_kotor,
                total_biaya_operasional_periode, laba_bersih
            ))
        
        mysql.connection.commit()
        cursor.close()
        
    except Exception as e:
        print(f"Error calculating laba rugi: {str(e)}")

def update_arus_kas_otomatis(proposal_id, bulan_tahun):
    """
    Fungsi untuk otomatis update data arus kas ketika ada perubahan data
    Dipanggil setiap kali ada perubahan di: penjualan, produksi, biaya operasional, Biaya Lain-lain, alat produksi
    """
    try:
        if not hasattr(mysql, 'connection') or mysql.connection is None:
            return False
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        
        # Parse bulan yang dipilih
        if not bulan_tahun:
            return False
        
        # Parse bulan yang dipilih
        
        try:
            selected_date = datetime.strptime(bulan_tahun, '%Y-%m')
            start_date = selected_date.replace(day=1)
            end_date = (selected_date.replace(day=1) + timedelta(days=32)).replace(day=1) - timedelta(days=1)
        except ValueError:
            return False
        
        # 1. Total Penerimaan Kas dari Penjualan (PER BULAN YANG DIPILIH)
        cursor.execute("""
            SELECT COALESCE(SUM(total), 0) as total_penjualan
            FROM penjualan 
            WHERE proposal_id = %s 
            AND DATE(tanggal_penjualan) BETWEEN %s AND %s
        """, (proposal_id, start_date.date(), end_date.date()))
        total_penjualan = cursor.fetchone()['total_penjualan']
        
        # 2. Total Biaya Produksi (Supplier) - PER BULAN YANG DIPILIH
        cursor.execute("""
            SELECT COALESCE(SUM(total_biaya), 0) as total_biaya_produksi
            FROM produksi 
            WHERE proposal_id = %s 
            AND DATE(tanggal_produksi) BETWEEN %s AND %s
        """, (proposal_id, start_date.date(), end_date.date()))
        total_biaya_produksi = cursor.fetchone()['total_biaya_produksi']
        
        # 3. Total Biaya Operasional - PER BULAN YANG DIPILIH
        cursor.execute("""
            SELECT COALESCE(SUM(total_harga), 0) as total_biaya_operasional
            FROM biaya_operasional 
            WHERE proposal_id = %s 
            AND DATE(tanggal_beli) BETWEEN %s AND %s
        """, (proposal_id, start_date.date(), end_date.date()))
        total_biaya_operasional = cursor.fetchone()['total_biaya_operasional']
        
        # 4. Total Biaya Lain-lain - PER BULAN YANG DIPILIH
        cursor.execute("""
            SELECT COALESCE(SUM(total_harga), 0) as total_biaya_non_operasional
            FROM biaya_non_operasional 
            WHERE proposal_id = %s 
            AND DATE(tanggal_transaksi) BETWEEN %s AND %s
        """, (proposal_id, start_date.date(), end_date.date()))
        total_biaya_non_operasional = cursor.fetchone()['total_biaya_non_operasional']
        
        # 5. Total Harga Jual Alat Produksi - PER BULAN YANG DIPILIH
        cursor.execute("""
            SELECT COALESCE(SUM(harga_jual), 0) as total_harga_jual_alat
            FROM alat_produksi 
            WHERE proposal_id = %s 
            AND DATE(tanggal_beli) BETWEEN %s AND %s
            AND harga_jual > 0
        """, (proposal_id, start_date.date(), end_date.date()))
        total_harga_jual_alat = cursor.fetchone()['total_harga_jual_alat']
        
        # 6. Total Harga Alat Produksi (Pembelian) - PER BULAN YANG DIPILIH
        cursor.execute("""
            SELECT COALESCE(SUM(total_alat_produksi), 0) as total_harga_alat
            FROM alat_produksi 
            WHERE proposal_id = %s 
            AND DATE(tanggal_beli) BETWEEN %s AND %s
        """, (proposal_id, start_date.date(), end_date.date()))
        total_harga_alat = cursor.fetchone()['total_harga_alat']
        
        # Perhitungan Kas Bersih - PER BULAN YANG DIPILIH (BUKAN KUMULATIF)
        kas_bersih_operasional = total_penjualan - total_biaya_produksi - total_biaya_operasional - total_biaya_non_operasional
        kas_bersih_investasi = total_harga_jual_alat - total_harga_alat
        kas_bersih_pembiayaan = 0  # Selalu 0 sesuai permintaan
        total_kas_bersih = kas_bersih_operasional + kas_bersih_investasi + kas_bersih_pembiayaan
        
        # Update data arus kas ke database - OTOMATIS
        cursor.execute("""
            INSERT INTO arus_kas (
                proposal_id, 
                bulan_tahun,
                total_penjualan,
                total_biaya_produksi,
                total_biaya_operasional,
                total_biaya_non_operasional,
                kas_bersih_operasional,
                total_harga_jual_alat,
                total_harga_alat,
                kas_bersih_investasi,
                kas_bersih_pembiayaan,
                total_kas_bersih,
                created_at,
                updated_at
            ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, NOW(), NOW())
            ON DUPLICATE KEY UPDATE
                total_penjualan = VALUES(total_penjualan),
                total_biaya_produksi = VALUES(total_biaya_produksi),
                total_biaya_operasional = VALUES(total_biaya_operasional),
                total_biaya_non_operasional = VALUES(total_biaya_non_operasional),
                kas_bersih_operasional = VALUES(kas_bersih_operasional),
                total_harga_jual_alat = VALUES(total_harga_jual_alat),
                total_harga_alat = VALUES(total_harga_alat),
                kas_bersih_investasi = VALUES(kas_bersih_investasi),
                kas_bersih_pembiayaan = VALUES(kas_bersih_pembiayaan),
                total_kas_bersih = VALUES(total_kas_bersih),
                updated_at = NOW()
        """, (
            proposal_id,
            bulan_tahun,
            total_penjualan,
            total_biaya_produksi,
            total_biaya_operasional,
            total_biaya_non_operasional,
            kas_bersih_operasional,
            total_harga_jual_alat,
            total_harga_alat,
            kas_bersih_investasi,
            kas_bersih_pembiayaan,
            total_kas_bersih
        ))
        mysql.connection.commit()
        cursor.close()
        
        return True
    except Exception as e:
        print(f"Error updating arus kas: {str(e)}")
        return False

def generate_excel_laba_rugi(data, proposal, total_pendapatan, total_biaya_produksi, 
                             total_biaya_operasional, total_laba_kotor, total_laba_bersih, 
                             start_date, end_date, filename):
    try:
        # Check if openpyxl is available
        try:
            import io
            import base64
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
            from openpyxl.drawing.image import Image
            import os
        except ImportError:
            return jsonify({'success': False, 'message': 'Library openpyxl tidak tersedia. Silakan install dengan: pip install openpyxl'})
        
        wb = Workbook()
        wb.properties.creator = "AYMP Unjaya Laba Rugi"
        ws = wb.active
        if ws:
            ws.title = "Laporan Laba Rugi"
            
            # Font styling (hitam putih saja)
            header_font = Font(name='Times New Roman', bold=True, size=12)
            title_font = Font(name='Times New Roman', bold=True, size=12)
            subtitle_font = Font(name='Times New Roman', bold=True, size=12)
            address_font = Font(name='Times New Roman', size=10)
            cell_font = Font(name='Times New Roman', size=12)
            
            # Border untuk tabel
            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            
            # KOP SURAT - Gunakan header dari Kop Surat.pdf yang asli
            from utils import create_pdf_header_from_kop_surat_pdf
            header_template = create_pdf_header_from_kop_surat_pdf()
            header_data = header_template['header_data']
            
            # Teks Kop Surat
            ws['A1'] = header_data[0][1]  # YAYASAN KARTIKA EKA PAKSI
            ws['A1'].font = subtitle_font
            ws['A1'].alignment = Alignment(horizontal='center')
            ws.merge_cells('A1:H1')
            
            ws['A2'] = header_data[1][1]  # UNIVERSITAS JENDERAL ACHMAD YANI
            ws['A2'].font = subtitle_font
            ws['A2'].alignment = Alignment(horizontal='center')
            ws.merge_cells('A2:H2')
            
            ws['A3'] = header_data[2][1]  # YOGYAKARTA
            ws['A3'].font = subtitle_font
            ws['A3'].alignment = Alignment(horizontal='center')
            ws.merge_cells('A3:H3')
            
            ws['A4'] = header_data[3][1]  # Alamat lengkap
            ws['A4'].font = address_font
            ws['A4'].alignment = Alignment(horizontal='center')
            ws.merge_cells('A4:H4')
            
            ws['A5'] = header_data[4][1]  # Telp, Fax, Website
            ws['A5'].font = address_font
            ws['A5'].alignment = Alignment(horizontal='center')
            ws.merge_cells('A5:H5')
            
            ws['A6'] = header_data[5][1]  # Email
            ws['A6'].font = address_font
            ws['A6'].alignment = Alignment(horizontal='center')
            ws.merge_cells('A6:H6')
            
            # Garis pemisah (gunakan border atas pada baris berikutnya)
            for col in range(1, 9):  # Kolom A sampai H
                ws.cell(row=7, column=col).border = Border(
                    top=Side(style='medium', color='000000')
                )
            
            # Title laporan
            ws['A8'] = f"LAPORAN LABA RUGI"
            ws['A8'].font = title_font
            ws['A8'].alignment = Alignment(horizontal='center')
            ws.merge_cells('A8:H8')
            
            ws['A9'] = f"{proposal['judul_usaha']}"
            ws['A9'].font = title_font
            ws['A9'].alignment = Alignment(horizontal='center')
            ws.merge_cells('A9:H9')
            
            # Period
            ws['A10'] = f"Periode: {start_date} s/d {end_date}"
            ws['A10'].font = cell_font
            ws['A10'].alignment = Alignment(horizontal='center')
            ws.merge_cells('A10:H10')
            
            # Empty row
            ws['A11'] = ""
            
            # Table headers
            headers = ['No', 'Tanggal', 'Nama Produk', 'Pendapatan', 'Biaya Produksi', 'Laba Kotor', 'Biaya Operasional', 'Laba Bersih']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=12, column=col, value=header)
                cell.font = header_font
                cell.alignment = Alignment(horizontal='center')
                cell.border = thin_border
            
            # Data rows
            for idx, row_data in enumerate(data, 1):
                row_num = 12 + idx
                ws.cell(row=row_num, column=1, value=idx).font = cell_font
                ws.cell(row=row_num, column=1).border = thin_border
                ws.cell(row=row_num, column=1).alignment = Alignment(horizontal='center')
                
                ws.cell(row=row_num, column=2, value=str(row_data['tanggal_produksi'])).font = cell_font
                ws.cell(row=row_num, column=2).border = thin_border
                ws.cell(row=row_num, column=2).alignment = Alignment(horizontal='center')
                
                ws.cell(row=row_num, column=3, value=row_data['nama_produk']).font = cell_font
                ws.cell(row=row_num, column=3).border = thin_border
                
                ws.cell(row=row_num, column=4, value=f"Rp {row_data['pendapatan']:,}").font = cell_font
                ws.cell(row=row_num, column=4).border = thin_border
                ws.cell(row=row_num, column=4).alignment = Alignment(horizontal='right')
                
                ws.cell(row=row_num, column=5, value=f"Rp {row_data['total_biaya_produksi']:,}").font = cell_font
                ws.cell(row=row_num, column=5).border = thin_border
                ws.cell(row=row_num, column=5).alignment = Alignment(horizontal='right')
                
                ws.cell(row=row_num, column=6, value=f"Rp {row_data['laba_rugi_kotor']:,}").font = cell_font
                ws.cell(row=row_num, column=6).border = thin_border
                ws.cell(row=row_num, column=6).alignment = Alignment(horizontal='right')
                
                ws.cell(row=row_num, column=7, value=f"Rp {row_data['biaya_operasional']:,}").font = cell_font
                ws.cell(row=row_num, column=7).border = thin_border
                ws.cell(row=row_num, column=7).alignment = Alignment(horizontal='right')
                
                ws.cell(row=row_num, column=8, value=f"Rp {row_data['laba_rugi_bersih']:,}").font = cell_font
                ws.cell(row=row_num, column=8).border = thin_border
                ws.cell(row=row_num, column=8).alignment = Alignment(horizontal='right')
            
            # Summary row
            summary_row = len(data) + 13
            ws.cell(row=summary_row, column=1, value="TOTAL").font = header_font
            ws.cell(row=summary_row, column=1).border = thin_border
            ws.cell(row=summary_row, column=1).alignment = Alignment(horizontal='center')
            
            ws.cell(row=summary_row, column=2, value="").border = thin_border
            ws.cell(row=summary_row, column=3, value="").border = thin_border
            
            ws.cell(row=summary_row, column=4, value=f"Rp {total_pendapatan:,}").font = header_font
            ws.cell(row=summary_row, column=4).border = thin_border
            ws.cell(row=summary_row, column=4).alignment = Alignment(horizontal='right')
            
            ws.cell(row=summary_row, column=5, value=f"Rp {total_biaya_produksi:,}").font = header_font
            ws.cell(row=summary_row, column=5).border = thin_border
            ws.cell(row=summary_row, column=5).alignment = Alignment(horizontal='right')
            
            ws.cell(row=summary_row, column=6, value=f"Rp {total_laba_kotor:,}").font = header_font
            ws.cell(row=summary_row, column=6).border = thin_border
            ws.cell(row=summary_row, column=6).alignment = Alignment(horizontal='right')
            
            ws.cell(row=summary_row, column=7, value=f"Rp {total_biaya_operasional:,}").font = header_font
            ws.cell(row=summary_row, column=7).border = thin_border
            ws.cell(row=summary_row, column=7).alignment = Alignment(horizontal='right')
            
            ws.cell(row=summary_row, column=8, value=f"Rp {total_laba_bersih:,}").font = header_font
            ws.cell(row=summary_row, column=8).border = thin_border
            ws.cell(row=summary_row, column=8).alignment = Alignment(horizontal='right')
            
            # Auto-adjust column widths
            ws.column_dimensions['A'].width = 8
            ws.column_dimensions['B'].width = 15
            ws.column_dimensions['C'].width = 25
            ws.column_dimensions['D'].width = 18
            ws.column_dimensions['E'].width = 18
            ws.column_dimensions['F'].width = 15
            ws.column_dimensions['G'].width = 20
            ws.column_dimensions['H'].width = 18
            
            # Atur tinggi baris untuk kop surat tanpa logo
            ws.row_dimensions[1].height = 16  # Spacing rapat
            ws.row_dimensions[2].height = 16  # Spacing rapat
            ws.row_dimensions[3].height = 14  # Spacing rapat
            ws.row_dimensions[4].height = 14  # Spacing rapat
            ws.row_dimensions[5].height = 14  # Spacing rapat
            ws.row_dimensions[6].height = 14  # Spacing rapat
            
            # Save to bytes
            output = io.BytesIO()
            wb.save(output)
            output.seek(0)
            file_bytes = output.getvalue()
            file_b64 = base64.b64encode(file_bytes).decode('utf-8')
            
            # Bersihkan nama file dari karakter yang tidak aman
            import re
            safe_filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
            safe_filename = safe_filename.replace(' ', '_')
            
            return jsonify({
                'success': True,
                'file_data': file_b64,
                'filename': f"{safe_filename}.xlsx"
            })
        else:
            return jsonify({'success': False, 'message': 'Gagal membuat worksheet Excel'})
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error generating Excel: {str(e)}'})

def generate_pdf_laba_rugi(data, proposal, total_pendapatan, total_biaya_produksi, 
                           total_biaya_operasional, total_laba_kotor, total_laba_bersih, 
                           start_date, end_date, filename):
    try:
        # Check if reportlab is available
        try:
            from reportlab.lib.pagesizes import A4, landscape
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image, HRFlowable
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib import colors
            from reportlab.lib.units import inch
            import io
            import base64
            import os
        except ImportError:
            return jsonify({'success': False, 'message': 'Library reportlab tidak tersedia. Silakan install dengan: pip install reportlab'})
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, 
                               pagesize=landscape(A4), 
                               topMargin=0.3*inch, 
                               bottomMargin=0.3*inch,
                               leftMargin=0.3*inch,
                               rightMargin=0.3*inch)
        # doc.author = "AYMP Unjaya Laba Rugi"
        elements = []
        
        # Styles dengan Times New Roman
        styles = getSampleStyleSheet()
        
        # KOP SURAT
        # Ambil header dari Kop Surat.pdf yang asli
        from utils import create_pdf_header_from_kop_surat_pdf
        header_template = create_pdf_header_from_kop_surat_pdf()
        header_data = header_template['header_data']
        logo_path = header_template['logo_path']
        shield_logo_path = header_template['shield_logo_path']
        
        # Kop surat dengan logo jika tersedia
        simple_kop = ParagraphStyle(
                'SimpleKop',
                parent=styles['Normal'],
                fontName='Times-Bold',
                fontSize=12,
                alignment=1,
                spaceAfter=0,
                spaceBefore=0,
                leading=12
            )
            
        simple_address = ParagraphStyle(
                'SimpleAddress',
                parent=styles['Normal'],
                fontName='Times-Roman',
                fontSize=10,
                alignment=1,
                spaceAfter=0,
                spaceBefore=0,
                leading=10
            )
        
        # Header dengan kop surat horizontal - TANPA LOGO
        # Buat table untuk tata letak horizontal: [TEXT] (tanpa logo)
        
        # Buat header dengan styling link yang benar
        from utils import create_header_with_styled_links
        styled_headers = create_header_with_styled_links(header_data)
        
        # Tambahkan header yang sudah di-styling
        for header in styled_headers:
            elements.append(header)
        
        elements.append(Spacer(1, 15))
        
        # Garis pemisah horizontal
        elements.append(Spacer(1, 6))  # Sedikit spasi sebelum garis
        elements.append(HRFlowable(width="100%", thickness=1, lineCap='round', color=colors.black))
        elements.append(Spacer(1, 6))  # Sedikit spasi setelah garis
        
        # Title laporan
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Normal'],
            fontName='Times-Bold',
            fontSize=14,
            alignment=1,
            spaceAfter=12
        )
        
        elements.append(Paragraph("LAPORAN LABA RUGI", title_style))
        elements.append(Paragraph(proposal['judul_usaha'], title_style))
        
        # Period
        period_style = ParagraphStyle(
            'Period',
            parent=styles['Normal'],
            fontName='Times-Roman',
            fontSize=12,
            alignment=1,
            spaceAfter=20
        )
        elements.append(Paragraph(f"Periode: {start_date} s/d {end_date}", period_style))
        
        # Table data
        table_data = [['No', 'Tanggal', 'Nama Produk', 'Pendapatan', 'Biaya Produksi', 'Laba Kotor', 'Biaya Operasional', 'Laba Bersih']]
        
        for idx, row_data in enumerate(data, 1):
            table_data.append([
                str(idx),
                str(row_data['tanggal_produksi']),
                str(row_data['nama_produk']),
                f"Rp {row_data['pendapatan']:,}",
                f"Rp {row_data['total_biaya_produksi']:,}",
                f"Rp {row_data['laba_rugi_kotor']:,}",
                f"Rp {row_data['biaya_operasional']:,}",
                f"Rp {row_data['laba_rugi_bersih']:,}"
            ])
        
        # Summary row
        table_data.append([
            'TOTAL',
            '',
            '',
            f"Rp {total_pendapatan:,}",
            f"Rp {total_biaya_produksi:,}",
            f"Rp {total_laba_kotor:,}",
            f"Rp {total_biaya_operasional:,}",
            f"Rp {total_laba_bersih:,}"
        ])
        
        # Create table dengan lebar kolom yang disesuaikan untuk landscape
        col_widths = [0.6*inch, 1*inch, 1.5*inch, 1.2*inch, 1.2*inch, 1*inch, 1.5*inch, 1.2*inch]
        table = Table(table_data, colWidths=col_widths)
        
        # Table style dengan Times New Roman
        table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Times-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('FONTNAME', (0, 1), (-1, -2), 'Times-Roman'),
            ('FONTSIZE', (0, 1), (-1, -2), 12),
            ('FONTNAME', (0, -1), (-1, -1), 'Times-Bold'),
            ('FONTSIZE', (0, -1), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
        ]))
        
        elements.append(table)
        
        def set_metadata(canvas, doc):
            canvas.setAuthor("AYMP Unjaya Laba Rugi")
            canvas.setTitle("Laporan Laba Rugi")
        doc.build(elements, onFirstPage=set_metadata, onLaterPages=set_metadata)
        buffer.seek(0)
        file_bytes = buffer.getvalue()
        file_b64 = base64.b64encode(file_bytes).decode('utf-8')
        
        # Bersihkan nama file dari karakter yang tidak aman
        import re
        safe_filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
        safe_filename = safe_filename.replace(' ', '_')
        
        return jsonify({
            'success': True,
            'file_data': file_b64,
            'filename': f"{safe_filename}.pdf"
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error generating PDF: {str(e)}'})
        
def generate_word_laba_rugi(data, proposal, total_pendapatan, total_biaya_produksi, 
                            total_biaya_operasional, total_laba_kotor, total_laba_bersih, 
                            start_date, end_date, filename):
        # Check if python-docx is available
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.table import WD_TABLE_ALIGNMENT
            import io
            import base64
            import os
        except ImportError:
            return jsonify({'success': False, 'message': 'Library python-docx tidak tersedia. Silakan install dengan: pip install python-docx'})
        
        doc = Document()
        doc.core_properties.author = "AYMP Unjaya Laba Rugi"
        # Set page margins dan orientation ke landscape
        from docx.enum.section import WD_ORIENTATION
        sections = doc.sections
        for section in sections:
            section.orientation = WD_ORIENTATION.LANDSCAPE
            section.page_width = Inches(11.69)  # A4 landscape width
            section.page_height = Inches(8.27)  # A4 landscape height
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # KOP SURAT - Header konsisten tanpa bergantung file eksternal
        
        # Baris 1: YAYASAN KARTIKA EKA PAKSI
        header1 = doc.add_paragraph()
        header1_run = header1.add_run("YAYASAN KARTIKA EKA PAKSI")
        header1_run.font.name = 'Times New Roman'
        header1_run.font.size = Pt(12)
        header1_run.bold = True
        header1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Baris 2: UNIVERSITAS JENDERAL ACHMAD YANI YOGYAKARTA
        header2 = doc.add_paragraph()
        header2_run = header2.add_run("UNIVERSITAS JENDERAL ACHMAD YANI YOGYAKARTA")
        header2_run.font.name = 'Times New Roman'
        header2_run.font.size = Pt(12)
        header2_run.bold = True
        header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Baris 3: (spasi)
        doc.add_paragraph()
        
        # Baris 4: Alamat
        header4 = doc.add_paragraph()
        header4_run = header4.add_run("Jl Siliwangi Ringroad Barat, Banyuraden, Gamping, Sleman, Yogyakarta 55293")
        header4_run.font.name = 'Times New Roman'
        header4_run.font.size = Pt(10)
        header4_run.bold = False
        header4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Baris 5: Telp, Fax, Website dengan link biru
        header5 = doc.add_paragraph()
        header5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Tambahkan teks sebelum link
        header5_run1 = header5.add_run("Telp. (0274) 552489, 552851 Fax. (0274) 557228 Website: ")
        header5_run1.font.name = 'Times New Roman'
        header5_run1.font.size = Pt(10)
        header5_run1.bold = False
        
        # Tambahkan link website dengan warna biru dan garis bawah
        header5_run2 = header5.add_run("www.unjaya.ac.id")
        header5_run2.font.name = 'Times New Roman'
        header5_run2.font.size = Pt(10)
        header5_run2.bold = False
        header5_run2.font.color.rgb = RGBColor(0, 0, 255)  # Warna biru
        header5_run2.font.underline = True
        
        # Baris 6: Email dengan link biru
        header6 = doc.add_paragraph()
        header6.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Tambahkan teks sebelum link
        header6_run1 = header6.add_run("Email: ")
        header6_run1.font.name = 'Times New Roman'
        header6_run1.font.size = Pt(10)
        header6_run1.bold = False
        
        # Tambahkan link email dengan warna biru dan garis bawah
        header6_run2 = header6.add_run("info@unjaya.ac.id")
        header6_run2.font.name = 'Times New Roman'
        header6_run2.font.size = Pt(10)
        header6_run2.bold = False
        header6_run2.font.color.rgb = RGBColor(0, 0, 255)  # Warna biru
        header6_run2.font.underline = True
        
        # Garis pemisah horizontal
        separator = doc.add_paragraph()
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Tambahkan garis horizontal dengan border
        from docx.oxml.ns import qn
        separator_format = separator.paragraph_format
        separator_format.space_before = Pt(6)
        separator_format.space_after = Pt(6)
        
        # Buat paragraph kosong dengan border bawah sebagai garis
        separator_border = doc.add_paragraph()
        separator_border.paragraph_format.space_before = Pt(0)
        separator_border.paragraph_format.space_after = Pt(0)
        
        # Set border bawah untuk paragraph
        p = separator_border._element
        pPr = p.get_or_add_pPr()
        pBdr = pPr.first_child_found_in("w:pBdr")
        if pBdr is None:
            from docx.oxml import parse_xml
            pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/></w:pBdr>')
            pPr.append(pBdr)
        
        # Title laporan
        title = doc.add_paragraph()
        title_run = title.add_run("LAPORAN LABA RUGI")
        title_run.font.name = 'Times New Roman'
        title_run.font.size = Pt(14)
        title_run.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Judul usaha
        judul = doc.add_paragraph()
        judul_run = judul.add_run(proposal['judul_usaha'])
        judul_run.font.name = 'Times New Roman'
        judul_run.font.size = Pt(12)
        judul_run.bold = True
        judul.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Period
        period = doc.add_paragraph()
        period_run = period.add_run(f'Periode: {start_date} s/d {end_date}')
        period_run.font.name = 'Times New Roman'
        period_run.font.size = Pt(12)
        period.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Empty line
        
        # Create table
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        header_cells = table.rows[0].cells
        headers = ['No', 'Tanggal', 'Nama Produk', 'Pendapatan', 'Biaya Produksi', 'Laba Kotor', 'Biaya Operasional', 'Laba Bersih']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Format header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data rows
        for idx, row_data in enumerate(data, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = str(row_data['tanggal_produksi'])
            row_cells[2].text = str(row_data['nama_produk'])
            row_cells[3].text = f"Rp {row_data['pendapatan']:,}"
            row_cells[4].text = f"Rp {row_data['total_biaya_produksi']:,}"
            row_cells[5].text = f"Rp {row_data['laba_rugi_kotor']:,}"
            row_cells[6].text = f"Rp {row_data['biaya_operasional']:,}"
            row_cells[7].text = f"Rp {row_data['laba_rugi_bersih']:,}"
            
            # Format data cells
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary row
        summary_cells = table.add_row().cells
        summary_cells[0].text = "TOTAL"
        summary_cells[1].text = ""
        summary_cells[2].text = ""
        summary_cells[3].text = f"Rp {total_pendapatan:,}"
        summary_cells[4].text = f"Rp {total_biaya_produksi:,}"
        summary_cells[5].text = f"Rp {total_laba_kotor:,}"
        summary_cells[6].text = f"Rp {total_biaya_operasional:,}"
        summary_cells[7].text = f"Rp {total_laba_bersih:,}"
        
        # Format summary cells
        for cell in summary_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        file_bytes = buffer.getvalue()
        file_b64 = base64.b64encode(file_bytes).decode('utf-8')
        
        # Bersihkan nama file dari karakter yang tidak aman
        import re
        safe_filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
        safe_filename = safe_filename.replace(' ', '_')
        
        return jsonify({
            'success': True,
            'file_data': file_b64,
            'filename': f"{safe_filename}.docx"
        })

def generate_excel_arus_kas(data, proposal, bulan_tahun, nama_bulan, filename):
    try:
        # Check if openpyxl is available
        try:
            import io
            import base64
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
            from openpyxl.drawing.image import Image
            import os
        except ImportError:
            return jsonify({'success': False, 'message': 'Library openpyxl tidak tersedia. Silakan install dengan: pip install openpyxl'})
        
        wb = Workbook()
        wb.properties.creator = "AYMP Unjaya Arus Kas"
        ws = wb.active
        if ws:
            ws.title = "Laporan Arus Kas"
            
            # Font styling (hitam putih saja)
            header_font = Font(name='Times New Roman', bold=True, size=11)
            title_font = Font(name='Times New Roman', bold=True, size=12)
            subtitle_font = Font(name='Times New Roman', bold=True, size=12)
            address_font = Font(name='Times New Roman', size=10)
            cell_font = Font(name='Times New Roman', size=10)
            
            # Border untuk tabel
            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            
            # KOP SURAT
            # Teks Kop Surat
            # Gunakan kop surat dinamis
            kop_surat = get_kop_surat_content()
            
            ws['A1'] = kop_surat['yayasan']
            ws['A1'].font = subtitle_font
            ws['A1'].alignment = Alignment(horizontal='center')
            ws.merge_cells('A1:H1')
            
            ws['A2'] = kop_surat['universitas']
            ws['A2'].font = subtitle_font
            ws['A2'].alignment = Alignment(horizontal='center')
            ws.merge_cells('A2:H2')
            
            ws['A3'] = kop_surat['alamat']
            ws['A3'].font = address_font
            ws['A3'].alignment = Alignment(horizontal='center')
            ws.merge_cells('A3:H3')
            
            ws['A4'] = kop_surat['kontak']
            ws['A4'].font = address_font
            ws['A4'].alignment = Alignment(horizontal='center')
            ws.merge_cells('A4:H4')
            
            ws['A5'] = kop_surat['email']
            ws['A5'].font = address_font
            ws['A5'].alignment = Alignment(horizontal='center')
            ws.merge_cells('A5:H5')
            
            # Garis pemisah
            for col in range(1, 9):
                ws.cell(row=6, column=col).border = Border(
                    top=Side(style='medium', color='000000')
                )
            
            # Title laporan
            # Baris 7: Nama Usaha
            ws['A7'] = f"{proposal['judul_usaha']}"
            ws['A7'].font = Font(name='Times New Roman', bold=True, size=14)
            ws['A7'].alignment = Alignment(horizontal='center')
            ws.merge_cells('A7:H7')

            # Baris 8: Laporan Arus Kas
            ws['A8'] = "Laporan Arus Kas"
            ws['A8'].font = Font(name='Times New Roman', bold=True, size=14)
            ws['A8'].alignment = Alignment(horizontal='center')
            ws.merge_cells('A8:H8')

            # Baris 9: Untuk Periode yang Berakhir [Nama Bulan Tahun]
            ws['A9'] = f"Untuk Periode yang Berakhir {nama_bulan}"
            ws['A9'].font = Font(name='Times New Roman', size=12)
            ws['A9'].alignment = Alignment(horizontal='center')
            ws.merge_cells('A9:H9')
                    
            # Empty row
            ws['A10'] = ""
            
            # Arus Kas dari Kegiatan Operasional
            ws['A11'] = "Kas dari Kegiatan Operasional"
            ws['A11'].font = header_font
            ws['A11'].alignment = Alignment(horizontal='left')
            
            # Item operasional
            ws['A12'] = "a. Penerimaan Kas dari penjualan"
            ws['A12'].font = cell_font
            ws['A12'].alignment = Alignment(horizontal='left', indent=1)
            ws['H12'] = f"Rp {data['total_penjualan']:,}"
            ws['H12'].font = cell_font
            ws['H12'].alignment = Alignment(horizontal='right')
            
            ws['A13'] = "b. Pengeluaran Kas untuk Supplier"
            ws['A13'].font = cell_font
            ws['A13'].alignment = Alignment(horizontal='left', indent=1)
            ws['H13'] = f"Rp {data['total_biaya_produksi']:,}"
            ws['H13'].font = cell_font
            ws['H13'].alignment = Alignment(horizontal='right')
            
            ws['A14'] = "c. Pengeluaran Kas untuk Operasional Lainnya"
            ws['A14'].font = cell_font
            ws['A14'].alignment = Alignment(horizontal='left', indent=1)
            ws['H14'] = f"Rp {data['total_biaya_operasional']:,}"
            ws['H14'].font = cell_font
            ws['H14'].alignment = Alignment(horizontal='right')
            
            ws['A15'] = "d. Pengeluaran Lainnya"
            ws['A15'].font = cell_font
            ws['A15'].alignment = Alignment(horizontal='left', indent=1)
            ws['H15'] = f"Rp {data['total_biaya_non_operasional']:,}"
            ws['H15'].font = cell_font
            ws['H15'].alignment = Alignment(horizontal='right')
            
            # Kas bersih operasional
            ws['A16'] = "Kas bersih dari Kegiatan Operasional"
            ws['A16'].font = header_font
            ws['A16'].alignment = Alignment(horizontal='left')
            ws['H16'] = f"Rp {data['kas_bersih_operasional']:,}"
            ws['H16'].font = header_font
            ws['H16'].alignment = Alignment(horizontal='right')
            
            # Empty row
            ws['A17'] = ""
            
            # Arus Kas dari Kegiatan Investasi
            ws['A18'] = "Kas dari Kegiatan Investasi (jika ada)"
            ws['A18'].font = header_font
            ws['A18'].alignment = Alignment(horizontal='left')
            
            # Item investasi
            ws['A19'] = "a. Penerimaan Kas dari penjualan harta tetap"
            ws['A19'].font = cell_font
            ws['A19'].alignment = Alignment(horizontal='left', indent=1)
            ws['H19'] = f"Rp {data['total_harga_jual_alat']:,}"
            ws['H19'].font = cell_font
            ws['H19'].alignment = Alignment(horizontal='right')
            
            ws['A20'] = "b. Pengeluaran Kas dari pembelian harta tetap"
            ws['A20'].font = cell_font
            ws['A20'].alignment = Alignment(horizontal='left', indent=1)
            ws['H20'] = f"Rp {data['total_harga_alat']:,}"
            ws['H20'].font = cell_font
            ws['H20'].alignment = Alignment(horizontal='right')
            
            # Kas bersih investasi
            ws['A21'] = "Kas bersih dari Kegiatan Investasi"
            ws['A21'].font = header_font
            ws['A21'].alignment = Alignment(horizontal='left')
            ws['H21'] = f"Rp {data['kas_bersih_investasi']:,}"
            ws['H21'].font = header_font
            ws['H21'].alignment = Alignment(horizontal='right')
            
            # Empty row
            ws['A22'] = ""
            
            # Arus Kas dari Kegiatan Pembiayaan
            ws['A23'] = "Kas dari Kegiatan Pembiayaan (jika ada)"
            ws['A23'].font = header_font
            ws['A23'].alignment = Alignment(horizontal='left')
            
            # Item pembiayaan
            ws['A24'] = "a. Penerimaan Kas dari Penerbitan Saham"
            ws['A24'].font = cell_font
            ws['A24'].alignment = Alignment(horizontal='left', indent=1)
            ws['H24'] = "Rp 0"
            ws['H24'].font = cell_font
            ws['H24'].alignment = Alignment(horizontal='right')
            
            ws['A25'] = "b. Pengeluaran Kas untuk Pembayaran Cicilan Hutang Bank"
            ws['A25'].font = cell_font
            ws['A25'].alignment = Alignment(horizontal='left', indent=1)
            ws['H25'] = "Rp 0"
            ws['H25'].font = cell_font
            ws['H25'].alignment = Alignment(horizontal='right')
            
            ws['A26'] = "c. Penerimaan Kas dari Pinjaman Bank"
            ws['A26'].font = cell_font
            ws['A26'].alignment = Alignment(horizontal='left', indent=1)
            ws['H26'] = "Rp 0"
            ws['H26'].font = cell_font
            ws['H26'].alignment = Alignment(horizontal='right')
            
            # Kas bersih pembiayaan
            ws['A27'] = "Kas bersih dari Kegiatan Pembiayaan"
            ws['A27'].font = header_font
            ws['H27'] = f"Rp {data['kas_bersih_pembiayaan']:,}"
            ws['H27'].font = header_font
            ws['H27'].alignment = Alignment(horizontal='right')
            
            # Empty row
            ws['A28'] = ""
            
            # Total Kas Bersih
            ws['A29'] = "Total Kas Bersih"
            ws['A29'].font = header_font
            ws['A29'].alignment = Alignment(horizontal='left')
            ws['H29'] = f"Rp {data['total_kas_bersih']:,}"
            ws['H29'].font = header_font
            ws['H29'].alignment = Alignment(horizontal='right')
            
            # Auto-adjust column widths
            ws.column_dimensions['A'].width = 50
            ws.column_dimensions['H'].width = 20
            
            # Atur tinggi baris
            ws.row_dimensions[1].height = 16
            ws.row_dimensions[2].height = 16
            ws.row_dimensions[3].height = 14
            ws.row_dimensions[4].height = 14
            ws.row_dimensions[5].height = 14
            
            # Save to bytes
            output = io.BytesIO()
            wb.save(output)
            output.seek(0)
            file_bytes = output.getvalue()
            file_b64 = base64.b64encode(file_bytes).decode('utf-8')
            return jsonify({
                'success': True,
                'file_data': file_b64,
                'filename': f"{filename}.xlsx"
            })
        else:
            return jsonify({'success': False, 'message': 'Gagal membuat worksheet Excel'})
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error generating Excel: {str(e)}'})

def generate_pdf_arus_kas(data, proposal, bulan_tahun, nama_bulan, filename):
    try:
        # Check if reportlab is available
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            import io
            import base64
        except ImportError:
            return jsonify({'success': False, 'message': 'Library reportlab tidak tersedia. Silakan install dengan: pip install reportlab'})
        
        buffer = io.BytesIO()
        # Kurangi margin agar muat dalam 1 halaman
        doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=0.8*inch, rightMargin=0.8*inch, topMargin=0.6*inch, bottomMargin=0.6*inch)
        story = []
        
        # Styles yang lebih compact
        styles = getSampleStyleSheet()
        
        # Style untuk kop surat
        kop_style = ParagraphStyle(
            'KopStyle',
            parent=styles['Normal'],
            fontSize=12,
            spaceAfter=0,
            spaceBefore=0,
            leading=12,
            alignment=1,  # Center
            fontName='Times-Bold'
        )
        
        # Style untuk alamat
        address_style = ParagraphStyle(
            'AddressStyle',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=0,
            spaceBefore=0,
            leading=10,
            alignment=1,  # Center
            fontName='Times-Roman'
        )
        
        # Style untuk judul usaha
        judul_usaha_style = ParagraphStyle(
            'JudulUsahaStyle',
            parent=styles['Normal'],
            fontSize=14,
            spaceAfter=2,
            alignment=1,  # Center
            fontName='Times-Roman',
            bold=True
        )
        
        # Style untuk judul laporan
        judul_laporan_style = ParagraphStyle(
            'JudulLaporanStyle',
            parent=styles['Normal'],
            fontSize=14,
            spaceAfter=2,
            alignment=1,  # Center
            fontName='Times-Roman',
            bold=True
        )
        
        # Style untuk periode
        periode_style = ParagraphStyle(
            'PeriodeStyle',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            alignment=1,  # Center
            fontName='Times-Roman'
        )
        
        # Style untuk header section
        section_header_style = ParagraphStyle(
            'SectionHeaderStyle',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=2,
            alignment=0,  # Left
            fontName='Times-Roman',
            bold=True
        )
        
        # Kop Surat dengan tata letak horizontal - LOGO WAJIB ADA
        # Ambil header dari Kop Surat.pdf yang asli
        from utils import create_pdf_header_from_kop_surat_pdf
        header_template = create_pdf_header_from_kop_surat_pdf()
        header_data = header_template['header_data']
        logo_path = header_template['logo_path']
        shield_logo_path = header_template['shield_logo_path']
        
        # Header dengan kop surat horizontal - TANPA LOGO
        # Buat table untuk tata letak horizontal: [TEXT] (tanpa logo)
        
        # Buat header dengan styling link yang benar
        from utils import create_header_with_styled_links
        styled_headers = create_header_with_styled_links(header_data)
        
        # Tambahkan header yang sudah di-styling
        for header in styled_headers:
            story.append(header)
        
        story.append(Spacer(1, 15))
        
        # Garis pemisah horizontal
        from reportlab.platypus import HRFlowable
        story.append(Spacer(1, 6))  # Sedikit spasi sebelum garis
        story.append(HRFlowable(width="100%", thickness=1, lineCap='round', color=colors.black))
        story.append(Spacer(1, 6))  # Sedikit spasi setelah garis
        
        # Judul
        story.append(Paragraph(f"{proposal['judul_usaha']}", judul_usaha_style))
        story.append(Paragraph("Laporan Arus Kas", judul_laporan_style))
        story.append(Paragraph(f"Untuk Periode yang Berakhir {nama_bulan}", periode_style))
        
        # Arus Kas dari Kegiatan Operasional
        story.append(Paragraph("Kas dari Kegiatan Operasional", section_header_style))
        
        # Table untuk operasional
        operasional_data = [
            ["a. Penerimaan Kas dari penjualan", f"Rp {data['total_penjualan']:,}"],
            ["b. Pengeluaran Kas untuk Supplier", f"Rp {data['total_biaya_produksi']:,}"],
            ["c. Pengeluaran Kas untuk Operasional Lainnya", f"Rp {data['total_biaya_operasional']:,}"],
            ["d. Pengeluaran Lainnya", f"Rp {data['total_biaya_non_operasional']:,}"],
            ["Kas bersih dari Kegiatan Operasional", f"Rp {data['kas_bersih_operasional']:,}"]
        ]
        
        operasional_table = Table(operasional_data, colWidths=[4*inch, 1.5*inch])
        operasional_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (0, -1), 'LEFT'),  # Left align text
            ('ALIGN', (1, 0), (1, -1), 'RIGHT'),  # Right align money
            ('FONTNAME', (0, 0), (-1, -2), 'Times-Roman'),
            ('FONTSIZE', (0, 0), (-1, -2), 10),
            ('FONTNAME', (0, -1), (-1, -1), 'Times-Bold'),  # Bold for subtotal
            ('FONTSIZE', (0, -1), (-1, -1), 10),
            ('LEFTPADDING', (0, 0), (0, -1), 20),  # Indent for items
            ('BOTTOMPADDING', (0, 0), (-1, -1), 1),
            ('TOPPADDING', (0, 0), (-1, -1), 1),
        ]))
        story.append(operasional_table)
        story.append(Spacer(1, 3))
        
        # Arus Kas dari Kegiatan Investasi
        story.append(Paragraph("Kas dari Kegiatan Investasi (jika ada)", section_header_style))
        
        # Table untuk investasi
        investasi_data = [
            ["a. Penerimaan Kas dari penjualan harta tetap", f"Rp {data['total_harga_jual_alat']:,}"],
            ["b. Pengeluaran Kas dari pembelian harta tetap", f"Rp {data['total_harga_alat']:,}"],
            ["Kas bersih dari Kegiatan Investasi", f"Rp {data['kas_bersih_investasi']:,}"]
        ]
        
        investasi_table = Table(investasi_data, colWidths=[4*inch, 1.5*inch])
        investasi_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (0, -1), 'LEFT'),  # Left align text
            ('ALIGN', (1, 0), (1, -1), 'RIGHT'),  # Right align money
            ('FONTNAME', (0, 0), (-1, -2), 'Times-Roman'),
            ('FONTSIZE', (0, 0), (-1, -2), 10),
            ('FONTNAME', (0, -1), (-1, -1), 'Times-Bold'),  # Bold for subtotal
            ('FONTSIZE', (0, -1), (-1, -1), 10),
            ('LEFTPADDING', (0, 0), (0, -1), 20),  # Indent for items
            ('BOTTOMPADDING', (0, 0), (-1, -1), 1),
            ('TOPPADDING', (0, 0), (-1, -1), 1),
        ]))
        story.append(investasi_table)
        story.append(Spacer(1, 3))
        
        # Arus Kas dari Kegiatan Pembiayaan
        story.append(Paragraph("Kas dari Kegiatan Pembiayaan (jika ada)", section_header_style))
        
        # Table untuk pembiayaan
        pembiayaan_data = [
            ["a. Penerimaan Kas dari Penerbitan Saham", "Rp 0"],
            ["b. Pengeluaran Kas untuk Pembayaran Cicilan Hutang Bank", "Rp 0"],
            ["c. Penerimaan Kas dari Pinjaman Bank", "Rp 0"],
            ["Kas bersih dari Kegiatan Pembiayaan", f"Rp {data['kas_bersih_pembiayaan']:,}"]
        ]
        
        pembiayaan_table = Table(pembiayaan_data, colWidths=[4*inch, 1.5*inch])
        pembiayaan_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (0, -1), 'LEFT'),  # Left align text
            ('ALIGN', (1, 0), (1, -1), 'RIGHT'),  # Right align money
            ('FONTNAME', (0, 0), (-1, -2), 'Times-Roman'),
            ('FONTSIZE', (0, 0), (-1, -2), 10),
            ('FONTNAME', (0, -1), (-1, -1), 'Times-Bold'),  # Bold for subtotal
            ('FONTSIZE', (0, -1), (-1, -1), 10),
            ('LEFTPADDING', (0, 0), (0, -1), 20),  # Indent for items
            ('BOTTOMPADDING', (0, 0), (-1, -1), 1),
            ('TOPPADDING', (0, 0), (-1, -1), 1),
        ]))
        story.append(pembiayaan_table)
        story.append(Spacer(1, 3))
        
        # Total Kas Bersih
        total_data = [["Total Kas Bersih", f"Rp {data['total_kas_bersih']:,}"]]
        total_table = Table(total_data, colWidths=[4*inch, 1.5*inch])
        total_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (0, -1), 'LEFT'),  # Left align text
            ('ALIGN', (1, 0), (1, -1), 'RIGHT'),  # Right align money
            ('FONTNAME', (0, 0), (-1, -1), 'Times-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 1),
            ('TOPPADDING', (0, 0), (-1, -1), 1),
        ]))
        story.append(total_table)
        
        doc.build(story)
        buffer.seek(0)
        file_bytes = buffer.getvalue()
        file_b64 = base64.b64encode(file_bytes).decode('utf-8')
        
        # Bersihkan nama file dari karakter yang tidak aman
        import re
        safe_filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
        safe_filename = safe_filename.replace(' ', '_')
        
        return jsonify({
            'success': True,
            'file_data': file_b64,
            'filename': f"{safe_filename}.pdf"
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error generating PDF: {str(e)}'})

def generate_word_arus_kas(data, proposal, bulan_tahun, nama_bulan, filename):
    try:
        # Check if python-docx is available
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.table import WD_TABLE_ALIGNMENT
            import io
            import base64
            import os
        except ImportError:
            return jsonify({'success': False, 'message': 'Library python-docx tidak tersedia. Silakan install dengan: pip install python-docx'})
        
        doc = Document()
        doc.core_properties.author = "AYMP Unjaya Arus Kas"
        
        # Set page margins yang lebih kecil agar muat dalam 1 halaman
        from docx.enum.section import WD_ORIENTATION
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.4)
            section.bottom_margin = Inches(0.4)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
        
        # KOP SURAT - Header konsisten tanpa bergantung file eksternal
        
        # Baris 1: YAYASAN KARTIKA EKA PAKSI
        header1 = doc.add_paragraph()
        header1_run = header1.add_run("YAYASAN KARTIKA EKA PAKSI")
        header1_run.font.name = 'Times New Roman'
        header1_run.font.size = Pt(12)
        header1_run.bold = True
        header1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Baris 2: UNIVERSITAS JENDERAL ACHMAD YANI YOGYAKARTA
        header2 = doc.add_paragraph()
        header2_run = header2.add_run("UNIVERSITAS JENDERAL ACHMAD YANI YOGYAKARTA")
        header2_run.font.name = 'Times New Roman'
        header2_run.font.size = Pt(12)
        header2_run.bold = True
        header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Baris 3: (spasi)
        doc.add_paragraph()
        
        # Baris 4: Alamat
        header4 = doc.add_paragraph()
        header4_run = header4.add_run("Jl Siliwangi Ringroad Barat, Banyuraden, Gamping, Sleman, Yogyakarta 55293")
        header4_run.font.name = 'Times New Roman'
        header4_run.font.size = Pt(10)
        header4_run.bold = False
        header4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Baris 5: Telp, Fax, Website dengan link biru
        header5 = doc.add_paragraph()
        header5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Tambahkan teks sebelum link
        header5_run1 = header5.add_run("Telp. (0274) 552489, 552851 Fax. (0274) 557228 Website: ")
        header5_run1.font.name = 'Times New Roman'
        header5_run1.font.size = Pt(10)
        header5_run1.bold = False
        
        # Tambahkan link website dengan warna biru dan garis bawah
        header5_run2 = header5.add_run("www.unjaya.ac.id")
        header5_run2.font.name = 'Times New Roman'
        header5_run2.font.size = Pt(10)
        header5_run2.bold = False
        header5_run2.font.color.rgb = RGBColor(0, 0, 255)  # Warna biru
        header5_run2.font.underline = True
        
        # Baris 6: Email dengan link biru
        header6 = doc.add_paragraph()
        header6.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Tambahkan teks sebelum link
        header6_run1 = header6.add_run("Email: ")
        header6_run1.font.name = 'Times New Roman'
        header6_run1.font.size = Pt(10)
        header6_run1.bold = False
        
        # Tambahkan link email dengan warna biru dan garis bawah
        header6_run2 = header6.add_run("info@unjaya.ac.id")
        header6_run2.font.name = 'Times New Roman'
        header6_run2.font.size = Pt(10)
        header6_run2.bold = False
        header6_run2.font.color.rgb = RGBColor(0, 0, 255)  # Warna biru
        header6_run2.font.underline = True
        
        # Garis pemisah horizontal
        separator = doc.add_paragraph()
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Tambahkan garis horizontal dengan border
        from docx.oxml.ns import qn
        separator_format = separator.paragraph_format
        separator_format.space_before = Pt(6)
        separator_format.space_after = Pt(6)
        
        # Buat paragraph kosong dengan border bawah sebagai garis
        separator_border = doc.add_paragraph()
        separator_border.paragraph_format.space_before = Pt(0)
        separator_border.paragraph_format.space_after = Pt(0)
        
        # Set border bawah untuk paragraph
        p = separator_border._element
        pPr = p.get_or_add_pPr()
        pBdr = pPr.first_child_found_in("w:pBdr")
        if pBdr is None:
            from docx.oxml import parse_xml
            pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/></w:pBdr>')
            pPr.append(pBdr)
        
        # Judul usaha
        judul_usaha = doc.add_paragraph()
        judul_usaha.paragraph_format.space_after = Pt(2)
        judul_usaha_run = judul_usaha.add_run(proposal['judul_usaha'])
        judul_usaha_run.font.name = 'Times New Roman'
        judul_usaha_run.font.size = Pt(14)
        judul_usaha_run.bold = True
        judul_usaha.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Judul laporan
        judul_laporan = doc.add_paragraph()
        judul_laporan.paragraph_format.space_after = Pt(2)
        judul_laporan_run = judul_laporan.add_run("Laporan Arus Kas")
        judul_laporan_run.font.name = 'Times New Roman'
        judul_laporan_run.font.size = Pt(14)
        judul_laporan_run.bold = True
        judul_laporan.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Periode
        periode = doc.add_paragraph()
        periode.paragraph_format.space_after = Pt(6)
        periode_run = periode.add_run(f"Untuk Periode yang Berakhir {nama_bulan}")
        periode_run.font.name = 'Times New Roman'
        periode_run.font.size = Pt(12)
        periode.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Arus Kas dari Kegiatan Operasional
        operasional_title = doc.add_paragraph()
        operasional_title_run = operasional_title.add_run("Kas dari Kegiatan Operasional")
        operasional_title_run.font.name = 'Times New Roman'
        operasional_title_run.font.size = Pt(11)
        operasional_title_run.bold = True
        
        # Item operasional
        items = [
            ("a. Penerimaan Kas dari penjualan", f"Rp {data['total_penjualan']:,}"),
            ("b. Pengeluaran Kas untuk Supplier", f"Rp {data['total_biaya_produksi']:,}"),
            ("c. Pengeluaran Kas untuk Operasional Lainnya", f"Rp {data['total_biaya_operasional']:,}"),
            ("d. Pengeluaran Lainnya", f"Rp {data['total_biaya_non_operasional']:,}")
        ]
        
        for item, value in items:
            item_para = doc.add_paragraph()
            item_para.paragraph_format.left_indent = Inches(0.5)
            item_para.paragraph_format.space_after = Pt(2)
            item_run = item_para.add_run(item)
            item_run.font.name = 'Times New Roman'
            item_run.font.size = Pt(10)
            
            # Tambahkan tab dan nilai dalam paragraph yang sama
            item_para.add_run("\t")
            value_run = item_para.add_run(value)
            value_run.font.name = 'Times New Roman'
            value_run.font.size = Pt(10)
            
            # Set tab stops untuk alignment
            tab_stops = item_para.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(7), WD_TAB_ALIGNMENT.RIGHT)
        
        # Kas bersih operasional
        kas_operasional_para = doc.add_paragraph()
        title_run = kas_operasional_para.add_run("Kas bersih dari Kegiatan Operasional")
        title_run.font.name = 'Times New Roman'
        title_run.font.size = Pt(11)
        title_run.bold = True
        
        # Tambahkan tab untuk alignment
        kas_operasional_para.add_run("\t")
        kas_operasional_value_run = kas_operasional_para.add_run(f"Rp {data['kas_bersih_operasional']:,}")
        kas_operasional_value_run.font.name = 'Times New Roman'
        kas_operasional_value_run.font.size = Pt(11)
        kas_operasional_value_run.bold = True
        
        # Set tab stops untuk alignment
        tab_stops = kas_operasional_para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(7), WD_TAB_ALIGNMENT.RIGHT)
        
        # Arus Kas dari Kegiatan Investasi
        investasi_title = doc.add_paragraph()
        investasi_title_run = investasi_title.add_run("Kas dari Kegiatan Investasi (jika ada)")
        investasi_title_run.font.name = 'Times New Roman'
        investasi_title_run.font.size = Pt(11)
        investasi_title_run.bold = True
        
        # Item investasi
        investasi_items = [
            ("a. Penerimaan Kas dari penjualan harta tetap", f"Rp {data['total_harga_jual_alat']:,}"),
            ("b. Pengeluaran Kas dari pembelian harta tetap", f"Rp {data['total_harga_alat']:,}")
        ]
        
        for item, value in investasi_items:
            item_para = doc.add_paragraph()
            item_para.paragraph_format.left_indent = Inches(0.5)
            item_para.paragraph_format.space_after = Pt(2)
            item_run = item_para.add_run(item)
            item_run.font.name = 'Times New Roman'
            item_run.font.size = Pt(10)
            
            # Tambahkan tab dan nilai dalam paragraph yang sama
            item_para.add_run("\t")
            value_run = item_para.add_run(value)
            value_run.font.name = 'Times New Roman'
            value_run.font.size = Pt(10)
            
            # Set tab stops untuk alignment
            tab_stops = item_para.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(7), WD_TAB_ALIGNMENT.RIGHT)
        
        # Kas bersih investasi
        kas_investasi_para = doc.add_paragraph()
        title_run = kas_investasi_para.add_run("Kas bersih dari Kegiatan Investasi")
        title_run.font.name = 'Times New Roman'
        title_run.font.size = Pt(11)
        title_run.bold = True
        
        # Tambahkan tab untuk alignment
        kas_investasi_para.add_run("\t")
        kas_investasi_value_run = kas_investasi_para.add_run(f"Rp {data['kas_bersih_investasi']:,}")
        kas_investasi_value_run.font.name = 'Times New Roman'
        kas_investasi_value_run.font.size = Pt(11)
        kas_investasi_value_run.bold = True
        
        # Set tab stops untuk alignment
        tab_stops = kas_investasi_para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(7), WD_TAB_ALIGNMENT.RIGHT)
        
        # Arus Kas dari Kegiatan Pembiayaan
        pembiayaan_title = doc.add_paragraph()
        pembiayaan_title_run = pembiayaan_title.add_run("Kas dari Kegiatan Pembiayaan (jika ada)")
        pembiayaan_title_run.font.name = 'Times New Roman'
        pembiayaan_title_run.font.size = Pt(11)
        pembiayaan_title_run.bold = True
        
        # Item pembiayaan
        pembiayaan_items = [
            ("a. Penerimaan Kas dari Penerbitan Saham", "Rp 0"),
            ("b. Pengeluaran Kas untuk Pembayaran Cicilan Hutang Bank", "Rp 0"),
            ("c. Penerimaan Kas dari Pinjaman Bank", "Rp 0")
        ]
        
        for item, value in pembiayaan_items:
            item_para = doc.add_paragraph()
            item_para.paragraph_format.left_indent = Inches(0.5)
            item_para.paragraph_format.space_after = Pt(2)
            item_run = item_para.add_run(item)
            item_run.font.name = 'Times New Roman'
            item_run.font.size = Pt(10)
            
            # Tambahkan tab dan nilai dalam paragraph yang sama
            item_para.add_run("\t")
            value_run = item_para.add_run(value)
            value_run.font.name = 'Times New Roman'
            value_run.font.size = Pt(10)
            
            # Set tab stops untuk alignment
            tab_stops = item_para.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(7), WD_TAB_ALIGNMENT.RIGHT)
        
        # Kas bersih pembiayaan
        kas_pembiayaan_para = doc.add_paragraph()
        title_run = kas_pembiayaan_para.add_run("Kas bersih dari Kegiatan Pembiayaan")
        title_run.font.name = 'Times New Roman'
        title_run.font.size = Pt(11)
        title_run.bold = True
        
        # Tambahkan tab untuk alignment
        kas_pembiayaan_para.add_run("\t")
        kas_pembiayaan_value_run = kas_pembiayaan_para.add_run(f"Rp {data['kas_bersih_pembiayaan']:,}")
        kas_pembiayaan_value_run.font.name = 'Times New Roman'
        kas_pembiayaan_value_run.font.size = Pt(11)
        kas_pembiayaan_value_run.bold = True
        
        # Set tab stops untuk alignment
        tab_stops = kas_pembiayaan_para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(7), WD_TAB_ALIGNMENT.RIGHT)
        
        # Total Kas Bersih
        total_para = doc.add_paragraph()
        title_run = total_para.add_run("Total Kas Bersih")
        title_run.font.name = 'Times New Roman'
        title_run.font.size = Pt(11)
        title_run.bold = True
        
        # Tambahkan tab untuk alignment
        total_para.add_run("\t")
        total_value_run = total_para.add_run(f"Rp {data['total_kas_bersih']:,}")
        total_value_run.font.name = 'Times New Roman'
        total_value_run.font.size = Pt(11)
        total_value_run.bold = True
        
        # Set tab stops untuk alignment
        tab_stops = total_para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(7), WD_TAB_ALIGNMENT.RIGHT)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        file_bytes = buffer.getvalue()
        file_b64 = base64.b64encode(file_bytes).decode('utf-8')
        
        # Bersihkan nama file dari karakter yang tidak aman
        import re
        safe_filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
        safe_filename = safe_filename.replace(' ', '_')
        
        return jsonify({
            'success': True,
            'file_data': file_b64,
            'filename': f"{safe_filename}.docx"
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error generating Word: {str(e)}'})



@app.route('/logout')
def logout():
    # End session tracking untuk mahasiswa
    if 'user_type' in session and session['user_type'] == 'mahasiswa':
        mahasiswa_info = get_mahasiswa_info_from_session()
        if mahasiswa_info:
            end_mahasiswa_session(mahasiswa_info['id'])
            # Log aktivitas logout
            log_mahasiswa_activity(
                mahasiswa_info['id'], 
                mahasiswa_info['nim'], 
                mahasiswa_info['nama_ketua'], 
                'logout', 
                'sistem', 
                'authentication', 
                'Logout dari sistem AYMP'
            )
    
    # End session tracking untuk pembimbing
    elif 'user_type' in session and session['user_type'] == 'pembimbing':
        pembimbing_info = get_pembimbing_info_from_session()
        if pembimbing_info:
            end_pembimbing_session(pembimbing_info['id'])
            # Log aktivitas logout
            log_pembimbing_activity(
                pembimbing_info['id'], 
                pembimbing_info['nip'], 
                pembimbing_info['nama'], 
                'logout', 
                'sistem', 
                'authentication', 
                'Logout dari sistem AYMP'
            )
    
    session.clear()
    flash('Anda telah berhasil logout!', 'success')
    return redirect(url_for('index'))

@app.route('/uploads/<path:filename>')
def uploaded_file(filename):
    """Serve uploaded files"""
    try:
        # Normalisasi path untuk menangani karakter khusus
        file_path = os.path.join('static/uploads', filename)
        
        # Cek apakah file ada
        if not os.path.exists(file_path):
            return "File tidak ditemukan", 404
        
        # Deteksi mimetype berdasarkan ekstensi file
        ext = os.path.splitext(filename)[1].lower()
        if ext == '.pdf':
            return send_file(file_path, mimetype='application/pdf', as_attachment=False)
        else:
            # Untuk file lain, gunakan mimetype otomatis
            import mimetypes
            mimetype = mimetypes.guess_type(file_path)[0] or 'application/octet-stream'
            return send_file(file_path, mimetype=mimetype, as_attachment=False)
            
    except Exception as e:
        return f"Error: {str(e)}", 500



@app.context_processor
def inject_proposals():
    if 'user_type' in session and session['user_type'] == 'mahasiswa':
        if not hasattr(mysql, 'connection') or mysql.connection is None:
            return dict(proposals=[])
        try:
            cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
            cursor.execute('SELECT * FROM proposal WHERE nim = %s', (session['nim'],))
            proposals = cursor.fetchall()
            
            # Cek apakah ada proposal dengan status_admin = 'lolos'
            has_lolos_proposal = any(proposal.get('status_admin') == 'lolos' for proposal in proposals)
            
            cursor.close()
            return dict(proposals=proposals, has_lolos_proposal=has_lolos_proposal)
        except Exception:
            return dict(proposals=[], has_lolos_proposal=False)
    return dict(proposals=[], has_lolos_proposal=False)




# Monitoring Mahasiswa


def update_laporan_kemajuan_from_anggaran(cursor, proposal_id, tabel_laporan, tabel_anggaran):
    """
    Update laporan kemajuan dengan data anggaran terbaru yang disetujui
    """
    try:
        print(f"Debug: Starting update for proposal {proposal_id}, tabel_laporan: {tabel_laporan}, tabel_anggaran: {tabel_anggaran}")
        
        # Ambil data anggaran yang disetujui
        cursor.execute(f'''
            SELECT id, kegiatan_utama, kegiatan, penanggung_jawab, target_capaian, 
                   nama_barang, satuan, keterangan
            FROM {tabel_anggaran} 
            WHERE id_proposal = %s AND status = 'disetujui'
            ORDER BY kegiatan_utama, kegiatan, nama_barang
        ''', (proposal_id,))
        
        anggaran_data = cursor.fetchall()
        print(f"Debug: Found {len(anggaran_data)} approved anggaran records")
        
        # Hapus SEMUA data laporan kemajuan yang ada (tidak peduli status)
        cursor.execute(f'DELETE FROM {tabel_laporan} WHERE id_proposal = %s', (proposal_id,))
        deleted_count = cursor.rowcount
        print(f"Debug: Deleted {deleted_count} existing laporan kemajuan records")
        
        # Insert data baru dari anggaran yang disetujui
        inserted_count = 0
        for anggaran in anggaran_data:
            cursor.execute(f'''
                INSERT INTO {tabel_laporan} (
                    id_proposal, kegiatan_utama, kegiatan, penanggung_jawab, target_capaian,
                    nama_barang, kuantitas, satuan, harga_satuan, jumlah, keterangan, status
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            ''', (
                proposal_id, 
                anggaran['kegiatan_utama'], 
                anggaran['kegiatan'], 
                anggaran['penanggung_jawab'], 
                anggaran['target_capaian'], 
                anggaran['nama_barang'],
                0,  # kuantitas reset ke 0
                anggaran['satuan'], 
                0,  # harga_satuan reset ke 0
                0,  # jumlah reset ke 0
                '',  # keterangan selalu kosong
                'draf'  # status reset ke draf
            ))
            inserted_count += 1
        
        print(f"Debug: Inserted {inserted_count} new laporan kemajuan records for proposal {proposal_id}")
        
        # Commit perubahan
        cursor.connection.commit()
        
    except Exception as e:
        print(f"Error updating laporan kemajuan: {e}")
        raise e


@app.route('/test/update_laporan_kemajuan/<int:proposal_id>', methods=['POST'])
def test_update_laporan_kemajuan(proposal_id):
    """
    Endpoint untuk testing update laporan kemajuan
    """
    if not hasattr(mysql, 'connection') or mysql.connection is None:
        return jsonify({'success': False, 'message': 'Koneksi ke database gagal!'})
    
    try:
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        
        # Update laporan kemajuan awal
        update_laporan_kemajuan_from_anggaran(cursor, proposal_id, 'laporan_kemajuan_awal', 'anggaran_awal')
        
        # Update laporan kemajuan bertumbuh
        update_laporan_kemajuan_from_anggaran(cursor, proposal_id, 'laporan_kemajuan_bertumbuh', 'anggaran_bertumbuh')
        
        cursor.close()
        
        return jsonify({'success': True, 'message': f'Laporan kemajuan berhasil diupdate untuk proposal {proposal_id}!'})
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error: {str(e)}'})




# Fungsi untuk memperbarui data terkait ketika ada perubahan bahan baku
def update_related_data_on_bahan_baku_change(bahan_baku_id):
    """
    Memperbarui data produksi, produk_bahan_baku, dan laba rugi ketika ada perubahan bahan baku
    """
    if not hasattr(mysql, 'connection') or mysql.connection is None:
        print("Koneksi database tidak tersedia")
        return False
        
    try:
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        
        # Ambil semua produksi yang menggunakan bahan baku ini
        cursor.execute('''
            SELECT DISTINCT p.id as produksi_id, p.proposal_id, pbb.quantity_digunakan, pbb.harga_satuan
            FROM produk_bahan_baku pbb
            JOIN produksi p ON pbb.produksi_id = p.id
            WHERE pbb.bahan_baku_id = %s
        ''', (bahan_baku_id,))
        
        affected_produksi = cursor.fetchall()
        
        # Ambil data bahan baku yang diupdate
        cursor.execute('SELECT nama_bahan, harga_satuan, quantity FROM bahan_baku WHERE id = %s', (bahan_baku_id,))
        bahan_baku_data = cursor.fetchone()
        
        if not bahan_baku_data:
            return False
        
        # Update semua produk_bahan_baku yang menggunakan bahan baku ini
        for produksi in affected_produksi:
            produksi_id = produksi['produksi_id']
            proposal_id = produksi['proposal_id']
            old_quantity = produksi['quantity_digunakan']
            old_harga_satuan = produksi['harga_satuan']
            
            # Hitung subtotal baru
            new_subtotal = old_quantity * bahan_baku_data['harga_satuan']
            
            # Update produk_bahan_baku dengan harga satuan baru
            cursor.execute('''
                UPDATE produk_bahan_baku 
                SET harga_satuan = %s, subtotal = %s
                WHERE produksi_id = %s AND bahan_baku_id = %s
            ''', (bahan_baku_data['harga_satuan'], new_subtotal, produksi_id, bahan_baku_id))
            
            # Recalculate total biaya produksi
            cursor.execute('''
                SELECT SUM(subtotal) as total_biaya
                FROM produk_bahan_baku 
                WHERE produksi_id = %s
            ''', (produksi_id,))
            
            total_biaya_result = cursor.fetchone()
            new_total_biaya = total_biaya_result['total_biaya'] if total_biaya_result['total_biaya'] else 0
            
            # Ambil data produksi untuk recalculate
            cursor.execute('''
                SELECT jumlah_produk, persentase_laba
                FROM produksi 
                WHERE id = %s
            ''', (produksi_id,))
            
            produksi_data = cursor.fetchone()
            if produksi_data:
                jumlah_produk = produksi_data['jumlah_produk']
                persentase_laba = produksi_data['persentase_laba']
                
                # Hitung HPP baru
                new_hpp = new_total_biaya / jumlah_produk if jumlah_produk > 0 else 0
                
                # Hitung harga jual baru
                new_harga_jual = new_hpp + (persentase_laba / 100 * new_hpp)
                
                # Update produksi
                cursor.execute('''
                    UPDATE produksi 
                    SET total_biaya = %s, hpp = %s, harga_jual = %s
                    WHERE id = %s
                ''', (new_total_biaya, new_hpp, new_harga_jual, produksi_id))
                
                # Update laba rugi untuk proposal ini
                calculate_and_save_laba_rugi(proposal_id)
                
                # Update penjualan untuk produk ini
                try:
                    update_penjualan_on_produksi_change(proposal_id, produksi_data['nama_produk'])
                except Exception as e:
                    print(f"Warning: Error updating penjualan for produk {produksi_data['nama_produk']}: {str(e)}")
        
        mysql.connection.commit()
        cursor.close()
        return True
        
    except Exception as e:
        print(f"Error updating related data: {str(e)}")
        return False

# Fungsi untuk memperbarui laba rugi ketika ada perubahan produksi
def update_laba_rugi_on_produksi_change(proposal_id):
    """
    Memperbarui laba rugi dan penjualan ketika ada perubahan data produksi
    """
    if not hasattr(mysql, 'connection') or mysql.connection is None:
        print("Koneksi database tidak tersedia")
        return False
        
    try:
        # Update laba rugi
        calculate_and_save_laba_rugi(proposal_id)
        
        # Update penjualan
        try:
            update_penjualan_on_produksi_change(proposal_id)
        except Exception as e:
            print(f"Warning: Error updating penjualan: {str(e)}")
        
        return True
    except Exception as e:
        print(f"Error updating laba rugi: {str(e)}")
        return False

# Fungsi untuk memperbarui data penjualan ketika ada perubahan produksi
def update_penjualan_on_produksi_change(proposal_id, nama_produk=None):
    """
    Memperbarui data penjualan ketika ada perubahan data produksi
    """
    if not hasattr(mysql, 'connection') or mysql.connection is None:
        print("Koneksi database tidak tersedia")
        return False
        
    try:
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        
        # Ambil data penjualan yang perlu diupdate
        if nama_produk:
            # Update penjualan untuk produk tertentu
            cursor.execute('''
                SELECT id, quantity, harga_jual, total
                FROM penjualan 
                WHERE proposal_id = %s AND nama_produk = %s
            ''', (proposal_id, nama_produk))
        else:
            # Update semua penjualan untuk proposal ini
            cursor.execute('''
                SELECT id, nama_produk, quantity, harga_jual, total
                FROM penjualan 
                WHERE proposal_id = %s
            ''', (proposal_id,))
        
        penjualan_list = cursor.fetchall()
        
        for penjualan in penjualan_list:
            # Ambil harga jual terbaru dari produksi
            cursor.execute('''
                SELECT harga_jual 
                FROM produksi 
                WHERE proposal_id = %s AND nama_produk = %s
            ''', (proposal_id, penjualan['nama_produk']))
            
            produksi_data = cursor.fetchone()
            if produksi_data:
                new_harga_jual = produksi_data['harga_jual']
                new_total = penjualan['quantity'] * new_harga_jual
                
                # Update harga jual dan total di penjualan
                cursor.execute('''
                    UPDATE penjualan 
                    SET harga_jual = %s, total = %s
                    WHERE id = %s
                ''', (new_harga_jual, new_total, penjualan['id']))
        
        mysql.connection.commit()
        cursor.close()
        return True
        
    except Exception as e:
        print(f"Error updating penjualan: {str(e)}")
        return False

# Fungsi untuk memperbarui data penjualan ketika ada perubahan bahan baku
def update_penjualan_on_bahan_baku_change(bahan_baku_id):
    """
    Memperbarui data penjualan ketika ada perubahan data bahan baku
    """
    if not hasattr(mysql, 'connection') or mysql.connection is None:
        print("Koneksi database tidak tersedia")
        return False
        
    try:
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        
        # Ambil semua produksi yang menggunakan bahan baku ini
        cursor.execute('''
            SELECT DISTINCT p.proposal_id, p.nama_produk
            FROM produk_bahan_baku pbb
            JOIN produksi p ON pbb.produksi_id = p.id
            WHERE pbb.bahan_baku_id = %s
        ''', (bahan_baku_id,))
        
        affected_produksi = cursor.fetchall()
        
        # Update penjualan untuk setiap produksi yang terpengaruh
        for produksi in affected_produksi:
            update_penjualan_on_produksi_change(
                produksi['proposal_id'], 
                produksi['nama_produk']
            )
        
        cursor.close()
        return True
        
    except Exception as e:
        print(f"Error updating penjualan on bahan baku change: {str(e)}")
        return False


def generate_excel_neraca(mahasiswa_info, neraca_data):
    """
    Generate laporan neraca dalam format Excel
    """
    try:
        import openpyxl
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        from datetime import datetime
        
        # Buat workbook baru
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Laporan Neraca"
        
        # Header
        ws['A1'] = "LAPORAN NERACA"
        ws['A1'].font = Font(name='Times New Roman', size=12, bold=True)
        ws.merge_cells('A1:D1')
        ws['A1'].alignment = Alignment(horizontal='center')
        ws.row_dimensions[1].height = 15  # Kurangi tinggi baris
        
        ws['A2'] = f"{mahasiswa_info['judul_usaha']}"
        ws['A2'].font = Font(name='Times New Roman', size=12, bold=True)
        ws['A2'].alignment = Alignment(horizontal='center')
        ws.merge_cells('A2:D2')
        ws.row_dimensions[2].height = 15  # Kurangi tinggi baris
        
        ws['A3'] = f"Per Tanggal: {neraca_data.get('tanggal_neraca', datetime.now().strftime('%d %B %Y'))}"
        ws['A3'].font = Font(name='Times New Roman', size=12, bold=True)
        ws['A3'].alignment = Alignment(horizontal='center')
        ws.merge_cells('A3:D3')
        ws.row_dimensions[3].height = 15  # Kurangi tinggi baris
        
        # Header tabel dengan kolom DEBIT dan KREDIT
        row = 5
        ws[f'A{row}'] = "NAMA AKUN"
        ws[f'A{row}'].font = Font(name='Times New Roman', bold=True, size=12)
        ws[f'B{row}'] = "NOMINAL"
        ws[f'B{row}'].font = Font(name='Times New Roman', bold=True, size=12)
        ws[f'B{row}'].alignment = Alignment(horizontal='center')
        ws.merge_cells(f'B{row}:C{row}')
        
        row += 1
        ws[f'B{row}'] = "DEBIT"
        ws[f'B{row}'].font = Font(name='Times New Roman', bold=True, size=12)
        ws[f'B{row}'].alignment = Alignment(horizontal='center')
        ws[f'C{row}'] = "KREDIT"
        ws[f'C{row}'].font = Font(name='Times New Roman', bold=True, size=12)
        ws[f'C{row}'].alignment = Alignment(horizontal='center')
        
        # Buat tabel tunggal yang menyambung untuk Excel
        # A. ASET
        row += 1
        ws[f'A{row}'] = "A. ASET"
        ws[f'A{row}'].font = Font(name='Times New Roman', bold=True, size=10)
        ws.merge_cells(f'A{row}:D{row}')
        
        # 1. Aset Lancar
        row += 1
        ws[f'A{row}'] = "1. Aset Lancar"
        ws[f'A{row}'].font = Font(name='Times New Roman', bold=True, size=10)
        ws.merge_cells(f'A{row}:D{row}')
        
        row += 1
        ws[f'A{row}'] = "   Kas dan Setara Kas"
        ws[f'B{row}'] = f"Rp {neraca_data.get('kas_dan_setara_kas', 0):,.0f}"
        ws[f'B{row}'].alignment = Alignment(horizontal='right')
        
        row += 1
        ws[f'A{row}'] = "   Piutang Usaha"
        ws[f'B{row}'] = f"Rp {neraca_data.get('piutang_usaha', 0):,.0f}"
        ws[f'B{row}'].alignment = Alignment(horizontal='right')
        
        row += 1
        ws[f'A{row}'] = "   Persediaan"
        ws[f'B{row}'] = f"Rp {neraca_data.get('persediaan', 0):,.0f}"
        ws[f'B{row}'].alignment = Alignment(horizontal='right')
        
        row += 1
        ws[f'A{row}'] = "   Beban Dibayar di Muka"
        ws[f'B{row}'] = f"Rp {neraca_data.get('beban_dibayar_dimuka', 0):,.0f}"
        ws[f'B{row}'].alignment = Alignment(horizontal='right')
        
        row += 1
        ws[f'A{row}'] = "   dst."
        ws[f'B{row}'] = f"Rp {neraca_data.get('aset_lancar_lainnya', 0):,.0f}"
        ws[f'B{row}'].alignment = Alignment(horizontal='right')
        
        row += 1
        ws[f'A{row}'] = "Total Aset Lancar:"
        ws[f'A{row}'].font = Font(name='Times New Roman', bold=True, size=10)
        ws[f'B{row}'] = f"Rp {neraca_data.get('total_aset_lancar', 0):,.0f}"
        ws[f'B{row}'].font = Font(name='Times New Roman', bold=True, size=10)
        ws[f'B{row}'].alignment = Alignment(horizontal='right')
        
        # 2. Aset Tetap
        row += 1
        ws[f'A{row}'] = "2. Aset Tetap"
        ws[f'A{row}'].font = Font(name='Times New Roman', bold=True, size=10)
        ws.merge_cells(f'A{row}:D{row}')
        
        row += 1
        ws[f'A{row}'] = "   Tanah"
        ws[f'B{row}'] = f"Rp {neraca_data.get('tanah', 0):,.0f}"
        ws[f'B{row}'].alignment = Alignment(horizontal='right')
        
        row += 1
        ws[f'A{row}'] = "   Bangunan"
        ws[f'B{row}'] = f"Rp {neraca_data.get('bangunan', 0):,.0f}"
        ws[f'B{row}'].alignment = Alignment(horizontal='right')
        
        row += 1
        ws[f'A{row}'] = "   Kendaraan"
        ws[f'B{row}'] = f"Rp {neraca_data.get('kendaraan', 0):,.0f}"
        ws[f'B{row}'].alignment = Alignment(horizontal='right')
        
        row += 1
        ws[f'A{row}'] = "   Peralatan dan Mesin"
        ws[f'B{row}'] = f"Rp {neraca_data.get('peralatan_dan_mesin', 0):,.0f}"
        ws[f'B{row}'].alignment = Alignment(horizontal='right')
        
        row += 1
        ws[f'A{row}'] = "   Akumulasi Penyusutan"
        ws[f'B{row}'] = f"Rp {neraca_data.get('akumulasi_penyusutan', 0):,.0f}"
        ws[f'B{row}'].alignment = Alignment(horizontal='right')
        
        row += 1
        ws[f'A{row}'] = "Total Aset Tetap (Setelah Penyusutan):"
        ws[f'A{row}'].font = Font(name='Times New Roman', bold=True, size=10)
        ws[f'B{row}'] = f"Rp {neraca_data.get('total_aset_tetap_bersih', 0):,.0f}"
        ws[f'B{row}'].font = Font(name='Times New Roman', bold=True, size=10)
        ws[f'B{row}'].alignment = Alignment(horizontal='right')
        
        # 3. Aset Lain-lain
        row += 1
        ws[f'A{row}'] = "3. Aset Lain-lain"
        ws[f'A{row}'].font = Font(name='Times New Roman', bold=True, size=10)
        ws.merge_cells(f'A{row}:D{row}')
        
        row += 1
        ws[f'A{row}'] = "   Investasi Jangka Panjang"
        ws[f'B{row}'] = f"Rp {neraca_data.get('investasi_jangka_panjang', 0):,.0f}"
        ws[f'B{row}'].alignment = Alignment(horizontal='right')
        
        row += 1
        ws[f'A{row}'] = "   Hak Paten dan Merek Dagang"
        ws[f'B{row}'] = f"Rp {neraca_data.get('hak_paten_merek', 0):,.0f}"
        ws[f'B{row}'].alignment = Alignment(horizontal='right')
        
        row += 1
        ws[f'A{row}'] = "   dst."
        ws[f'B{row}'] = f"Rp {neraca_data.get('aset_lainnya', 0):,.0f}"
        ws[f'B{row}'].alignment = Alignment(horizontal='right')
        
        row += 1
        ws[f'A{row}'] = "Total Aset Lain-lain:"
        ws[f'A{row}'].font = Font(name='Times New Roman', bold=True, size=10)
        ws[f'B{row}'] = f"Rp {neraca_data.get('total_aset_lainnya', 0):,.0f}"
        ws[f'B{row}'].font = Font(name='Times New Roman', bold=True, size=10)
        ws[f'B{row}'].alignment = Alignment(horizontal='right')
        
        # Total Aset Keseluruhan
        row += 1
        ws[f'A{row}'] = "Total Aset Keseluruhan:"
        ws[f'A{row}'].font = Font(name='Times New Roman', bold=True, size=10, color="0000FF")
        ws[f'B{row}'] = f"Rp {neraca_data.get('total_aset', 0):,.0f}"
        ws[f'B{row}'].font = Font(name='Times New Roman', bold=True, size=10, color="0000FF")
        ws[f'B{row}'].alignment = Alignment(horizontal='right')
        
        # B. LIABILITAS
        row += 1
        ws[f'A{row}'] = "B. LIABILITAS (KEWAJIBAN)"
        ws[f'A{row}'].font = Font(name='Times New Roman', bold=True, size=10)
        ws.merge_cells(f'A{row}:D{row}')
        
        # 1. Liabilitas Jangka Pendek
        row += 1
        ws[f'A{row}'] = "1. Liabilitas Jangka Pendek"
        ws[f'A{row}'].font = Font(name='Times New Roman', bold=True, size=10)
        ws.merge_cells(f'A{row}:D{row}')
        
        row += 1
        ws[f'A{row}'] = "   Utang Usaha"
        ws[f'C{row}'] = f"Rp {neraca_data.get('utang_usaha', 0):,.0f}"
        ws[f'C{row}'].alignment = Alignment(horizontal='right')
        
        row += 1
        ws[f'A{row}'] = "   Beban yang Masih Harus Dibayar"
        ws[f'C{row}'] = f"Rp {neraca_data.get('beban_harus_dibayar', 0):,.0f}"
        ws[f'C{row}'].alignment = Alignment(horizontal='right')
        
        row += 1
        ws[f'A{row}'] = "   Utang Pajak"
        ws[f'C{row}'] = f"Rp {neraca_data.get('utang_pajak', 0):,.0f}"
        ws[f'C{row}'].alignment = Alignment(horizontal='right')
        
        row += 1
        ws[f'A{row}'] = "   Utang Jangka Pendek Lainnya"
        ws[f'C{row}'] = f"Rp {neraca_data.get('utang_jangka_pendek_lainnya', 0):,.0f}"
        ws[f'C{row}'].alignment = Alignment(horizontal='right')
        
        row += 1
        ws[f'A{row}'] = "Total Liabilitas Jangka Pendek"
        ws[f'A{row}'].font = Font(name='Times New Roman', bold=True, size=10)
        ws[f'C{row}'] = f"Rp {neraca_data.get('total_liabilitas_jangka_pendek', 0):,.0f}"
        ws[f'C{row}'].font = Font(name='Times New Roman', bold=True, size=10)
        ws[f'C{row}'].alignment = Alignment(horizontal='right')
        
        # 2. Liabilitas Jangka Panjang
        row += 1
        ws[f'A{row}'] = "2. Liabilitas Jangka Panjang"
        ws[f'A{row}'].font = Font(name='Times New Roman', bold=True, size=10)
        ws.merge_cells(f'A{row}:D{row}')
        
        row += 1
        ws[f'A{row}'] = "   Utang Bank"
        ws[f'C{row}'] = f"Rp {neraca_data.get('utang_bank', 0):,.0f}"
        ws[f'C{row}'].alignment = Alignment(horizontal='right')
        
        row += 1
        ws[f'A{row}'] = "   Obligasi yang Diterbitkan"
        ws[f'C{row}'] = f"Rp 0"
        ws[f'C{row}'].alignment = Alignment(horizontal='right')
        
        row += 1
        ws[f'A{row}'] = "   dst."
        ws[f'C{row}'] = f"Rp 0"
        ws[f'C{row}'].alignment = Alignment(horizontal='right')
        
        row += 1
        ws[f'A{row}'] = "Total Liabilitas Jangka Panjang"
        ws[f'A{row}'].font = Font(name='Times New Roman', bold=True, size=10)
        ws[f'C{row}'] = f"Rp {neraca_data.get('total_liabilitas_jangka_panjang', 0):,.0f}"
        ws[f'C{row}'].font = Font(name='Times New Roman', bold=True, size=10)
        ws[f'C{row}'].alignment = Alignment(horizontal='right')
        
        # Total Liabilitas
        row += 1
        ws[f'A{row}'] = "Total Liabilitas"
        ws[f'A{row}'].font = Font(name='Times New Roman', bold=True, size=11)
        ws[f'C{row}'] = f"Rp {neraca_data.get('total_liabilitas', 0):,.0f}"
        ws[f'C{row}'].font = Font(name='Times New Roman', bold=True, size=11)
        ws[f'C{row}'].alignment = Alignment(horizontal='right')
        
        # C. EKUITAS
        row += 1
        ws[f'A{row}'] = "C. EKUITAS"
        ws[f'A{row}'].font = Font(name='Times New Roman', bold=True, size=11)
        ws.merge_cells(f'A{row}:D{row}')
        
        row += 1
        ws[f'A{row}'] = "   Modal Saham"
        ws[f'C{row}'] = f"Rp {neraca_data.get('modal_disetor', 0):,.0f}"
        ws[f'C{row}'].alignment = Alignment(horizontal='right')
        
        row += 1
        ws[f'A{row}'] = "   Laba Ditahan"
        ws[f'C{row}'] = f"Rp {neraca_data.get('laba_ditahan', 0):,.0f}"
        ws[f'C{row}'].alignment = Alignment(horizontal='right')
        
        row += 1
        ws[f'A{row}'] = "   dst."
        ws[f'C{row}'] = f"Rp 0"
        ws[f'C{row}'].alignment = Alignment(horizontal='right')
        
        row += 1
        ws[f'A{row}'] = "Total Ekuitas"
        ws[f'A{row}'].font = Font(name='Times New Roman', bold=True, size=11)
        ws[f'C{row}'] = f"Rp {neraca_data.get('total_ekuitas', 0):,.0f}"
        ws[f'C{row}'].font = Font(name='Times New Roman', bold=True, size=11)
        ws[f'C{row}'].alignment = Alignment(horizontal='right')
        
        # Total Liabilitas dan Ekuitas
        row += 1
        ws[f'A{row}'] = "Total Liabilitas dan Ekuitas"
        ws[f'A{row}'].font = Font(name='Times New Roman', bold=True, size=11, color="0000FF")
        ws[f'C{row}'] = f"Rp {neraca_data.get('total_liabilitas_dan_ekuitas', 0):,.0f}"
        ws[f'C{row}'].font = Font(name='Times New Roman', bold=True, size=11, color="0000FF")
        ws[f'C{row}'].alignment = Alignment(horizontal='right')
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            # Cari cell pertama yang bukan MergedCell untuk mendapatkan column_letter
            column_letter = None
            for cell in column:
                if hasattr(cell, 'column_letter'):
                    column_letter = cell.column_letter
                    break
            
            if column_letter:
                for cell in column:
                    try:
                        if hasattr(cell, 'value') and cell.value and len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width
        
        # Simpan ke buffer langsung tanpa menyimpan ke file sistem
        import io
        excel_buffer = io.BytesIO()
        wb.save(excel_buffer)
        excel_buffer.seek(0)
        
        return excel_buffer
        
    except Exception as e:
        print(f"❌ Error generating Excel neraca: {e}")
        import traceback
        print(traceback.format_exc())
        return None


def generate_pdf_neraca(mahasiswa_info, neraca_data):
    """
    Generate laporan neraca dalam format PDF dengan kop surat
    """
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from datetime import datetime
        
        # Buat buffer untuk PDF
        import io
        pdf_buffer = io.BytesIO()
        
        doc = SimpleDocTemplate(pdf_buffer, pagesize=A4)
        story = []
        
        # Ambil header dari Kop Surat.pdf yang asli
        from utils import create_pdf_header_from_kop_surat_pdf
        header_template = create_pdf_header_from_kop_surat_pdf()
        header_data = header_template['header_data']
        logo_path = header_template['logo_path']
        shield_logo_path = header_template['shield_logo_path']
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1  # Center
        )
        
        # Header dengan kop surat horizontal - TANPA LOGO
        # Buat table untuk tata letak horizontal: [TEXT] (tanpa logo)
        
        # Buat header dengan styling link yang benar
        from utils import create_header_with_styled_links
        styled_headers = create_header_with_styled_links(header_data)
        
        # Tambahkan header yang sudah di-styling
        for header in styled_headers:
            story.append(header)
        
        story.append(Spacer(1, 15))
        story.append(Spacer(1, 20))
        
        # Judul laporan
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Heading1'],
            fontName='Times-Bold',
            fontSize=12,
            alignment=1,  # Center
            spaceAfter=0,
            spaceBefore=0
        )
        story.append(Paragraph("LAPORAN NERACA", title_style))
        
        # Judul usaha dengan alignment center
        judul_usaha_style = ParagraphStyle(
            'JudulUsaha',
            parent=styles['Heading2'],
            fontName='Times-Bold',
            fontSize=12,
            alignment=1,  # Center
            spaceAfter=0,
            spaceBefore=0
        )
        story.append(Paragraph(f"{mahasiswa_info['judul_usaha']}", judul_usaha_style))
        
        # Tanggal dengan alignment center
        tanggal_style = ParagraphStyle(
            'Tanggal',
            parent=styles['Heading2'],
            fontName='Times-Bold',
            fontSize=12,
            alignment=1,  # Center
            spaceAfter=12,
            spaceBefore=0
        )
        story.append(Paragraph(f"Per Tanggal: {neraca_data.get('tanggal_neraca', datetime.now().strftime('%d %B %Y'))}", tanggal_style))
        story.append(Spacer(1, 20))
        # Gabungkan semua data menjadi satu tabel besar
        all_data = []
        
        # Header
        all_data.extend([
            ['NAMA AKUN', 'NOMINAL', ''],
            ['', 'DEBIT', 'KREDIT']
        ])
        
        # ASET
        all_data.append(['A. ASET', '', ''])
        
        # Aset Lancar
        all_data.extend([
            ['1. Aset Lancar', '', ''],
            ['   Kas dan Setara Kas', f"Rp {neraca_data.get('kas_dan_setara_kas', 0):,.0f}", ''],
            ['   Piutang Usaha', f"Rp {neraca_data.get('piutang_usaha', 0):,.0f}", ''],
            ['   Persediaan', f"Rp {neraca_data.get('persediaan', 0):,.0f}", ''],
            ['   Beban Dibayar di Muka', f"Rp {neraca_data.get('beban_dibayar_dimuka', 0):,.0f}", ''],
            ['   dst.', f"Rp {neraca_data.get('aset_lancar_lainnya', 0):,.0f}", ''],
            ['Total Aset Lancar:', f"Rp {neraca_data.get('total_aset_lancar', 0):,.0f}", '']
        ])
        
        # Aset Tetap
        all_data.extend([
            ['2. Aset Tetap', '', ''],
            ['   Tanah', f"Rp {neraca_data.get('tanah', 0):,.0f}", ''],
            ['   Bangunan', f"Rp {neraca_data.get('bangunan', 0):,.0f}", ''],
            ['   Kendaraan', f"Rp {neraca_data.get('kendaraan', 0):,.0f}", ''],
            ['   Peralatan dan Mesin', f"Rp {neraca_data.get('peralatan_dan_mesin', 0):,.0f}", ''],
            ['   Akumulasi Penyusutan', f"Rp {neraca_data.get('akumulasi_penyusutan', 0):,.0f}", ''],
            ['Total Aset Tetap (Setelah Penyusutan):', f"Rp {neraca_data.get('total_aset_tetap_bersih', 0):,.0f}", '']
        ])
        
        # Aset Lain-lain
        all_data.extend([
            ['3. Aset Lain-lain', '', ''],
            ['   Investasi Jangka Panjang', f"Rp {neraca_data.get('investasi_jangka_panjang', 0):,.0f}", ''],
            ['   Hak Paten dan Merek Dagang', f"Rp {neraca_data.get('hak_paten_merek', 0):,.0f}", ''],
            ['   dst.', f"Rp {neraca_data.get('aset_lainnya', 0):,.0f}", ''],
            ['Total Aset Lain-lain:', f"Rp {neraca_data.get('total_aset_lainnya', 0):,.0f}", '']
        ])
        
        # Total Aset
        all_data.append(['Total Aset Keseluruhan:', f"Rp {neraca_data.get('total_aset', 0):,.0f}", ''])
        
        # LIABILITAS
        all_data.append(['B. LIABILITAS (KEWAJIBAN)', '', ''])
        
        # Liabilitas Jangka Pendek
        all_data.extend([
            ['1. Liabilitas Jangka Pendek', '', ''],
            ['   Utang Usaha', '', f"Rp {neraca_data.get('utang_usaha', 0):,.0f}"],
            ['   Beban yang Masih Harus Dibayar', '', f"Rp {neraca_data.get('beban_harus_dibayar', 0):,.0f}"],
            ['   Utang Pajak', '', f"Rp {neraca_data.get('utang_pajak', 0):,.0f}"],
            ['   Utang Jangka Pendek Lainnya', '', f"Rp {neraca_data.get('utang_jangka_pendek_lainnya', 0):,.0f}"],
            ['Total Liabilitas Jangka Pendek', '', f"Rp {neraca_data.get('total_liabilitas_jangka_pendek', 0):,.0f}"]
        ])
        
        # Liabilitas Jangka Panjang
        all_data.extend([
            ['2. Liabilitas Jangka Panjang', '', ''],
            ['   Utang Bank', '', f"Rp {neraca_data.get('utang_bank', 0):,.0f}"],
            ['   Obligasi yang Diterbitkan', '', f"Rp 0"],
            ['   dst.', '', f"Rp 0"],
            ['Total Liabilitas Jangka Panjang', '', f"Rp {neraca_data.get('total_liabilitas_jangka_panjang', 0):,.0f}"]
        ])
        
        # Total Liabilitas
        all_data.append(['Total Liabilitas', '', f"Rp {neraca_data.get('total_liabilitas', 0):,.0f}"])
        
        # EKUITAS
        all_data.append(['C. EKUITAS', '', ''])
        
        all_data.extend([
            ['   Modal Saham', '', f"Rp {neraca_data.get('modal_disetor', 0):,.0f}"],
            ['   Laba Ditahan', '', f"Rp {neraca_data.get('laba_ditahan', 0):,.0f}"],
            ['   dst.', '', f"Rp 0"],
            ['Total Ekuitas', '', f"Rp {neraca_data.get('total_ekuitas', 0):,.0f}"]
        ])
        
        # Total Liabilitas dan Ekuitas
        all_data.append(['Total Liabilitas dan Ekuitas', '', f"Rp {neraca_data.get('total_liabilitas_dan_ekuitas', 0):,.0f}"])
        
        # Buat tabel tunggal yang menyambung
        main_table = Table(all_data, colWidths=[3*inch, 2*inch, 1*inch])
        
        # Definisikan style untuk tabel
        table_style = [
            ('FONTNAME', (0, 0), (-1, -1), 'Times-Roman'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
            ('LEFTPADDING', (0, 0), (0, -1), 6),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('SPAN', (1, 0), (2, 0)),  # Merge NOMINAL di header
        ]
        
        # Tambahkan style khusus untuk baris-baris tertentu
        for i, row in enumerate(all_data):
            # Header (baris 0-1) - NAMA AKUN, NOMINAL, DEBIT, KREDIT
            if i <= 1:
                table_style.extend([
                    ('FONTNAME', (0, i), (-1, i), 'Times-Bold'),
                    ('FONTSIZE', (0, i), (-1, i), 12),
                    ('ALIGN', (1, i), (2, i), 'CENTER'),
                ])
            
            # Judul utama (A. ASET, B. LIABILITAS, C. EKUITAS)
            elif any(row[0].startswith(prefix) for prefix in ['A. ASET', 'B. LIABILITAS', 'C. EKUITAS']):
                table_style.extend([
                    ('FONTNAME', (0, i), (-1, i), 'Times-Bold'),
                    ('FONTSIZE', (0, i), (-1, i), 10),
                ])
            
            # Sub-judul (1. Aset Lancar, 2. Aset Tetap, dll)
            elif any(row[0].startswith(prefix) for prefix in ['1. ', '2. ', '3. ']):
                table_style.extend([
                    ('FONTNAME', (0, i), (-1, i), 'Times-Bold'),
                    ('FONTSIZE', (0, i), (-1, i), 10),
                ])
            
            # Total rows
            elif 'Total' in row[0]:
                table_style.extend([
                    ('FONTNAME', (0, i), (-1, i), 'Times-Bold'),
                    ('FONTSIZE', (0, i), (-1, i), 10),
                ])
                
                # Warna biru untuk total utama
                if 'Total Aset Keseluruhan' in row[0] or 'Total Liabilitas dan Ekuitas' in row[0]:
                    table_style.extend([
                        ('TEXTCOLOR', (0, i), (-1, i), colors.blue),
                    ])
        
        main_table.setStyle(TableStyle(table_style))
        story.append(main_table)
        
        def set_metadata(canvas, doc):
            canvas.setAuthor("AYMP Unjaya Neraca")
            canvas.setTitle("Laporan Neraca")
            
        doc.build(story, onFirstPage=set_metadata, onLaterPages=set_metadata)
        pdf_buffer.seek(0)
        
        return pdf_buffer
        
    except Exception as e:
        print(f"❌ Error generating PDF neraca: {e}")
        import traceback
        print(traceback.format_exc())
        return None


def generate_word_neraca(mahasiswa_info, neraca_data):
    """
    Generate laporan neraca dalam format Word dengan header asli dari Kop Surat.docx
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        import os
        from datetime import datetime
        
        # Buat dokumen baru
        doc = Document()
        
        # KOP SURAT - Header konsisten tanpa bergantung file eksternal
        
        # Baris 1: YAYASAN KARTIKA EKA PAKSI
        header1 = doc.add_paragraph()
        header1_run = header1.add_run("YAYASAN KARTIKA EKA PAKSI")
        header1_run.font.name = 'Times New Roman'
        header1_run.font.size = Pt(12)
        header1_run.bold = True
        header1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header1.paragraph_format.space_after = Pt(0)  # Kurangi spacing
        
        # Baris 2: UNIVERSITAS JENDERAL ACHMAD YANI YOGYAKARTA
        header2 = doc.add_paragraph()
        header2_run = header2.add_run("UNIVERSITAS JENDERAL ACHMAD YANI YOGYAKARTA")
        header2_run.font.name = 'Times New Roman'
        header2_run.font.size = Pt(12)
        header2_run.bold = True
        header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header2.paragraph_format.space_after = Pt(3)  # Spacing minimal
        
        # Baris 3: Alamat (tanpa spasi tambahan)
        header3 = doc.add_paragraph()
        header3_run = header3.add_run("Jl Siliwangi Ringroad Barat, Banyuraden, Gamping, Sleman, Yogyakarta 55293")
        header3_run.font.name = 'Times New Roman'
        header3_run.font.size = Pt(10)
        header3_run.bold = False
        header3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header3.paragraph_format.space_after = Pt(2)  # Spacing minimal
        
        # Baris 4: Telp, Fax, Website dengan link biru
        header4 = doc.add_paragraph()
        header4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header4.paragraph_format.space_after = Pt(0)  # Tidak ada spacing
        
        # Tambahkan teks sebelum link
        header4_run1 = header4.add_run("Telp. (0274) 552489, 552851 Fax. (0274) 557228 Website: ")
        header4_run1.font.name = 'Times New Roman'
        header4_run1.font.size = Pt(10)
        header4_run1.bold = False
        
        # Tambahkan link website dengan warna biru dan garis bawah
        header4_run2 = header4.add_run("www.unjaya.ac.id")
        header4_run2.font.name = 'Times New Roman'
        header4_run2.font.size = Pt(10)
        header4_run2.bold = False
        header4_run2.font.color.rgb = RGBColor(0, 0, 255)  # Warna biru
        header4_run2.font.underline = True
        
        # Baris 5: Email dengan link biru
        header5 = doc.add_paragraph()
        header5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header5.paragraph_format.space_after = Pt(6)  # Spacing sebelum garis
        
        # Tambahkan teks sebelum link
        header5_run1 = header5.add_run("Email: ")
        header5_run1.font.name = 'Times New Roman'
        header5_run1.font.size = Pt(10)
        header5_run1.bold = False
        
        # Tambahkan link email dengan warna biru dan garis bawah
        header5_run2 = header5.add_run("info@unjaya.ac.id")
        header5_run2.font.name = 'Times New Roman'
        header5_run2.font.size = Pt(10)
        header5_run2.bold = False
        header5_run2.font.color.rgb = RGBColor(0, 0, 255)  # Warna biru
        header5_run2.font.underline = True

        # Garis pemisah horizontal - lebih tipis dan mepet
        separator = doc.add_paragraph()
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        separator.paragraph_format.space_before = Pt(3)  # Kurangi spacing sebelum garis
        separator.paragraph_format.space_after = Pt(3)   # Kurangi spacing setelah garis
        
        # Set border bawah untuk paragraph separator
        p = separator._element
        pPr = p.get_or_add_pPr()
        pBdr = pPr.first_child_found_in("w:pBdr")
        if pBdr is None:
            from docx.oxml import parse_xml
            pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/></w:pBdr>')
            pPr.append(pBdr)

        # Title laporan
        title = doc.add_paragraph()
        title_run = title.add_run("LAPORAN NERACA")
        title_run.font.name = 'Times New Roman'
        title_run.font.size = Pt(12)
        title_run.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(0)
        title.paragraph_format.space_before = Pt(6)  # Spacing minimal dari garis
        
        # Informasi perusahaan
        judul = doc.add_paragraph()
        judul_run = judul.add_run(mahasiswa_info.get('judul_usaha', 'Nama Perusahaan'))
        judul_run.font.name = 'Times New Roman'
        judul_run.font.size = Pt(12)
        judul_run.bold = True
        judul.alignment = WD_ALIGN_PARAGRAPH.CENTER
        judul.paragraph_format.space_after = Pt(0)
        judul.paragraph_format.space_before = Pt(0)
        
        # Periode
        periode = doc.add_paragraph()
        periode_run = periode.add_run(f"Per Tanggal: {neraca_data.get('tanggal_neraca', datetime.now().strftime('%d %B %Y'))}")
        periode_run.font.name = 'Times New Roman'
        periode_run.font.size = Pt(12)
        periode_run.bold = True
        periode.alignment = WD_ALIGN_PARAGRAPH.CENTER
        periode.paragraph_format.space_after = Pt(6)  # Kurangi spacing
        periode.paragraph_format.space_before = Pt(0)
        
        # Spacer minimal
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(3)  # Spacing minimal
        
        # Buat tabel tunggal yang menyambung untuk Word
        # Hitung total baris yang dibutuhkan
        total_rows = 2  # Header
        total_rows += 1  # A. ASET
        total_rows += 7  # Aset Lancar (1. + 5 item + total)
        total_rows += 7  # Aset Tetap (2. + 5 item + total)
        total_rows += 5  # Aset Lain-lain (3. + 3 item + total)
        total_rows += 1  # Total Aset Keseluruhan
        total_rows += 1  # B. LIABILITAS
        total_rows += 6  # Liabilitas Jangka Pendek (1. + 4 item + total)
        total_rows += 5  # Liabilitas Jangka Panjang (2. + 3 item + total)
        total_rows += 1  # Total Liabilitas
        total_rows += 1  # C. EKUITAS
        total_rows += 4  # Ekuitas (3 item + total)
        total_rows += 1  # Total Liabilitas dan Ekuitas
        
        # Buat tabel utama
        main_table = doc.add_table(rows=total_rows, cols=3)
        main_table.style = 'Table Grid'
        main_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Header
        main_table.cell(0, 0).text = "NAMA AKUN"
        main_table.cell(0, 1).text = "NOMINAL"
        main_table.cell(0, 2).text = ""
        main_table.cell(1, 0).text = ""
        main_table.cell(1, 1).text = "DEBIT"
        main_table.cell(1, 2).text = "KREDIT"
        
        # Merge NOMINAL cell
        main_table.cell(0, 1).merge(main_table.cell(0, 2))
        
        # Isi data tabel
        row_idx = 2
        
        # A. ASET
        main_table.cell(row_idx, 0).text = "A. ASET"
        row_idx += 1
        
        # 1. Aset Lancar
        main_table.cell(row_idx, 0).text = "1. Aset Lancar"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "   Kas dan Setara Kas"
        main_table.cell(row_idx, 1).text = f"Rp {neraca_data.get('kas_dan_setara_kas', 0):,}"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "   Piutang Usaha"
        main_table.cell(row_idx, 1).text = f"Rp {neraca_data.get('piutang_usaha', 0):,}"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "   Persediaan"
        main_table.cell(row_idx, 1).text = f"Rp {neraca_data.get('persediaan', 0):,}"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "   Beban Dibayar di Muka"
        main_table.cell(row_idx, 1).text = f"Rp {neraca_data.get('beban_dibayar_dimuka', 0):,}"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "   dst."
        main_table.cell(row_idx, 1).text = f"Rp {neraca_data.get('aset_lancar_lainnya', 0):,}"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "Total Aset Lancar:"
        main_table.cell(row_idx, 1).text = f"Rp {neraca_data.get('total_aset_lancar', 0):,}"
        row_idx += 1
        
        # 2. Aset Tetap
        main_table.cell(row_idx, 0).text = "2. Aset Tetap"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "   Tanah"
        main_table.cell(row_idx, 1).text = f"Rp {neraca_data.get('tanah', 0):,}"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "   Bangunan"
        main_table.cell(row_idx, 1).text = f"Rp {neraca_data.get('bangunan', 0):,}"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "   Kendaraan"
        main_table.cell(row_idx, 1).text = f"Rp {neraca_data.get('kendaraan', 0):,}"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "   Peralatan dan Mesin"
        main_table.cell(row_idx, 1).text = f"Rp {neraca_data.get('peralatan_dan_mesin', 0):,}"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "   Akumulasi Penyusutan"
        main_table.cell(row_idx, 1).text = f"Rp {neraca_data.get('akumulasi_penyusutan', 0):,}"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "Total Aset Tetap (Setelah Penyusutan):"
        main_table.cell(row_idx, 1).text = f"Rp {neraca_data.get('total_aset_tetap_bersih', 0):,}"
        row_idx += 1
        
        # 3. Aset Lain-lain
        main_table.cell(row_idx, 0).text = "3. Aset Lain-lain"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "   Investasi Jangka Panjang"
        main_table.cell(row_idx, 1).text = f"Rp {neraca_data.get('investasi_jangka_panjang', 0):,}"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "   Hak Paten dan Merek Dagang"
        main_table.cell(row_idx, 1).text = f"Rp {neraca_data.get('hak_paten_merek', 0):,}"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "   dst."
        main_table.cell(row_idx, 1).text = f"Rp {neraca_data.get('aset_lainnya', 0):,}"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "Total Aset Lain-lain:"
        main_table.cell(row_idx, 1).text = f"Rp {neraca_data.get('total_aset_lainnya', 0):,}"
        row_idx += 1
        
        # Total Aset Keseluruhan
        main_table.cell(row_idx, 0).text = "Total Aset Keseluruhan:"
        main_table.cell(row_idx, 1).text = f"Rp {neraca_data.get('total_aset', 0):,}"
        row_idx += 1
        
        # B. LIABILITAS
        main_table.cell(row_idx, 0).text = "B. LIABILITAS (KEWAJIBAN)"
        row_idx += 1
        
        # 1. Liabilitas Jangka Pendek
        main_table.cell(row_idx, 0).text = "1. Liabilitas Jangka Pendek"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "   Utang Usaha"
        main_table.cell(row_idx, 2).text = f"Rp {neraca_data.get('utang_usaha', 0):,}"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "   Beban yang Masih Harus Dibayar"
        main_table.cell(row_idx, 2).text = f"Rp {neraca_data.get('beban_harus_dibayar', 0):,}"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "   Utang Pajak"
        main_table.cell(row_idx, 2).text = f"Rp {neraca_data.get('utang_pajak', 0):,}"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "   Utang Jangka Pendek Lainnya"
        main_table.cell(row_idx, 2).text = f"Rp {neraca_data.get('utang_jangka_pendek_lainnya', 0):,}"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "Total Liabilitas Jangka Pendek"
        main_table.cell(row_idx, 2).text = f"Rp {neraca_data.get('total_liabilitas_jangka_pendek', 0):,}"
        row_idx += 1
        
        # 2. Liabilitas Jangka Panjang
        main_table.cell(row_idx, 0).text = "2. Liabilitas Jangka Panjang"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "   Utang Bank"
        main_table.cell(row_idx, 2).text = f"Rp {neraca_data.get('utang_bank', 0):,}"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "   Obligasi yang Diterbitkan"
        main_table.cell(row_idx, 2).text = f"Rp 0"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "   dst."
        main_table.cell(row_idx, 2).text = f"Rp 0"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "Total Liabilitas Jangka Panjang"
        main_table.cell(row_idx, 2).text = f"Rp {neraca_data.get('total_liabilitas_jangka_panjang', 0):,}"
        row_idx += 1
        
        # Total Liabilitas
        main_table.cell(row_idx, 0).text = "Total Liabilitas"
        main_table.cell(row_idx, 2).text = f"Rp {neraca_data.get('total_liabilitas', 0):,}"
        row_idx += 1
        
        # C. EKUITAS
        main_table.cell(row_idx, 0).text = "C. EKUITAS"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "   Modal Saham"
        main_table.cell(row_idx, 2).text = f"Rp {neraca_data.get('modal_disetor', 0):,}"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "   Laba Ditahan"
        main_table.cell(row_idx, 2).text = f"Rp {neraca_data.get('laba_ditahan', 0):,}"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "   dst."
        main_table.cell(row_idx, 2).text = f"Rp 0"
        row_idx += 1
        main_table.cell(row_idx, 0).text = "Total Ekuitas"
        main_table.cell(row_idx, 2).text = f"Rp {neraca_data.get('total_ekuitas', 0):,}"
        row_idx += 1
        
        # Total Liabilitas dan Ekuitas
        main_table.cell(row_idx, 0).text = "Total Liabilitas dan Ekuitas"
        main_table.cell(row_idx, 2).text = f"Rp {neraca_data.get('total_liabilitas_dan_ekuitas', 0):,}"
        
        # Styling tabel
        for i, row in enumerate(main_table.rows):
            for j, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)  # Default font size untuk isi tabel
                        
                        # Header (baris 0-1) - NAMA AKUN, NOMINAL, DEBIT, KREDIT
                        if i <= 1:
                            run.font.size = Pt(12)
                            run.bold = True
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Judul utama (A. ASET, B. LIABILITAS, C. EKUITAS)
                        elif any(cell.text.startswith(prefix) for prefix in ['A. ASET', 'B. LIABILITAS', 'C. EKUITAS']):
                            run.font.size = Pt(10)
                            run.bold = True
                        
                        # Sub-judul (1. Aset Lancar, 2. Aset Tetap, dll)
                        elif any(cell.text.startswith(prefix) for prefix in ['1. ', '2. ', '3. ']):
                            run.font.size = Pt(10)
                            run.bold = True
                        
                        # Total rows
                        elif 'Total' in cell.text:
                            run.font.size = Pt(10)
                            run.bold = True
        
                            # Warna biru untuk total utama
                            if 'Total Aset Keseluruhan' in cell.text or 'Total Liabilitas dan Ekuitas' in cell.text:
                                run.font.color.rgb = RGBColor(0, 0, 255)
                        
                        # Alignment untuk kolom nominal
                        if (j == 1 or j == 2) and cell.text and cell.text.startswith('Rp'):
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Simpan ke buffer
        import io
        word_buffer = io.BytesIO()
        doc.save(word_buffer)
        word_buffer.seek(0)
        return word_buffer
    except Exception as e:
        print(f"❌ Error generate_word_neraca (header asli): {e}")
        import traceback
        print(traceback.format_exc())
        return None


# Import dan registrasi blueprint mahasiswa di akhir untuk menghindari circular import
from mahasiswa import mahasiswa_bp
from pembimbing import pembimbing_bp
from admin import admin_bp
from reviewer import reviewer_bp

app.register_blueprint(mahasiswa_bp, url_prefix='/mahasiswa')
app.register_blueprint(pembimbing_bp, url_prefix='/pembimbing')
app.register_blueprint(admin_bp, url_prefix='/admin')
app.register_blueprint(reviewer_bp, url_prefix='/reviewer')

# Global error handler untuk menampilkan error di terminal
@app.errorhandler(Exception)
def handle_exception(e):
    import traceback
    print("=" * 50)
    print("GLOBAL EXCEPTION HANDLER")
    print("=" * 50)
    print(f"Error: {str(e)}")
    print("Traceback:")
    traceback.print_exc()
    print("=" * 50)
    return "Internal Server Error", 500

# Error handler untuk 404
@app.errorhandler(404)
def not_found_error(error):
    print(f"404 Error: {request.url}")
    return "Page not found", 404

# Error handler untuk 500
@app.errorhandler(500)
def internal_error(error):
    import traceback
    print("500 Internal Server Error:")
    traceback.print_exc()
    return "Internal Server Error", 500

def hitung_neraca_real_time(proposal_id, cursor):
    """
    Fungsi helper untuk menghitung laporan neraca secara real-time
    dari data transaksi terbaru sesuai spesifikasi database
    """
    try:
        # 1. HITUNG ASET LANCAR
        # Kas dan Setara Kas - dari total_penjualan database arus_kas
        cursor.execute('''
            SELECT COALESCE(SUM(total_penjualan), 0) as total_penjualan
            FROM arus_kas 
            WHERE proposal_id = %s
        ''', (proposal_id,))
        result = cursor.fetchone()
        kas_dan_setara_kas = Decimal(str(result['total_penjualan'] or 0))
        
        # Piutang Usaha - hasilnya Rp.0
        piutang_usaha = Decimal('0')
        
        # Persediaan - total harga bahan baku yang tidak digunakan di produk_bahan_baku selama 1 bulan
        # Hitung total harga bahan baku
        cursor.execute('''
            SELECT COALESCE(SUM(total_harga), 0) as total_harga_bahan_baku
            FROM bahan_baku 
            WHERE proposal_id = %s
        ''', (proposal_id,))
        result = cursor.fetchone()
        total_harga_bahan_baku = Decimal(str(result['total_harga_bahan_baku'] or 0))
        
        # Hitung total harga bahan baku yang digunakan dalam produksi
        cursor.execute('''
            SELECT COALESCE(SUM(pbb.subtotal), 0) as total_harga_digunakan
            FROM produk_bahan_baku pbb
            JOIN produksi p ON pbb.produksi_id = p.id
            WHERE p.proposal_id = %s
        ''', (proposal_id,))
        result = cursor.fetchone()
        total_harga_digunakan = Decimal(str(result['total_harga_digunakan'] or 0))
        
        persediaan = total_harga_bahan_baku - total_harga_digunakan
        if persediaan < 0:
            persediaan = Decimal('0')
        
        # Beban Dibayar Dimuka - hasilnya Rp.0
        beban_dibayar_dimuka = Decimal('0')
        
        total_aset_lancar = kas_dan_setara_kas + piutang_usaha + persediaan + beban_dibayar_dimuka
        
        # 2. HITUNG ASET TETAP
        # Tanah - hasilnya Rp.0
        tanah = Decimal('0')
        
        # Bangunan - hasilnya Rp.0
        bangunan = Decimal('0')
        
        # Kendaraan - hasilnya Rp.0
        kendaraan = Decimal('0')
        
        # Peralatan dan Mesin - total dari total_alat_produksi
        cursor.execute('''
            SELECT COALESCE(SUM(total_alat_produksi), 0) as total_peralatan_mesin
            FROM alat_produksi 
            WHERE proposal_id = %s
        ''', (proposal_id,))
        result = cursor.fetchone()
        peralatan_dan_mesin = Decimal(str(result['total_peralatan_mesin'] or 0))
        
        # Akumulasi Penyusutan - 10% dari total aset tetap per tahun
        total_aset_tetap_sebelum_penyusutan = tanah + bangunan + kendaraan + peralatan_dan_mesin
        akumulasi_penyusutan = total_aset_tetap_sebelum_penyusutan * Decimal('0.1')  # 10% per tahun
        total_aset_tetap_bersih = total_aset_tetap_sebelum_penyusutan - akumulasi_penyusutan
        
        # 3. HITUNG ASET LAIN-LAIN
        # Investasi Jangka Panjang - hasilnya Rp.0
        investasi_jangka_panjang = Decimal('0')
        
        # Hak Paten dan Merek Dagang - hasilnya NIB
        cursor.execute('''
            SELECT COALESCE(nib, 0) as hak_paten_merek_dagang
            FROM proposal 
            WHERE id = %s
        ''', (proposal_id,))
        result = cursor.fetchone()
        hak_paten_merek_dagang = Decimal(str(result['hak_paten_merek_dagang'] if result and result['hak_paten_merek_dagang'] else 0))
        
        total_aset_lain_lain = investasi_jangka_panjang + hak_paten_merek_dagang
        
        # Total Aset Keseluruhan
        total_aset = total_aset_lancar + total_aset_tetap_bersih + total_aset_lain_lain
        
        # 4. HITUNG LIABILITAS (KEWAJIBAN) - hasilnya Rp.0 semua
        utang_usaha = Decimal('0')
        beban_harus_dibayar = Decimal('0')
        utang_pajak = Decimal('0')
        utang_jangka_pendek_lainnya = Decimal('0')
        total_liabilitas_jangka_pendek = utang_usaha + beban_harus_dibayar + utang_pajak + utang_jangka_pendek_lainnya
        
        utang_bank = Decimal('0')
        obligasi_diterbitkan = Decimal('0')
        total_liabilitas_jangka_panjang = utang_bank + obligasi_diterbitkan
        
        total_liabilitas = total_liabilitas_jangka_pendek + total_liabilitas_jangka_panjang
        
        # 5. HITUNG EKUITAS - hasilnya Rp.0 semua
        modal_saham = Decimal('0')
        laba_ditahan = Decimal('0')
        modal_disetor = Decimal('0')
        total_ekuitas = modal_saham + laba_ditahan + modal_disetor
        
        # Total Liabilitas dan Ekuitas
        total_liabilitas_dan_ekuitas = total_liabilitas + total_ekuitas
        
        # Simpan data ke tabel laporan_neraca
        try:
            # Hapus data lama jika ada
            cursor.execute('DELETE FROM laporan_neraca WHERE proposal_id = %s', (proposal_id,))
            
            # Insert data baru dengan field mapping yang benar
            cursor.execute('''
                INSERT INTO laporan_neraca (
                    proposal_id, tanggal_neraca, kas_setara_kas, piutang_usaha, persediaan,
                    beban_dibayar_dimuka, total_aset_lancar, tanah, bangunan, kendaraan,
                    peralatan_mesin, akumulasi_penyusutan, total_aset_tetap, investasi_jangka_panjang,
                    hak_paten_merek, total_aset_lain, total_aset_keseluruhan, utang_usaha,
                    beban_harus_dibayar, utang_pajak, utang_jangka_pendek_lain, total_liabilitas_jangka_pendek,
                    utang_bank, obligasi_diterbitkan, total_liabilitas_jangka_panjang, total_liabilitas,
                    modal_saham, laba_ditahan, modal_disetor, total_ekuitas, total_liabilitas_ekuitas
                ) VALUES (
                    %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s
                )
            ''', (
                proposal_id, datetime.now().date(), kas_dan_setara_kas, piutang_usaha, persediaan,
                beban_dibayar_dimuka, total_aset_lancar, tanah, bangunan, kendaraan,
                peralatan_dan_mesin, akumulasi_penyusutan, total_aset_tetap_bersih, investasi_jangka_panjang,
                hak_paten_merek_dagang, total_aset_lain_lain, total_aset, utang_usaha,
                beban_harus_dibayar, utang_pajak, utang_jangka_pendek_lainnya, total_liabilitas_jangka_pendek,
                utang_bank, obligasi_diterbitkan, total_liabilitas_jangka_panjang, total_liabilitas,
                modal_saham, laba_ditahan, modal_disetor, total_ekuitas, total_liabilitas_dan_ekuitas
            ))
            
            # Commit perubahan
            cursor.connection.commit()
            print(f"Berhasil menyimpan laporan neraca untuk proposal_id: {proposal_id}")
            
        except Exception as e:
            print(f"Error saving to laporan_neraca table: {e}")
            # Continue even if save fails
        
        return {
            'tanggal_neraca': datetime.now().strftime('%d %B %Y'),
            'kas_dan_setara_kas': float(kas_dan_setara_kas),
            'piutang_usaha': float(piutang_usaha),
            'persediaan': float(persediaan),
            'beban_dibayar_dimuka': float(beban_dibayar_dimuka),
            'total_aset_lancar': float(total_aset_lancar),
            'tanah': float(tanah),
            'bangunan': float(bangunan),
            'kendaraan': float(kendaraan),
            'peralatan_dan_mesin': float(peralatan_dan_mesin),
            'akumulasi_penyusutan': float(akumulasi_penyusutan),
            'total_aset_tetap_bersih': float(total_aset_tetap_bersih),
            'investasi_jangka_panjang': float(investasi_jangka_panjang),
            'hak_paten_merek_dagang': float(hak_paten_merek_dagang),
            'total_aset_lain_lain': float(total_aset_lain_lain),
            'total_aset': float(total_aset),
            'utang_usaha': float(utang_usaha),
            'beban_harus_dibayar': float(beban_harus_dibayar),
            'utang_pajak': float(utang_pajak),
            'utang_jangka_pendek_lainnya': float(utang_jangka_pendek_lainnya),
            'total_liabilitas_jangka_pendek': float(total_liabilitas_jangka_pendek),
            'utang_bank': float(utang_bank),
            'total_liabilitas_jangka_panjang': float(total_liabilitas_jangka_panjang),
            'total_liabilitas': float(total_liabilitas),
            'modal_disetor': float(modal_disetor),
            'laba_ditahan': float(laba_ditahan),
            'total_ekuitas': float(total_ekuitas),
            'total_liabilitas_dan_ekuitas': float(total_liabilitas_dan_ekuitas)
        }
        
    except Exception as e:
        print(f"Error in hitung_neraca_real_time: {e}")
        # Return default values if error
        return {
            'tanggal_neraca': datetime.now().strftime('%d %B %Y'),
            'kas_dan_setara_kas': 0.0, 'piutang_usaha': 0.0, 'persediaan': 0.0, 'beban_dibayar_dimuka': 0.0,
            'total_aset_lancar': 0.0, 'tanah': 0.0, 'bangunan': 0.0, 'kendaraan': 0.0, 'peralatan_dan_mesin': 0.0,
            'akumulasi_penyusutan': 0.0, 'total_aset_tetap_bersih': 0.0, 'investasi_jangka_panjang': 0.0,
            'hak_paten_merek_dagang': 0.0, 'total_aset_lain_lain': 0.0, 'total_aset': 0.0, 'utang_usaha': 0.0,
            'beban_harus_dibayar': 0.0, 'utang_pajak': 0.0, 'utang_jangka_pendek_lainnya': 0.0,
            'total_liabilitas_jangka_pendek': 0.0, 'utang_bank': 0.0,
            'total_liabilitas_jangka_panjang': 0.0, 'total_liabilitas': 0.0, 'modal_disetor': 0.0,
            'laba_ditahan': 0.0, 'total_ekuitas': 0.0, 'total_liabilitas_dan_ekuitas': 0.0
        }

# Global variable untuk tracking worker status
auto_fill_worker_running = False
auto_fill_worker_thread = None

# ✅ PERBAIKAN: Auto-start worker saat aplikasi Flask dimulai
def start_auto_fill_worker_on_startup():
    """Memulai worker auto-fill secara otomatis saat aplikasi Flask dimulai"""
    global auto_fill_worker_running, auto_fill_worker_thread
    
    try:
        # ✅ PERBAIKAN BARU: Cek apakah auto-fill dinonaktifkan untuk data lama
        if AUTO_FILL_CONFIG.get('DISABLE_AUTO_FILL_FOR_OLD_DATA', False):
            print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] AUTO-FILL DISABLED: Worker auto-fill dinonaktifkan untuk data lama")
            return
        
        if not auto_fill_worker_running:
            print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Memulai worker auto-fill secara otomatis...")
            auto_fill_worker_running = True
            auto_fill_worker_thread = threading.Thread(target=auto_fill_worker, daemon=True)
            auto_fill_worker_thread.start()
            print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Worker auto-fill berhasil dimulai!")
        else:
            print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Worker auto-fill sudah berjalan")
    except Exception as e:
        print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Error memulai worker auto-fill: {e}")

# Auto-fill worker yang berjalan otomatis di background
def auto_fill_worker():
    """
    Worker yang berjalan otomatis untuk mendeteksi dan mengisi sesi terlewat
    sesuai dengan jadwal yang telah ditentukan
    """
    global auto_fill_worker_running
    
    # Setup console logging sederhana tanpa file
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Worker auto-fill berhasil dimulai (hanya console logging)")
    
    # Timer untuk pengecekan prioritas
    last_priority_check = time.time()
    
    # Helper function untuk console logging
    def log_info(message):
        print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] INFO: {message}")
    
    def log_error(message):
        print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] ERROR: {message}")
    
    while auto_fill_worker_running:  # Gunakan global variable untuk kontrol
        try:
            current_time = time.time()
            
            # Pengecekan prioritas setiap 5 detik untuk sesi yang sangat urgent
            if AUTO_FILL_CONFIG['ENABLE_SMART_CHECKING'] and (current_time - last_priority_check) >= AUTO_FILL_CONFIG['PRIORITY_CHECK_INTERVAL']:
                log_info("Pengecekan prioritas: Memeriksa sesi yang sangat urgent...")
                check_urgent_sessions(log_info, log_error)
                last_priority_check = current_time
            
            log_info("Pengecekan rutin: Memeriksa sesi penilaian yang terlewat...")
            
            # Koneksi database menggunakan konfigurasi
            conn = MySQLdb.connect(
                host=AUTO_FILL_CONFIG['DB_HOST'],
                user=AUTO_FILL_CONFIG['DB_USER'],
                passwd=AUTO_FILL_CONFIG['DB_PASSWD'],
                db=AUTO_FILL_CONFIG['DB_NAME']
            )
            cursor = conn.cursor()
            
            try:
                # ✅ PERBAIKAN BARU: Auto-fill ENABLED untuk data baru
                if not AUTO_FILL_CONFIG.get('AUTO_FILL_ENABLED', True):
                    log_info("AUTO-FILL DISABLED: Auto-fill dinonaktifkan sepenuhnya")
                    time.sleep(60)  # Tunggu 1 menit sebelum cek lagi
                    continue
                
                # Ambil semua pengaturan jadwal yang aktif atau baru selesai (dalam 24 jam terakhir)
                cursor.execute("""
                    SELECT DISTINCT 
                        pj.id,
                        pj.pembimbing_nilai_mulai as waktu_mulai,
                        pj.pembimbing_nilai_selesai as waktu_selesai,
                        pj.pembimbing_interval_tipe as tipe_interval,
                        pj.pembimbing_interval_nilai as nilai_interval,
                        pj.pembimbing_hari_aktif as hari_aktif
                    FROM pengaturan_jadwal pj
                    WHERE pj.pembimbing_nilai_selesai >= DATE_SUB(NOW(), INTERVAL 24 HOUR)
                    AND pj.pembimbing_nilai_mulai IS NOT NULL
                    AND pj.pembimbing_nilai_selesai IS NOT NULL
                """)
                
                jadwal_list = cursor.fetchall()
                log_info(f"Ditemukan {len(jadwal_list)} jadwal penilaian yang aktif")
                
                # Proses dalam batch untuk performa lebih baik
                batch_size = AUTO_FILL_CONFIG['MAX_BATCH_SIZE']
                for i in range(0, len(jadwal_list), batch_size):
                    batch = jadwal_list[i:i + batch_size]
                    log_info(f"Memproses batch {i//batch_size + 1} ({len(batch)} jadwal)")
                    
                    for jadwal in batch:
                        jadwal_id = jadwal[0]
                        waktu_mulai = jadwal[1]
                        waktu_selesai = jadwal[2]
                        tipe_interval = jadwal[3]
                        nilai_interval = jadwal[4]
                        hari_aktif = jadwal[5]
                        
                        log_info(f"Memproses jadwal ID {jadwal_id}: Interval {tipe_interval} setiap {nilai_interval}")
                        
                        # Buat metadata kolom otomatis jika belum ada
                        create_metadata_kolom_for_active_schedule(
                            cursor, jadwal_id, tipe_interval, nilai_interval, 
                            waktu_mulai, waktu_selesai, log_info, log_error
                        )
                        
                        # Hitung sesi berdasarkan waktu saat ini
                        current_sesi_waktu = calculate_sesi_by_time_worker(
                            waktu_mulai, waktu_selesai, tipe_interval, 
                            nilai_interval, hari_aktif
                        )
                        
                        log_info(f"Sesi saat ini berdasarkan waktu: {current_sesi_waktu}")
                        
                        if current_sesi_waktu > 1:
                            # Cek apakah ada sesi yang terlewat untuk semua pembimbing-mahasiswa
                            auto_fill_all_pembimbing_mahasiswa_worker(
                                cursor, jadwal_id, current_sesi_waktu, waktu_mulai, waktu_selesai, 
                                tipe_interval, nilai_interval, hari_aktif, log_info, log_error
                            )
                        
                        # NEW: Auto-fill penilaian proposal, laporan kemajuan, dan laporan akhir
                        # Buat cursor baru dengan DictCursor untuk worker baru
                        cursor_dict = conn.cursor(MySQLdb.cursors.DictCursor)
                        try:
                            auto_fill_penilaian_proposal_worker(cursor_dict, log_info, log_error)
                            auto_fill_penilaian_laporan_kemajuan_worker(cursor_dict, log_info, log_error)
                            auto_fill_penilaian_laporan_akhir_worker(cursor_dict, log_info, log_error)
                            # NEW: Auto-update status mahasiswa dan proposal menjadi 'selesai'
                            auto_update_status_selesai_worker(cursor_dict, log_info, log_error)
                            
                            # NEW: Auto-reject anggaran, laporan kemajuan, dan laporan akhir berdasarkan jadwal
                            auto_reject_anggaran_worker(cursor_dict, log_info, log_error)
                            auto_reject_laporan_kemajuan_worker(cursor_dict, log_info, log_error)
                            auto_reject_laporan_akhir_worker(cursor_dict, log_info, log_error)
                        finally:
                            cursor_dict.close()
                
                conn.commit()
                log_info(f"Berhasil menyelesaikan pengecekan {len(jadwal_list)} jadwal")
                
                # ✅ PERBAIKAN: Cek apakah semua sesi sudah terisi untuk semua jadwal
                if AUTO_FILL_CONFIG.get('ENABLE_SMART_CHECKING', True):
                    all_sessions_completed = check_all_sessions_completed(cursor, jadwal_list, log_info, log_error)
                    if all_sessions_completed:
                        log_info("Semua sesi penilaian sudah selesai! Worker auto-fill akan berhenti otomatis.")
                        log_info("Untuk memulai ulang auto-fill, restart aplikasi Flask.")
                        break  # ✅ Berhenti otomatis jika semua sesi sudah terisi
                    else:
                        log_info("Masih ada sesi yang belum selesai atau jadwal masih aktif, worker tetap berjalan")
                
            except Exception as e:
                log_error(f"Error database di worker auto-fill: {e}")
                conn.rollback()
            finally:
                cursor.close()
                conn.close()
            
            # ✅ PENTING: Sleep setelah berhasil memproses semua jadwal
            # Ini mencegah worker berulang terus dengan cepat
            check_interval = AUTO_FILL_CONFIG['CHECK_INTERVAL']
            log_info(f"Pengecekan rutin selesai. Menunggu {check_interval} detik sebelum pengecekan berikutnya...")
            time.sleep(check_interval)
                
        except Exception as e:
            log_error(f"Error kritis di worker auto-fill: {e}")
            
            # ✅ PERBAIKAN: Cek apakah harus berhenti jika ada error
            if AUTO_FILL_CONFIG.get('STOP_WORKER_ON_ERROR', False):
                log_error("Worker auto-fill dihentikan karena error kritis")
                auto_fill_worker_running = False
                break
            
            # ✅ PERBAIKAN: Cek apakah auto-fill dinonaktifkan
            if not AUTO_FILL_CONFIG.get('AUTO_FILL_ENABLED', True):
                log_info("Worker auto-fill dihentikan karena auto-fill dinonaktifkan")
                auto_fill_worker_running = False
                break
            
            # Fallback ke sleep normal
            check_interval = AUTO_FILL_CONFIG['CHECK_INTERVAL']
            log_info(f"Menunggu {check_interval} detik sebelum pengecekan berikutnya...")
            time.sleep(check_interval)
    
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Worker auto-fill dihentikan.")

@app.route('/admin/start_auto_fill_worker', methods=['POST'])
def start_auto_fill_worker_endpoint():
    """Endpoint untuk memulai auto-fill worker tanpa restart aplikasi"""
    global auto_fill_worker_running, auto_fill_worker_thread
    
    if 'user_type' not in session or session['user_type'] != 'admin':
        return jsonify({'success': False, 'message': 'Anda harus login sebagai admin!'})
    
    try:
        if auto_fill_worker_running:
            return jsonify({'success': False, 'message': 'Auto-fill worker sudah berjalan!'})
        
        # Mulai worker
        auto_fill_worker_running = True
        auto_fill_worker_thread = threading.Thread(target=auto_fill_worker, daemon=True)
        auto_fill_worker_thread.start()
        
        print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Auto-fill worker dimulai melalui endpoint!")
        
        return jsonify({'success': True, 'message': 'Auto-fill worker berhasil dimulai!'})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error: {str(e)}'})

@app.route('/admin/auto_fill_worker_status', methods=['GET'])
def get_auto_fill_worker_status():
    """Endpoint untuk mengecek status auto-fill worker"""
    global auto_fill_worker_running, auto_fill_worker_thread
    
    if 'user_type' not in session or session['user_type'] != 'admin':
        return jsonify({'success': False, 'message': 'Anda harus login sebagai admin!'})
    
    try:
        status = {
            'worker_running': auto_fill_worker_running,
            'thread_alive': auto_fill_worker_thread.is_alive() if auto_fill_worker_thread else False,
            'thread_name': auto_fill_worker_thread.name if auto_fill_worker_thread else None,
            'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        }
        
        return jsonify({'success': True, 'status': status})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error: {str(e)}'})

@app.route('/admin/cleanup_incorrect_auto_fill', methods=['POST'])
def cleanup_incorrect_auto_fill():
    """Endpoint untuk membersihkan data auto-fill yang salah (metadata_kolom_id = NULL)"""
    if 'user_type' not in session or session['user_type'] != 'admin':
        return jsonify({'success': False, 'message': 'Anda harus login sebagai admin!'})
    
    try:
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        
        # Hapus data auto-fill yang salah (metadata_kolom_id = NULL)
        cursor.execute("""
            DELETE FROM detail_penilaian_mahasiswa 
            WHERE metadata_kolom_id IS NULL
        """)
        
        deleted_count = cursor.rowcount
        
        # Commit perubahan
        mysql.connection.commit()
        cursor.close()
        
        logger.info(f"Cleanup incorrect auto-fill: {deleted_count} records deleted")
        
        return jsonify({
            'success': True, 
            'message': f'Berhasil membersihkan {deleted_count} data auto-fill yang salah'
        })
        
    except Exception as e:
        logger.error(f"Error cleaning up incorrect auto-fill: {str(e)}")
        return jsonify({'success': False, 'message': f'Error: {str(e)}'})

@app.route('/admin/stop_auto_fill_worker', methods=['POST'])
def stop_auto_fill_worker_endpoint():
    """Endpoint untuk menghentikan auto-fill worker"""
    global auto_fill_worker_running
    
    if 'user_type' not in session or session['user_type'] != 'admin':
        return jsonify({'success': False, 'message': 'Anda harus login sebagai admin!'})
    
    try:
        auto_fill_worker_running = False
        print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Auto-fill worker dihentikan melalui endpoint!")
        
        return jsonify({'success': True, 'message': 'Auto-fill worker berhasil dihentikan!'})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error: {str(e)}'})

@app.route('/admin/auto_fill_status', methods=['GET'])
def get_auto_fill_status():
    """Endpoint untuk mengecek status auto-fill worker"""
    global auto_fill_worker_running
    
    if 'user_type' not in session or session['user_type'] != 'admin':
        return jsonify({'success': False, 'message': 'Anda harus login sebagai admin!'})
    
    return jsonify({
        'success': True, 
        'worker_running': auto_fill_worker_running,
        'message': 'Auto-fill worker sedang berjalan' if auto_fill_worker_running else 'Auto-fill worker tidak berjalan'
    })

@app.route('/admin/run_auto_fill_once', methods=['POST'])
def run_auto_fill_once():
    """Endpoint untuk menjalankan auto-fill satu kali"""
    if 'user_type' not in session or session['user_type'] != 'admin':
        return jsonify({'success': False, 'message': 'Anda harus login sebagai admin!'})
    
    try:
        # Jalankan auto-fill worker satu kali
        auto_fill_worker_single_run()
        return jsonify({'success': True, 'message': 'Auto-fill berhasil dijalankan satu kali!'})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error: {str(e)}'})

@app.route('/admin/run_auto_reject_once', methods=['POST'])
def run_auto_reject_once():
    """Endpoint untuk menjalankan auto-reject satu kali"""
    if 'user_type' not in session or session['user_type'] != 'admin':
        return jsonify({'success': False, 'message': 'Anda harus login sebagai admin!'})
    
    try:
        # Jalankan auto-reject worker satu kali
        auto_reject_worker_single_run()
        return jsonify({'success': True, 'message': 'Auto-reject berhasil dijalankan satu kali!'})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error: {str(e)}'})

def auto_reject_worker_single_run():
    """Jalankan auto-reject worker satu kali"""
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Menjalankan auto-reject manual...")
    
    def log_info(message):
        print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] INFO: {message}")
    
    def log_error(message):
        print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] ERROR: {message}")
    
    try:
        # Koneksi database
        conn = MySQLdb.connect(
            host=AUTO_FILL_CONFIG['DB_HOST'],
            user=AUTO_FILL_CONFIG['DB_USER'],
            passwd=AUTO_FILL_CONFIG['DB_PASSWD'],
            db=AUTO_FILL_CONFIG['DB_NAME']
        )
        cursor = conn.cursor(MySQLdb.cursors.DictCursor)
        
        try:
            # Jalankan auto-reject untuk semua jenis
            auto_reject_anggaran_worker(cursor, log_info, log_error)
            auto_reject_laporan_kemajuan_worker(cursor, log_info, log_error)
            auto_reject_laporan_akhir_worker(cursor, log_info, log_error)
            
            conn.commit()
            log_info("Auto-reject manual berhasil diselesaikan")
            
        except Exception as e:
            log_error(f"Error database: {e}")
            conn.rollback()
        finally:
            cursor.close()
            conn.close()
            
    except Exception as e:
        log_error(f"Error kritis: {e}")

def auto_fill_worker_single_run():
    """Jalankan auto-fill worker satu kali"""
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] Menjalankan auto-fill manual...")
    
    def log_info(message):
        print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] INFO: {message}")
    
    def log_error(message):
        print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] ERROR: {message}")
    
    try:
        # Koneksi database
        conn = MySQLdb.connect(
            host=AUTO_FILL_CONFIG['DB_HOST'],
            user=AUTO_FILL_CONFIG['DB_USER'],
            passwd=AUTO_FILL_CONFIG['DB_PASSWD'],
            db=AUTO_FILL_CONFIG['DB_NAME']
        )
        cursor = conn.cursor()
        
        try:
            # Ambil jadwal penilaian yang aktif
            cursor.execute("""
                SELECT DISTINCT 
                    pj.id,
                    pj.pembimbing_nilai_mulai as waktu_mulai,
                    pj.pembimbing_nilai_selesai as waktu_selesai,
                    pj.pembimbing_interval_tipe as tipe_interval,
                    pj.pembimbing_interval_nilai as nilai_interval,
                    pj.pembimbing_hari_aktif as hari_aktif
                FROM pengaturan_jadwal pj
                WHERE pj.pembimbing_nilai_selesai >= DATE_SUB(NOW(), INTERVAL 24 HOUR)
                AND pj.pembimbing_nilai_mulai IS NOT NULL
                AND pj.pembimbing_nilai_selesai IS NOT NULL
            """)
            
            jadwal_list = cursor.fetchall()
            log_info(f"Ditemukan {len(jadwal_list)} jadwal penilaian yang aktif")
            
            for jadwal in jadwal_list:
                jadwal_id = jadwal[0]
                waktu_mulai = jadwal[1]
                waktu_selesai = jadwal[2]
                tipe_interval = jadwal[3]
                nilai_interval = jadwal[4]
                hari_aktif = jadwal[5]
                
                log_info(f"Memproses jadwal ID {jadwal_id}: Interval {tipe_interval} setiap {nilai_interval}")
                
                # Buat metadata kolom otomatis jika belum ada
                create_metadata_kolom_for_active_schedule(
                    cursor, jadwal_id, tipe_interval, nilai_interval, 
                    waktu_mulai, waktu_selesai, log_info, log_error
                )
                
                # Hitung sesi berdasarkan waktu saat ini
                current_sesi_waktu = calculate_sesi_by_time_worker(
                    waktu_mulai, waktu_selesai, tipe_interval, 
                    nilai_interval, hari_aktif
                )
                
                log_info(f"Sesi saat ini berdasarkan waktu: {current_sesi_waktu}")
                
                if current_sesi_waktu > 1:
                    # Auto-fill sesi yang terlewat
                    auto_fill_all_pembimbing_mahasiswa_worker(
                        cursor, jadwal_id, current_sesi_waktu, waktu_mulai, waktu_selesai, 
                        tipe_interval, nilai_interval, hari_aktif, log_info, log_error
                    )
                
                # NEW: Auto-reject anggaran, laporan kemajuan, dan laporan akhir berdasarkan jadwal
                cursor_dict = conn.cursor(MySQLdb.cursors.DictCursor)
                try:
                    auto_reject_anggaran_worker(cursor_dict, log_info, log_error)
                    auto_reject_laporan_kemajuan_worker(cursor_dict, log_info, log_error)
                    auto_reject_laporan_akhir_worker(cursor_dict, log_info, log_error)
                    
                    # Auto-fill penilaian proposal, laporan kemajuan, dan laporan akhir
                    auto_fill_penilaian_proposal_worker(cursor_dict, log_info, log_error)
                    auto_fill_penilaian_laporan_kemajuan_worker(cursor_dict, log_info, log_error)
                    auto_fill_penilaian_laporan_akhir_worker(cursor_dict, log_info, log_error)
                finally:
                    cursor_dict.close()
            
            conn.commit()
            log_info(f"Berhasil menyelesaikan pengecekan {len(jadwal_list)} jadwal")
            
        except Exception as e:
            log_error(f"Error database: {e}")
            conn.rollback()
        finally:
            cursor.close()
            conn.close()
            
    except Exception as e:
        log_error(f"Error kritis: {e}")

def check_urgent_sessions(log_info, log_error):
    """
    Pengecekan cepat untuk sesi yang sangat urgent (setiap 5 detik)
    """
    try:
        # Koneksi database untuk pengecekan cepat
        conn = MySQLdb.connect(
            host=AUTO_FILL_CONFIG['DB_HOST'],
            user=AUTO_FILL_CONFIG['DB_USER'],
            passwd=AUTO_FILL_CONFIG['DB_PASSWD'],
            db=AUTO_FILL_CONFIG['DB_NAME']
        )
        cursor = conn.cursor()
        
        try:
            # Cek sesi yang sangat urgent (yang baru saja terlewat atau jadwal sudah selesai)
            cursor.execute("""
                SELECT DISTINCT 
                    pj.id,
                    pj.pembimbing_nilai_mulai as waktu_mulai,
                    pj.pembimbing_nilai_selesai as waktu_selesai,
                    pj.pembimbing_interval_tipe as tipe_interval,
                    pj.pembimbing_interval_nilai as nilai_interval
                FROM pengaturan_jadwal pj
                WHERE pj.pembimbing_nilai_selesai >= DATE_SUB(NOW(), INTERVAL 24 HOUR)
                AND pj.pembimbing_nilai_mulai IS NOT NULL
                AND pj.pembimbing_nilai_mulai <= NOW()
            """)
            
            urgent_schedules = cursor.fetchall()
            
            if urgent_schedules:
                log_info(f"Pengecekan prioritas: Ditemukan {len(urgent_schedules)} jadwal urgent")
                
                for schedule in urgent_schedules:
                    jadwal_id = schedule[0]
                    waktu_mulai = schedule[1]
                    waktu_selesai = schedule[2]
                    tipe_interval = schedule[3]
                    nilai_interval = schedule[4]
                    
                    # Hitung sesi urgent
                    current_sesi = calculate_sesi_by_time_worker(
                        waktu_mulai, waktu_selesai, tipe_interval, 
                        nilai_interval, None
                    )
                    
                    if current_sesi > 1:
                        # Auto-fill sesi urgent untuk semua pembimbing-mahasiswa
                        auto_fill_all_pembimbing_mahasiswa_worker(
                            cursor, jadwal_id, current_sesi, waktu_mulai, waktu_selesai, 
                            tipe_interval, nilai_interval, None, log_info, log_error
                        )
                
                conn.commit()
                log_info("Pengecekan prioritas berhasil diselesaikan")
            else:
                log_info("Pengecekan prioritas: Tidak ada jadwal urgent")
                
        except Exception as e:
            log_error(f"Error database pada pengecekan prioritas: {e}")
            conn.rollback()
        finally:
            cursor.close()
            conn.close()
            
    except Exception as e:
        log_error(f"Error pada pengecekan prioritas: {e}")

def check_all_sessions_completed(cursor, jadwal_list, log_info, log_error):
    """
    Cek apakah semua sesi penilaian sudah terisi untuk semua jadwal
    Returns: True jika semua sesi sudah terisi, False jika masih ada yang kosong
    """
    try:
        log_info("Memeriksa apakah semua sesi penilaian sudah selesai...")
        
        # ✅ PERBAIKAN: Cek apakah ada jadwal yang masih aktif
        now = datetime.now()
        active_schedules = []
        
        for jadwal in jadwal_list:
            jadwal_id = jadwal[0]
            waktu_mulai = jadwal[1]
            waktu_selesai = jadwal[2]
            tipe_interval = jadwal[3]
            nilai_interval = jadwal[4]
            hari_aktif = jadwal[5]
            
            # Parse waktu selesai
            waktu_selesai_parsed = parse_date_flexible_worker(waktu_selesai)
            
            if waktu_selesai_parsed and now < waktu_selesai_parsed:
                active_schedules.append(jadwal_id)
                log_info(f"Jadwal ID {jadwal_id} masih aktif (selesai: {waktu_selesai_parsed})")
        
        # ✅ PERBAIKAN: Jika ada jadwal yang masih aktif, anggap belum selesai
        if active_schedules:
            log_info(f"Ada {len(active_schedules)} jadwal yang masih aktif: {active_schedules}")
            return False
        
        # ✅ PERBAIKAN: Jika tidak ada jadwal aktif, cek apakah semua sesi sudah terisi
        for jadwal in jadwal_list:
            jadwal_id = jadwal[0]
            waktu_mulai = jadwal[1]
            waktu_selesai = jadwal[2]
            tipe_interval = jadwal[3]
            nilai_interval = jadwal[4]
            hari_aktif = jadwal[5]
            
            # Hitung total sesi yang seharusnya ada
            total_sesi = calculate_total_sessions_worker(waktu_mulai, waktu_selesai, tipe_interval, nilai_interval, hari_aktif)
            
            # Cek apakah semua sesi sudah terisi untuk semua pembimbing-mahasiswa
            if not check_schedule_sessions_completed(cursor, jadwal_id, total_sesi, log_info, log_error):
                log_info(f"Jadwal ID {jadwal_id} masih memiliki sesi yang belum selesai")
                return False
        
        log_info("Semua sesi penilaian sudah selesai!")
        return True
        
    except Exception as e:
        log_error(f"Error saat memeriksa semua sesi selesai: {e}")
        return False

def calculate_total_sessions_worker(waktu_mulai, waktu_selesai, tipe_interval, nilai_interval, hari_aktif):
    """
    Hitung total sesi berdasarkan jadwal
    """
    try:
        if not waktu_mulai or not waktu_selesai:
            return 1  # Default minimal 1 sesi
            
        # Parse waktu jika dalam bentuk string
        if isinstance(waktu_mulai, str):
            waktu_mulai = parse_date_flexible_worker(waktu_mulai)
        if isinstance(waktu_selesai, str):
            waktu_selesai = parse_date_flexible_worker(waktu_selesai)
        
        if not waktu_mulai or not waktu_selesai:
            return 1  # Default minimal 1 sesi
        
        # Pastikan nilai_interval valid
        if not nilai_interval or nilai_interval <= 0:
            nilai_interval = 1
        
        if tipe_interval == 'setiap_jam':
            # Hitung selisih jam
            diff_hours = (waktu_selesai - waktu_mulai).total_seconds() / 3600
            total_sesi = max(1, int(diff_hours / nilai_interval) + 1)
            return total_sesi
        elif tipe_interval == 'harian':
            # Hitung selisih hari
            diff_days = (waktu_selesai - waktu_mulai).days
            total_sesi = max(1, int(diff_days / nilai_interval) + 1)
            return total_sesi
        elif tipe_interval == 'mingguan':
            # Hitung selisih minggu
            diff_weeks = (waktu_selesai - waktu_mulai).days / 7
            total_sesi = max(1, int(diff_weeks / nilai_interval) + 1)
            return total_sesi
        elif tipe_interval == 'bulanan':
            # Hitung selisih bulan
            diff_months = (waktu_selesai.year - waktu_mulai.year) * 12 + (waktu_selesai.month - waktu_mulai.month)
            total_sesi = max(1, int(diff_months / nilai_interval) + 1)
            return total_sesi
        else:
            # Default untuk interval yang tidak dikenal
            return 1
            
    except Exception as e:
        # Log error untuk debugging
        print(f"Error in calculate_total_sessions_worker: {str(e)}")
        print(f"waktu_mulai: {waktu_mulai}, waktu_selesai: {waktu_selesai}")
        print(f"tipe_interval: {tipe_interval}, nilai_interval: {nilai_interval}")
        return 1  # Default minimal 1 sesi

def check_schedule_sessions_completed(cursor, jadwal_id, total_sesi, log_info, log_error):
    """
    Cek apakah semua sesi untuk jadwal tertentu sudah terisi
    """
    try:
        # ✅ PERBAIKAN: Ambil semua pembimbing-mahasiswa untuk jadwal ini
        cursor.execute("""
            SELECT DISTINCT 
                p.dosen_pembimbing,
                m.nim as nim_mahasiswa,
                p.id as proposal_id
            FROM proposal p
            INNER JOIN mahasiswa m ON p.nim = m.nim
            INNER JOIN pembimbing pb ON pb.nama = p.dosen_pembimbing
            WHERE p.status_admin = 'lolos'
            AND p.status IN ('disetujui', 'revisi')
            AND pb.status = 'aktif'
        """)
        
        pembimbing_mahasiswa_list = cursor.fetchall()
        
        # ✅ PERBAIKAN: Log jumlah pembimbing-mahasiswa yang ditemukan
        log_info(f"Found {len(pembimbing_mahasiswa_list)} pembimbing-mahasiswa pairs for schedule {jadwal_id}")
        
        # ✅ PERBAIKAN: Jika tidak ada pembimbing-mahasiswa, anggap belum selesai
        if not pembimbing_mahasiswa_list:
            log_info(f"No active pembimbing-mahasiswa pairs found for schedule {jadwal_id}, considering as not completed")
            return False
        
        for pembimbing_mahasiswa in pembimbing_mahasiswa_list:
            dosen_pembimbing = pembimbing_mahasiswa[0]
            nim_mahasiswa = pembimbing_mahasiswa[1]
            proposal_id = pembimbing_mahasiswa[2]
            
            # Ambil ID pembimbing
            cursor.execute("SELECT id FROM pembimbing WHERE nama = %s", (dosen_pembimbing,))
            pembimbing_result = cursor.fetchone()
            if not pembimbing_result:
                log_info(f"Pembimbing {dosen_pembimbing} not found in database")
                continue
                
            id_pembimbing = pembimbing_result[0]
            
            # ✅ PERBAIKAN: Cek apakah semua sesi sudah terisi
            for sesi in range(1, total_sesi + 1):
                cursor.execute("""
                    SELECT COUNT(*) FROM detail_penilaian_mahasiswa dpm
                    JOIN penilaian_mahasiswa pm ON dpm.id_penilaian_mahasiswa = pm.id
                    WHERE pm.id_pembimbing = %s AND pm.id_proposal = %s AND dpm.sesi_penilaian = %s
                """, (id_pembimbing, proposal_id, sesi))
                
                if cursor.fetchone()[0] == 0:
                    log_info(f"Sesi {sesi} belum terisi untuk pembimbing {dosen_pembimbing}, mahasiswa {nim_mahasiswa}")
                    return False
        
        log_info(f"All sessions completed for schedule {jadwal_id}")
        return True
        
    except Exception as e:
        log_error(f"Error checking schedule sessions completed: {e}")
        return False

def parse_date_flexible_worker(date_input):
    """
    Parse tanggal dengan format yang fleksibel untuk worker
    """
    if not date_input:
        return None
    if isinstance(date_input, datetime):
        return date_input
    if isinstance(date_input, str):
        try:
            return datetime.strptime(date_input, '%Y-%m-%d %H:%M:%S')
        except:
            try:
                return datetime.strptime(date_input, '%Y-%m-%d')
            except:
                return None
    return None

def calculate_smart_sleep_time(cursor, jadwal_list, log_info, log_error):
    """
    Hitung waktu sleep yang smart berdasarkan kondisi
    """
    try:
        # Jika tidak ada jadwal aktif, sleep lebih lama
        if not jadwal_list:
            return 300  # 5 menit
        
        # Cek apakah ada jadwal yang masih aktif
        now = datetime.now()
        active_schedules = 0
        
        for jadwal in jadwal_list:
            waktu_selesai = parse_date_flexible_worker(jadwal[2])
            if waktu_selesai and now < waktu_selesai:
                active_schedules += 1
        
        # Jika tidak ada jadwal aktif, sleep lebih lama
        if active_schedules == 0:
            return 600  # 10 menit
        else:
            return 60   # 1 menit
            
    except Exception as e:
        log_error(f"Error calculating smart sleep time: {e}")
        return 60  # Fallback ke 1 menit

def calculate_sesi_by_time_worker(waktu_mulai, waktu_selesai, tipe_interval, nilai_interval, hari_aktif):
    """
    Hitung sesi berdasarkan waktu saat ini untuk worker
    FIX: Lebih konservatif - hanya return sesi yang benar-benar sudah waktunya
    """
    try:
        waktu_mulai = waktu_mulai if isinstance(waktu_mulai, datetime) else datetime.strptime(str(waktu_mulai), '%Y-%m-%d %H:%M:%S')
        waktu_selesai = waktu_selesai if isinstance(waktu_selesai, datetime) else datetime.strptime(str(waktu_selesai), '%Y-%m-%d %H:%M:%S')
        
        now = datetime.now()
        
        # Jika waktu sekarang sebelum waktu mulai, return 0
        if now < waktu_mulai:
            return 0
        
        # Jika waktu sekarang setelah waktu selesai, hitung total sesi
        if now >= waktu_selesai:
            return calculate_total_sessions_worker(waktu_mulai, waktu_selesai, tipe_interval, nilai_interval, hari_aktif)
        
        # Hitung sesi berdasarkan waktu yang telah berlalu
        elapsed_time = now - waktu_mulai
        
        if tipe_interval == 'setiap_jam':
            # PERBAIKAN: Logic yang lebih konservatif untuk setiap_jam
            hours_elapsed = elapsed_time.total_seconds() / 3600
            
            # PERBAIKAN: Hanya return sesi baru jika sudah benar-benar lewat interval penuh
            # Contoh: mulai 13:35, interval 1 jam
            # - 14:00: hours_elapsed = 0.42 < 1, return 1 (masih sesi 1)
            # - 14:30: hours_elapsed = 0.92 < 1, return 1 (masih sesi 1)  
            # - 14:35: hours_elapsed = 1.0 >= 1, return 2 (sesi 2 mulai)
            # - 15:35: hours_elapsed = 2.0 >= 2, return 3 (sesi 3 mulai)
            
            if hours_elapsed < nilai_interval:
                return 1  # Masih dalam sesi pertama
            else:
                sesi = int(hours_elapsed / nilai_interval) + 1
                return max(1, sesi)
        elif tipe_interval == 'harian':
            # Interval per hari - LOGIKA YANG BENAR
            days_elapsed = elapsed_time.days
            # FIX: Hanya return sesi 2 jika benar-benar sudah lewat sesi 1
            if days_elapsed < nilai_interval:
                return 1  # Masih dalam sesi pertama
            else:
                sesi = int(days_elapsed / nilai_interval) + 1
                return max(1, sesi)
        elif tipe_interval == 'mingguan':
            # Interval per minggu - LOGIKA YANG BENAR
            weeks_elapsed = elapsed_time.days / 7
            # FIX: Hanya return sesi 2 jika benar-benar sudah lewat sesi 1
            if weeks_elapsed < nilai_interval:
                return 1  # Masih dalam sesi pertama
            else:
                sesi = int(weeks_elapsed / nilai_interval) + 1
                return max(1, sesi)
        elif tipe_interval == 'bulanan':
            # Interval per bulan (perkiraan) - LOGIKA YANG BENAR
            months_elapsed = (now.year - waktu_mulai.year) * 12 + (now.month - waktu_mulai.month)
            # FIX: Hanya return sesi 2 jika benar-benar sudah lewat sesi 1
            if months_elapsed < nilai_interval:
                return 1  # Masih dalam sesi pertama
            else:
                sesi = int(months_elapsed / nilai_interval) + 1
                return max(1, sesi)
        else:
            sesi = 1
        
        return max(1, sesi)
        
    except Exception as e:
        # Log error untuk worker
        return 1

def calculate_total_sessions_worker(waktu_mulai, waktu_selesai, tipe_interval, nilai_interval, hari_aktif):
    """
    Hitung total sesi yang seharusnya ada untuk worker
    """
    try:
        total_duration = waktu_selesai - waktu_mulai
        
        if tipe_interval == 'setiap_jam':
            total_hours = total_duration.total_seconds() / 3600
            return max(1, int(total_hours / nilai_interval) + 1)
        elif tipe_interval == 'harian':
            total_days = total_duration.days
            return max(1, int(total_days / nilai_interval) + 1)
        elif tipe_interval == 'mingguan':
            total_weeks = total_duration.days / 7
            return max(1, int(total_weeks / nilai_interval) + 1)
        elif tipe_interval == 'bulanan':
            total_months = (waktu_selesai.year - waktu_mulai.year) * 12 + (waktu_selesai.month - waktu_mulai.month)
            return max(1, int(total_months / nilai_interval) + 1)
        else:
            return 1
            
    except Exception as e:
        # Log error untuk worker
        return 1

def create_metadata_kolom_for_active_schedule(cursor, jadwal_id, tipe_interval, nilai_interval, waktu_mulai, waktu_selesai, log_info, log_error):
    """
    Buat metadata kolom otomatis untuk jadwal yang aktif jika belum ada
    """
    try:
        # Cek apakah sudah ada metadata kolom untuk jadwal ini
        cursor.execute("""
            SELECT COUNT(*) FROM metadata_kolom_penilaian 
            WHERE jadwal_id = %s
        """, (jadwal_id,))
        
        if cursor.fetchone()[0] > 0:
            log_info(f"Metadata kolom untuk jadwal {jadwal_id} sudah ada, skip pembuatan")
            return
        
        log_info(f"Membuat metadata kolom untuk jadwal {jadwal_id}: {tipe_interval} setiap {nilai_interval}")
        
        # Generate metadata berdasarkan tipe interval
        if tipe_interval == 'setiap_jam':
            # Hitung total jam
            total_hours = int((waktu_selesai - waktu_mulai).total_seconds() / 3600)
            for i in range(total_hours + 1):
                jam = waktu_mulai + timedelta(hours=i)
                nama_kolom = f"Jam {jam.strftime('%H:%M')}"
                
                cursor.execute("""
                    INSERT INTO metadata_kolom_penilaian 
                    (jadwal_id, interval_tipe, interval_nilai, urutan_kolom, nama_kolom, tanggal_mulai, tanggal_selesai)
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                """, (jadwal_id, tipe_interval, nilai_interval, i + 1, nama_kolom, jam.date(), jam.date()))
                
        elif tipe_interval == 'harian':
            # Hitung total hari
            total_days = (waktu_selesai - waktu_mulai).days + 1
            hari_names = ['Senin', 'Selasa', 'Rabu', 'Kamis', 'Jumat', 'Sabtu', 'Minggu']
            
            for i in range(total_days):
                current_date = waktu_mulai + timedelta(days=i)
                hari_index = current_date.weekday()  # 0=Senin, 1=Selasa, dst
                nama_hari = hari_names[hari_index]
                nama_kolom = f"{nama_hari} ({current_date.strftime('%d/%m')})"
                
                cursor.execute("""
                    INSERT INTO metadata_kolom_penilaian 
                    (jadwal_id, interval_tipe, interval_nilai, urutan_kolom, nama_kolom, tanggal_mulai, tanggal_selesai)
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                """, (jadwal_id, tipe_interval, nilai_interval, i + 1, nama_kolom, current_date.date(), current_date.date()))
                
        elif tipe_interval == 'mingguan':
            # Hitung total minggu
            total_weeks = int((waktu_selesai - waktu_mulai).days / 7) + 1
            
            for i in range(total_weeks):
                current_date = waktu_mulai + timedelta(weeks=i)
                nama_kolom = f"Minggu {i + 1} ({current_date.strftime('%d/%m')})"
                
                cursor.execute("""
                    INSERT INTO metadata_kolom_penilaian 
                    (jadwal_id, interval_tipe, interval_nilai, urutan_kolom, nama_kolom, tanggal_mulai, tanggal_selesai)
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                """, (jadwal_id, tipe_interval, nilai_interval, i + 1, nama_kolom, current_date.date(), current_date.date()))
                
        elif tipe_interval == 'bulanan':
            # Hitung total bulan
            total_months = (waktu_selesai.year - waktu_mulai.year) * 12 + (waktu_selesai.month - waktu_mulai.month) + 1
            bulan_names = ['Januari', 'Februari', 'Maret', 'April', 'Mei', 'Juni', 
                          'Juli', 'Agustus', 'September', 'Oktober', 'November', 'Desember']
            
            for i in range(total_months):
                current_date = waktu_mulai + timedelta(days=30*i)  # Perkiraan
                bulan_index = current_date.month - 1
                nama_bulan = bulan_names[bulan_index]
                nama_kolom = f"{nama_bulan} {current_date.year}"
                
                cursor.execute("""
                    INSERT INTO metadata_kolom_penilaian 
                    (jadwal_id, interval_tipe, interval_nilai, urutan_kolom, nama_kolom, tanggal_mulai, tanggal_selesai)
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                """, (jadwal_id, tipe_interval, nilai_interval, i + 1, nama_kolom, current_date.date(), current_date.date()))
        
        log_info(f"Berhasil membuat metadata kolom untuk jadwal {jadwal_id}")
        
    except Exception as e:
        log_error(f"Error membuat metadata kolom untuk jadwal {jadwal_id}: {str(e)}")

def auto_fill_all_pembimbing_mahasiswa_worker(cursor, jadwal_id, current_sesi_waktu, 
                                             waktu_mulai, waktu_selesai, tipe_interval, nilai_interval, hari_aktif, log_info, log_error):
    """
    Auto-fill sesi terlewat untuk semua pembimbing-mahasiswa berdasarkan jadwal global
    """
    try:
        # PERBAIKAN: Cek apakah jadwal sudah selesai
        now = datetime.now()
        jadwal_sudah_selesai = now >= waktu_selesai
        
        if jadwal_sudah_selesai:
            # Jika jadwal sudah selesai, auto-fill SEMUA sesi dari 1 sampai total_sesi
            total_sesi = calculate_total_sessions_worker(waktu_mulai, waktu_selesai, tipe_interval, nilai_interval, hari_aktif)
            sessions_to_check = total_sesi
            log_info(f"Jadwal sudah selesai. Auto-fill semua sesi dari 1 sampai {total_sesi}")
            log_info(f"Logic: current_sesi_waktu = {current_sesi_waktu}, jadwal_sudah_selesai = True")
            log_info(f"Auto-fill sessions 1 to {total_sesi}")
        else:
            # Jika jadwal masih berlangsung, auto-fill sesi yang sudah lewat saja
            sessions_to_check = current_sesi_waktu - 1 if current_sesi_waktu > 1 else 1
            log_info(f"Checking for missed sessions from 1 to {sessions_to_check} for all pembimbing-mahasiswa (current_sesi_waktu: {current_sesi_waktu})")
            log_info(f"Logic: current_sesi_waktu = {current_sesi_waktu} means we're in session {current_sesi_waktu}")
            log_info(f"Auto-fill sessions 1 to {sessions_to_check} because they have already passed")
            log_info(f"Session {current_sesi_waktu} is currently active and should not be auto-filled")
        
        # ✅ PERBAIKAN BARU: Filter proposal berdasarkan tanggal untuk menghindari data lama
        # Hanya ambil proposal yang dibuat DURING atau AFTER jadwal aktif
        cursor.execute("""
            SELECT DISTINCT 
                p.dosen_pembimbing,
                m.nim as nim_mahasiswa,
                p.id as proposal_id,
                p.tanggal_buat,
                p.tanggal_kirim
            FROM proposal p
            INNER JOIN mahasiswa m ON p.nim = m.nim
            INNER JOIN pembimbing pb ON pb.nama = p.dosen_pembimbing
            WHERE p.status_admin = 'lolos'
            AND p.status IN ('disetujui', 'revisi', 'selesai')
            AND pb.status = 'aktif'
            AND p.tanggal_buat >= %s  -- ✅ PERBAIKAN: Hanya proposal yang dibuat setelah jadwal aktif
            AND p.tanggal_buat <= %s  -- ✅ PERBAIKAN: Hanya proposal yang dibuat sebelum jadwal selesai
            AND NOT EXISTS (  -- ✅ PERBAIKAN BARU: Exclude data lama yang sudah memiliki penilaian
                SELECT 1 FROM penilaian_mahasiswa pm 
                WHERE pm.id_proposal = p.id 
                AND pm.jadwal_id IS NULL  -- Data lama
            )
        """, (waktu_mulai, waktu_selesai))
        
        pembimbing_mahasiswa_list = cursor.fetchall()
        log_info(f"Found {len(pembimbing_mahasiswa_list)} active pembimbing-mahasiswa pairs")
        
        # Cek apakah ada sesi yang terlewat untuk setiap pembimbing-mahasiswa
        for pembimbing_mahasiswa in pembimbing_mahasiswa_list:
            # Handle both tuple and dict formats
            if isinstance(pembimbing_mahasiswa, dict):
                dosen_pembimbing = pembimbing_mahasiswa['dosen_pembimbing']
                nim_mahasiswa = pembimbing_mahasiswa['nim_mahasiswa']
                proposal_id = pembimbing_mahasiswa['proposal_id']
            else:
                dosen_pembimbing = pembimbing_mahasiswa[0]
                nim_mahasiswa = pembimbing_mahasiswa[1]
                proposal_id = pembimbing_mahasiswa[2]
            
            # Ambil ID pembimbing dari nama
            cursor.execute("SELECT id FROM pembimbing WHERE nama = %s", (dosen_pembimbing,))
            pembimbing_result = cursor.fetchone()
            if not pembimbing_result:
                log_info(f"Pembimbing {dosen_pembimbing} not found, skipping...")
                continue
                
            # ✅ PERBAIKAN BARU: Cek apakah ini data lama berdasarkan jadwal_id
            cursor.execute("""
                SELECT jadwal_id FROM penilaian_mahasiswa 
                WHERE id_pembimbing = %s AND id_proposal = %s
            """, (id_pembimbing, proposal_id))
            
            penilaian_existing = cursor.fetchone()
            if penilaian_existing:
                jadwal_id = penilaian_existing['jadwal_id'] if isinstance(penilaian_existing, dict) else penilaian_existing[0]
                
                # ✅ PERBAIKAN BARU: Jika jadwal_id NULL, ini data lama - SKIP auto-fill
                if jadwal_id is None:
                    log_info(f"Data lama terdeteksi - jadwal_id NULL untuk pembimbing {dosen_pembimbing}, mahasiswa {nim_mahasiswa}, SKIP auto-fill")
                    continue
                
            # Handle both tuple and dict formats
            if isinstance(pembimbing_result, dict):
                id_pembimbing = pembimbing_result['id']
            else:
                id_pembimbing = pembimbing_result[0]
            
            # PERBAIKAN: Logika yang benar untuk jadwal yang sudah selesai
            if jadwal_sudah_selesai:
                # Jika jadwal sudah selesai, cek SEMUA sesi dari 1 sampai total_sesi
                log_info(f"Processing pembimbing {dosen_pembimbing}, mahasiswa {nim_mahasiswa} - checking ALL sessions 1 to {sessions_to_check} (jadwal sudah selesai)")
                log_info(f"Logic: jadwal_sudah_selesai = True, so check ALL sessions 1 to {sessions_to_check} for auto-fill")
            else:
                # Jika jadwal masih berlangsung, cek sesi yang sudah lewat saja
                log_info(f"Processing pembimbing {dosen_pembimbing}, mahasiswa {nim_mahasiswa} - checking sessions 1 to {sessions_to_check} (current_sesi_waktu: {current_sesi_waktu})")
                log_info(f"Logic: current_sesi_waktu = {current_sesi_waktu} means we're in session {current_sesi_waktu}, so check sessions 1 to {sessions_to_check} for auto-fill")
            
            for sesi in range(1, sessions_to_check + 1):
                log_info(f"Checking sesi {sesi} for pembimbing {dosen_pembimbing}, mahasiswa {nim_mahasiswa}")
                
                # ✅ PERBAIKAN: Cek apakah sesi ini sudah ada data untuk pembimbing-mahasiswa ini
                # Dan pastikan data tersebut memiliki metadata_kolom_id yang valid (bukan NULL)
                cursor.execute("""
                    SELECT COUNT(*) as count FROM detail_penilaian_mahasiswa dpm
                    JOIN penilaian_mahasiswa pm ON dpm.id_penilaian_mahasiswa = pm.id
                    WHERE pm.id_pembimbing = %s AND pm.id_proposal = %s AND dpm.sesi_penilaian = %s
                    AND dpm.metadata_kolom_id IS NOT NULL
                """, (id_pembimbing, proposal_id, sesi))
                
                result = cursor.fetchone()
                count = result['count'] if isinstance(result, dict) else result[0]
                
                # ✅ PERBAIKAN BARU: Cek apakah ini data lama yang sudah memiliki penilaian
                # Jika sudah ada data dengan metadata_kolom_id valid, skip auto-fill
                if count > 0:
                    log_info(f"Data lama terdeteksi untuk sesi {sesi} - sudah ada {count} data dengan metadata_kolom_id valid, SKIP auto-fill")
                    continue
                
                # ✅ PERBAIKAN BARU: Cek apakah proposal ini adalah data lama berdasarkan tanggal
                cursor.execute("""
                    SELECT p.tanggal_buat, p.tanggal_kirim 
                    FROM proposal p 
                    WHERE p.id = %s
                """, (proposal_id,))
                
                proposal_info = cursor.fetchone()
                if proposal_info:
                    tanggal_buat = proposal_info['tanggal_buat'] if isinstance(proposal_info, dict) else proposal_info[0]
                    tanggal_kirim = proposal_info['tanggal_kirim'] if isinstance(proposal_info, dict) else proposal_info[1]
                    
                    # ✅ PERBAIKAN BARU: Validasi ketat untuk data lama
                    # 1. Jika proposal dibuat sebelum jadwal aktif saat ini, ini adalah data lama
                    # TAPI: Jika proposal dibuat dalam 24 jam sebelum jadwal aktif, masih dianggap valid
                    if tanggal_buat and tanggal_buat < waktu_mulai:
                        # Cek apakah proposal dibuat dalam 24 jam sebelum jadwal aktif
                        from datetime import timedelta
                        threshold_time = waktu_mulai - timedelta(hours=24)
                        if tanggal_buat < threshold_time:
                            log_info(f"Data lama terdeteksi - proposal dibuat {tanggal_buat} terlalu lama sebelum jadwal aktif {waktu_mulai}, SKIP auto-fill")
                            continue
                        else:
                            log_info(f"Data valid - proposal dibuat {tanggal_buat} dalam 24 jam sebelum jadwal aktif {waktu_mulai}, lanjutkan auto-fill")
                    
                    # 2. Jika proposal dibuat setelah jadwal selesai, ini adalah data baru yang tidak relevan
                    if tanggal_buat and tanggal_buat > waktu_selesai:
                        log_info(f"Data baru terdeteksi - proposal dibuat {tanggal_buat} setelah jadwal selesai {waktu_selesai}, SKIP auto-fill")
                        continue
                    
                    # 3. ✅ PERBAIKAN BARU: Cek apakah ini data lama berdasarkan jadwal_id
                    cursor.execute("""
                        SELECT jadwal_id FROM penilaian_mahasiswa 
                        WHERE id_pembimbing = %s AND id_proposal = %s
                    """, (id_pembimbing, proposal_id))
                    
                    jadwal_result = cursor.fetchone()
                    if jadwal_result:
                        jadwal_id = jadwal_result[0] if isinstance(jadwal_result, dict) else jadwal_result[0]
                        if jadwal_id is None:
                            log_info(f"Data lama terdeteksi - jadwal_id IS NULL untuk proposal {proposal_id}, SKIP auto-fill")
                            continue
                    
                    # 3. ✅ PERBAIKAN BARU: Cek apakah proposal ini sudah memiliki penilaian yang valid
                    cursor.execute("""
                        SELECT COUNT(*) as count_valid FROM detail_penilaian_mahasiswa dpm
                        JOIN penilaian_mahasiswa pm ON dpm.id_penilaian_mahasiswa = pm.id
                        WHERE pm.id_pembimbing = %s AND pm.id_proposal = %s
                        AND dpm.metadata_kolom_id IS NOT NULL
                    """, (id_pembimbing, proposal_id))
                    
                    result_valid = cursor.fetchone()
                    count_valid = result_valid['count_valid'] if isinstance(result_valid, dict) else result_valid[0]
                    
                    # Jika sudah ada data valid, skip auto-fill untuk proposal ini
                    if count_valid > 0:
                        log_info(f"Proposal {proposal_id} sudah memiliki {count_valid} data penilaian valid, SKIP auto-fill")
                        continue
                    
                    # 4. ✅ PERBAIKAN BARU: Cek berdasarkan konfigurasi SKIP_OLD_DATA
                    if AUTO_FILL_CONFIG.get('SKIP_OLD_DATA', False):
                        # Cek apakah proposal dibuat lebih dari threshold hari yang ditentukan
                        threshold_days = AUTO_FILL_CONFIG.get('OLD_DATA_THRESHOLD_DAYS', 7)
                        threshold_date = datetime.now() - timedelta(days=threshold_days)
                        
                        if tanggal_buat and tanggal_buat < threshold_date:
                            log_info(f"Data lama terdeteksi berdasarkan threshold - proposal dibuat {tanggal_buat} (lebih dari {threshold_days} hari), SKIP auto-fill")
                            continue
                
                # ✅ PERBAIKAN: Cek juga apakah ada data dengan metadata_kolom_id = NULL (data auto-fill yang salah)
                cursor.execute("""
                    SELECT COUNT(*) as count_null FROM detail_penilaian_mahasiswa dpm
                    JOIN penilaian_mahasiswa pm ON dpm.id_penilaian_mahasiswa = pm.id
                    WHERE pm.id_pembimbing = %s AND pm.id_proposal = %s AND dpm.sesi_penilaian = %s
                    AND dpm.metadata_kolom_id IS NULL
                """, (id_pembimbing, proposal_id, sesi))
                
                result_null = cursor.fetchone()
                count_null = result_null['count_null'] if isinstance(result_null, dict) else result_null[0]
                
                # ✅ PERBAIKAN: Jika ada data dengan metadata_kolom_id = NULL, hapus dulu
                if count_null > 0:
                    log_info(f"Found {count_null} incorrect auto-fill data for sesi {sesi}, cleaning up...")
                    cursor.execute("""
                        DELETE FROM detail_penilaian_mahasiswa dpm
                        JOIN penilaian_mahasiswa pm ON dpm.id_penilaian_mahasiswa = pm.id
                        WHERE pm.id_pembimbing = %s AND pm.id_proposal = %s AND dpm.sesi_penilaian = %s
                        AND dpm.metadata_kolom_id IS NULL
                    """, (id_pembimbing, proposal_id, sesi))
                    log_info(f"Cleaned up {count_null} incorrect auto-fill data for sesi {sesi}")
                
                # ✅ PERBAIKAN BARU: Cek apakah proposal ini sudah memiliki penilaian yang valid untuk sesi ini
                cursor.execute("""
                    SELECT COUNT(*) as count_valid FROM detail_penilaian_mahasiswa dpm
                    JOIN penilaian_mahasiswa pm ON dpm.id_penilaian_mahasiswa = pm.id
                    WHERE pm.id_pembimbing = %s AND pm.id_proposal = %s AND dpm.sesi_penilaian = %s
                    AND dpm.metadata_kolom_id IS NOT NULL
                    AND dpm.skor > 0  -- Sudah ada penilaian manual
                """, (id_pembimbing, proposal_id, sesi))
                
                result_valid = cursor.fetchone()
                count_valid = result_valid['count_valid'] if isinstance(result_valid, dict) else result_valid[0]
                
                # Jika sudah ada penilaian manual untuk sesi ini, skip auto-fill
                if count_valid > 0:
                    log_info(f"Proposal {proposal_id} sudah memiliki {count_valid} penilaian manual untuk sesi {sesi}, SKIP auto-fill")
                    continue
                
                if count == 0:
                    # Sesi belum ada data yang valid, auto-fill dengan skor 0
                    log_info(f"Auto-filling sesi {sesi} for pembimbing {dosen_pembimbing}, mahasiswa {nim_mahasiswa}")
                    
                    # PERBAIKAN: Ambil pertanyaan berdasarkan snapshot jadwal mulai
                    # Gunakan pertanyaan yang aktif pada saat jadwal mulai
                    cursor.execute("""
                        SELECT p.id FROM pertanyaan_penilaian_mahasiswa p
                        WHERE p.status = 'Aktif' 
                        AND p.created_at <= %s
                        ORDER BY p.id
                    """, (waktu_mulai,))
                    pertanyaan_list = cursor.fetchall()
                    
                    # ✅ PERBAIKAN BARU: Ambil metadata yang sesuai dengan jadwal aktif
                    # Hanya gunakan metadata yang dibuat dalam rentang jadwal aktif
                    cursor.execute("""
                        SELECT id FROM metadata_kolom_penilaian 
                        WHERE jadwal_id = %s 
                        AND urutan_kolom = %s
                        AND created_at >= %s  -- Metadata dibuat setelah jadwal mulai
                        AND created_at <= %s  -- Metadata dibuat sebelum jadwal selesai
                        ORDER BY created_at DESC 
                        LIMIT 1
                    """, (jadwal_id, sesi, waktu_mulai, waktu_selesai))
                    
                    metadata_result = cursor.fetchone()
                    if not metadata_result:
                        log_info(f"No valid metadata found for sesi {sesi} in active schedule, jadwal_id {jadwal_id}, skipping...")
                        continue
                    
                    metadata_kolom_id = metadata_result['id'] if isinstance(metadata_result, dict) else metadata_result[0]
                    log_info(f"Using metadata_kolom_id {metadata_kolom_id} for sesi {sesi}")
                    
                    log_info(f"Found {len(pertanyaan_list)} active assessment questions")
                    
                    # Cek apakah sudah ada penilaian_mahasiswa untuk pembimbing-mahasiswa ini
                    cursor.execute("""
                        SELECT id FROM penilaian_mahasiswa 
                        WHERE id_pembimbing = %s AND id_proposal = %s
                    """, (id_pembimbing, proposal_id))
                    
                    penilaian_result = cursor.fetchone()
                    if penilaian_result:
                        # Handle both tuple and dict formats
                        if isinstance(penilaian_result, dict):
                            id_penilaian = penilaian_result['id']
                        else:
                            id_penilaian = penilaian_result[0]
                    else:
                        # Buat penilaian baru jika belum ada
                        cursor.execute("""
                            INSERT INTO penilaian_mahasiswa 
                            (id_proposal, id_pembimbing, nilai_akhir, komentar_pembimbing, status)
                            VALUES (%s, %s, %s, %s, %s)
                        """, (proposal_id, id_pembimbing, 
                              0,  # Default score 0
                              'Auto-filled: Jadwal sudah selesai',
                              'Draft'))
                        
                        id_penilaian = cursor.lastrowid
                    
                    # Buat detail penilaian dengan skor 0 untuk sesi yang terlewat
                    for pertanyaan in pertanyaan_list:
                        # Handle both tuple and dict formats
                        if isinstance(pertanyaan, dict):
                            id_pertanyaan = pertanyaan['id']
                        else:
                            id_pertanyaan = pertanyaan[0]
                        
                        # Cek apakah sudah ada detail penilaian untuk sesi ini
                        cursor.execute("""
                            SELECT COUNT(*) as count FROM detail_penilaian_mahasiswa 
                            WHERE id_penilaian_mahasiswa = %s AND id_pertanyaan = %s AND sesi_penilaian = %s
                        """, (id_penilaian, id_pertanyaan, sesi))
                        
                        result = cursor.fetchone()
                        count = result['count'] if isinstance(result, dict) else result[0]
                        
                        if count == 0:
                            # ✅ PERBAIKAN: Gunakan metadata_kolom_id yang sudah diambil sebelumnya
                            # metadata_kolom_id sudah diambil di atas dan divalidasi
                            if metadata_kolom_id is None:
                                log_info(f"Metadata kolom tidak valid untuk sesi {sesi}, skipping...")
                                continue
                            
                            # Buat detail penilaian dengan skor 0
                            cursor.execute("""
                                INSERT INTO detail_penilaian_mahasiswa 
                                (id_penilaian_mahasiswa, id_pertanyaan, skor, nilai, sesi_penilaian, is_locked, tanggal_input, metadata_kolom_id)
                                VALUES (%s, %s, %s, %s, %s, %s, NOW(), %s)
                            """, (id_penilaian, id_pertanyaan, 
                                  0, 0, sesi, 1, metadata_kolom_id))
                            
                            log_info(f"Auto-filled sesi {sesi}, pertanyaan {id_pertanyaan} with score 0, metadata_kolom_id: {metadata_kolom_id}")
                    
                    log_info(f"Completed auto-filling sesi {sesi} for pembimbing {dosen_pembimbing}, mahasiswa {nim_mahasiswa}")
                else:
                    log_info(f"Sesi {sesi} already exists for pembimbing {dosen_pembimbing}, mahasiswa {nim_mahasiswa}, skipping...")
                
    except Exception as e:
        log_error(f"Error in auto_fill_all_pembimbing_mahasiswa_worker: {str(e)}")
        # Log detail error untuk debugging
        import traceback
        log_error(f"Traceback: {traceback.format_exc()}")
        # Tidak raise exception, biarkan worker tetap berjalan
        return False

def auto_fill_penilaian_proposal_worker(cursor, log_info, log_error):
    """
    Auto-fill penilaian proposal dengan skor 0 jika jadwal sudah selesai dan belum ada penilaian
    """
    try:
        # Cek jadwal penilaian proposal
        cursor.execute('''
            SELECT reviewer_proposal_mulai, reviewer_proposal_selesai
            FROM pengaturan_jadwal 
            ORDER BY id DESC 
            LIMIT 1
        ''')
        
        jadwal = cursor.fetchone()
        if not jadwal or not jadwal.get('reviewer_proposal_selesai'):
            log_info("Tidak ada jadwal penilaian proposal, skip auto-fill")
            return
        
        jadwal_selesai = jadwal['reviewer_proposal_selesai']
        now = datetime.now()
        
        if now <= jadwal_selesai:
            log_info("Jadwal penilaian proposal belum selesai, skip auto-fill")
            return
        
        log_info(f"Jadwal penilaian proposal sudah selesai ({jadwal_selesai}), mulai auto-fill")
        
        # Ambil semua proposal yang sudah selesai direview tapi belum dinilai
        cursor.execute('''
            SELECT pr.id_proposal, pr.id_reviewer
            FROM proposal_reviewer pr
            WHERE pr.status_review = 'selesai_review'
            AND NOT EXISTS (
                SELECT 1 FROM penilaian_proposal pp 
                WHERE pp.id_proposal = pr.id_proposal 
                AND pp.id_reviewer = pr.id_reviewer
            )
        ''')
        
        proposal_reviewer_list = cursor.fetchall()
        log_info(f"Found {len(proposal_reviewer_list)} proposal-reviewer pairs without assessment")
        
        # Ambil pertanyaan berdasarkan snapshot
        from reviewer import get_pertanyaan_proposal_snapshot
        pertanyaan_list = get_pertanyaan_proposal_snapshot(cursor)
        
        if not pertanyaan_list:
            log_error("Tidak ada pertanyaan penilaian proposal aktif")
            return
        
        filled_count = 0
        
        for proposal_reviewer in proposal_reviewer_list:
            proposal_id = proposal_reviewer['id_proposal']
            reviewer_id = proposal_reviewer['id_reviewer']
            
            try:
                # Buat header penilaian dengan nilai 0
                cursor.execute('''
                    INSERT INTO penilaian_proposal 
                    (id_proposal, id_reviewer, total_skor, total_nilai, nilai_akhir, catatan, nilai_bantuan, persentase_bantuan, tanggal_penilaian)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, NOW())
                ''', (proposal_id, reviewer_id, 0, 0, 0, 'Auto-filled: Jadwal penilaian sudah selesai', 0, 0))
                
                # Buat detail penilaian dengan skor 0
                for pertanyaan in pertanyaan_list:
                    nilai = pertanyaan['bobot'] * 0  # bobot × 0 = 0
                    cursor.execute('''
                        INSERT INTO detail_penilaian_proposal 
                        (id_proposal, id_reviewer, id_pertanyaan, skor, bobot, nilai)
                        VALUES (%s, %s, %s, %s, %s, %s)
                    ''', (proposal_id, reviewer_id, pertanyaan['id'], 0, pertanyaan['bobot'], nilai))
                
                filled_count += 1
                log_info(f"Auto-filled proposal {proposal_id} for reviewer {reviewer_id}")
                
            except Exception as e:
                log_error(f"Error auto-filling proposal {proposal_id} for reviewer {reviewer_id}: {str(e)}")
        
        log_info(f"Auto-fill penilaian proposal completed: {filled_count} assessments filled")
        
    except Exception as e:
        log_error(f"Error in auto_fill_penilaian_proposal_worker: {str(e)}")

def auto_fill_penilaian_laporan_kemajuan_worker(cursor, log_info, log_error):
    """
    Auto-fill penilaian laporan kemajuan dengan skor 0 jika jadwal sudah selesai dan belum ada penilaian
    """
    try:
        # Cek jadwal penilaian laporan kemajuan
        cursor.execute('''
            SELECT reviewer_kemajuan_mulai, reviewer_kemajuan_selesai
            FROM pengaturan_jadwal 
            ORDER BY id DESC 
            LIMIT 1
        ''')
        
        jadwal = cursor.fetchone()
        if not jadwal or not jadwal.get('reviewer_kemajuan_selesai'):
            log_info("Tidak ada jadwal penilaian laporan kemajuan, skip auto-fill")
            return
        
        jadwal_selesai = jadwal['reviewer_kemajuan_selesai']
        now = datetime.now()
        
        if now <= jadwal_selesai:
            log_info("Jadwal penilaian laporan kemajuan belum selesai, skip auto-fill")
            return
        
        log_info(f"Jadwal penilaian laporan kemajuan sudah selesai ({jadwal_selesai}), mulai auto-fill")
        
        # Ambil semua proposal yang sudah selesai direview tapi belum dinilai laporan kemajuan
        cursor.execute('''
            SELECT pr.id_proposal, pr.id_reviewer
            FROM proposal_reviewer pr
            WHERE pr.status_review = 'selesai_review'
            AND NOT EXISTS (
                SELECT 1 FROM penilaian_laporan_kemajuan plk 
                WHERE plk.id_proposal = pr.id_proposal 
                AND plk.id_reviewer = pr.id_reviewer
            )
        ''')
        
        proposal_reviewer_list = cursor.fetchall()
        log_info(f"Found {len(proposal_reviewer_list)} proposal-reviewer pairs without laporan kemajuan assessment")
        
        # Ambil pertanyaan berdasarkan snapshot
        from reviewer import get_pertanyaan_laporan_kemajuan_snapshot
        pertanyaan_list = get_pertanyaan_laporan_kemajuan_snapshot(cursor)
        
        if not pertanyaan_list:
            log_error("Tidak ada pertanyaan penilaian laporan kemajuan aktif")
            return
        
        filled_count = 0
        
        for proposal_reviewer in proposal_reviewer_list:
            proposal_id = proposal_reviewer['id_proposal']
            reviewer_id = proposal_reviewer['id_reviewer']
            
            try:
                # Hitung nilai akhir dengan rumus baru
                total_bobot = sum(p['bobot'] for p in pertanyaan_list)
                total_nilai = 0  # Semua skor 0
                skor_maksimal_tertinggi = max([p['skor_maksimal'] for p in pertanyaan_list]) if pertanyaan_list else 100
                nilai_akhir = (total_nilai / (total_bobot * skor_maksimal_tertinggi)) * 100 if (total_bobot > 0 and skor_maksimal_tertinggi > 0) else 0
                
                # Buat header penilaian dengan nilai 0
                cursor.execute('''
                    INSERT INTO penilaian_laporan_kemajuan 
                    (id_proposal, id_reviewer, nilai_akhir, komentar_reviewer, rekomendasi_pendanaan, alasan_rekomendasi, tanggal_penilaian)
                    VALUES (%s, %s, %s, %s, %s, %s, NOW())
                ''', (proposal_id, reviewer_id, nilai_akhir, 'Auto-filled: Jadwal penilaian sudah selesai', 'berhenti pendanaan', 'Auto-filled: Skor 0'))
                
                penilaian_id = cursor.lastrowid
                
                # Buat detail penilaian dengan skor 0
                for pertanyaan in pertanyaan_list:
                    nilai_terbobot = pertanyaan['bobot'] * 0  # bobot × 0 = 0
                    cursor.execute('''
                        INSERT INTO detail_penilaian_laporan_kemajuan 
                        (id_penilaian_laporan_kemajuan, id_pertanyaan, skor_diberikan, bobot_pertanyaan, skor_maksimal, nilai_terbobot)
                        VALUES (%s, %s, %s, %s, %s, %s)
                    ''', (penilaian_id, pertanyaan['id'], 0, pertanyaan['bobot'], pertanyaan['skor_maksimal'], nilai_terbobot))
                
                filled_count += 1
                log_info(f"Auto-filled laporan kemajuan proposal {proposal_id} for reviewer {reviewer_id}")
                
            except Exception as e:
                log_error(f"Error auto-filling laporan kemajuan proposal {proposal_id} for reviewer {reviewer_id}: {str(e)}")
        
        log_info(f"Auto-fill penilaian laporan kemajuan completed: {filled_count} assessments filled")
        
    except Exception as e:
        log_error(f"Error in auto_fill_penilaian_laporan_kemajuan_worker: {str(e)}")

def auto_fill_penilaian_laporan_akhir_worker(cursor, log_info, log_error):
    """
    Auto-fill penilaian laporan akhir dengan skor 0 jika jadwal sudah selesai dan belum ada penilaian
    """
    try:
        # Cek jadwal penilaian laporan akhir
        cursor.execute('''
            SELECT reviewer_akhir_mulai, reviewer_akhir_selesai
            FROM pengaturan_jadwal 
            ORDER BY id DESC 
            LIMIT 1
        ''')
        
        jadwal = cursor.fetchone()
        if not jadwal or not jadwal.get('reviewer_akhir_selesai'):
            log_info("Tidak ada jadwal penilaian laporan akhir, skip auto-fill")
            return
        
        jadwal_selesai = jadwal['reviewer_akhir_selesai']
        now = datetime.now()
        
        if now <= jadwal_selesai:
            log_info("Jadwal penilaian laporan akhir belum selesai, skip auto-fill")
            return
        
        log_info(f"Jadwal penilaian laporan akhir sudah selesai ({jadwal_selesai}), mulai auto-fill")
        
        # Ambil semua proposal yang sudah selesai direview tapi belum dinilai laporan akhir
        cursor.execute('''
            SELECT pr.id_proposal, pr.id_reviewer
            FROM proposal_reviewer pr
            WHERE pr.status_review = 'selesai_review'
            AND NOT EXISTS (
                SELECT 1 FROM penilaian_laporan_akhir pla 
                WHERE pla.id_proposal = pr.id_proposal 
                AND pla.id_reviewer = pr.id_reviewer
            )
        ''')
        
        proposal_reviewer_list = cursor.fetchall()
        log_info(f"Found {len(proposal_reviewer_list)} proposal-reviewer pairs without laporan akhir assessment")
        
        # Ambil pertanyaan berdasarkan snapshot
        from reviewer import get_pertanyaan_laporan_akhir_snapshot
        pertanyaan_list = get_pertanyaan_laporan_akhir_snapshot(cursor)
        
        if not pertanyaan_list:
            log_error("Tidak ada pertanyaan penilaian laporan akhir aktif")
            return
        
        filled_count = 0
        
        for proposal_reviewer in proposal_reviewer_list:
            proposal_id = proposal_reviewer['id_proposal']
            reviewer_id = proposal_reviewer['id_reviewer']
            
            try:
                # Hitung nilai akhir dengan rumus baru
                total_bobot = sum(p['bobot'] for p in pertanyaan_list)
                total_nilai = 0  # Semua skor 0
                skor_maksimal_tertinggi = max([p['skor_maksimal'] for p in pertanyaan_list]) if pertanyaan_list else 100
                nilai_akhir = (total_nilai / (total_bobot * skor_maksimal_tertinggi)) * 100 if (total_bobot > 0 and skor_maksimal_tertinggi > 0) else 0
                
                # Buat header penilaian dengan nilai 0
                cursor.execute('''
                    INSERT INTO penilaian_laporan_akhir 
                    (id_proposal, id_reviewer, nilai_akhir, komentar_reviewer, rekomendasi_kelulusan, alasan_rekomendasi, tanggal_penilaian)
                    VALUES (%s, %s, %s, %s, %s, %s, NOW())
                ''', (proposal_id, reviewer_id, nilai_akhir, 'Auto-filled: Jadwal penilaian sudah selesai', 'tidak lolos', 'Auto-filled: Skor 0'))
                
                penilaian_id = cursor.lastrowid
                
                # Buat detail penilaian dengan skor 0
                for pertanyaan in pertanyaan_list:
                    nilai_terbobot = pertanyaan['bobot'] * 0  # bobot × 0 = 0
                    cursor.execute('''
                        INSERT INTO detail_penilaian_laporan_akhir 
                        (id_penilaian_laporan_akhir, id_pertanyaan, skor_diberikan, bobot_pertanyaan, skor_maksimal, nilai_terbobot)
                        VALUES (%s, %s, %s, %s, %s, %s)
                    ''', (penilaian_id, pertanyaan['id'], 0, pertanyaan['bobot'], pertanyaan['skor_maksimal'], nilai_terbobot))
                
                filled_count += 1
                log_info(f"Auto-filled laporan akhir proposal {proposal_id} for reviewer {reviewer_id}")
                
            except Exception as e:
                log_error(f"Error auto-filling laporan akhir proposal {proposal_id} for reviewer {reviewer_id}: {str(e)}")
        
        log_info(f"Auto-fill penilaian laporan akhir completed: {filled_count} assessments filled")
        
    except Exception as e:
        log_error(f"Error in auto_fill_penilaian_laporan_akhir_worker: {str(e)}")

def auto_reject_anggaran_worker(cursor, log_info, log_error):
    """
    Otomatis mengubah status anggaran menjadi 'ditolak' jika melewati proposal_selesai
    """
    try:
        # Cek jadwal proposal selesai
        cursor.execute('''
            SELECT proposal_selesai FROM pengaturan_jadwal 
            ORDER BY id DESC LIMIT 1
        ''')
        jadwal = cursor.fetchone()
        
        if not jadwal or not jadwal.get('proposal_selesai'):
            log_info("Tidak ada jadwal proposal selesai, skip auto-reject anggaran")
            return
        
        jadwal_selesai = jadwal['proposal_selesai']
        now = datetime.now()
        
        if now <= jadwal_selesai:
            log_info("Jadwal proposal belum selesai, skip auto-reject anggaran")
            return
        
        log_info(f"Jadwal proposal sudah selesai ({jadwal_selesai}), cek anggaran untuk auto-reject")
        
        # Cek anggaran awal yang belum disetujui
        cursor.execute('''
            SELECT id, id_proposal, kegiatan_utama, kegiatan, nama_barang
            FROM anggaran_awal 
            WHERE status IN ('draf', 'diajukan', 'revisi')
            AND tanggal_buat < %s
        ''', (jadwal_selesai,))
        
        anggaran_awal_to_reject = cursor.fetchall()
        log_info(f"Found {len(anggaran_awal_to_reject)} anggaran awal to reject")
        
        # Cek anggaran bertumbuh yang belum disetujui
        cursor.execute('''
            SELECT id, id_proposal, kegiatan_utama, kegiatan, nama_barang
            FROM anggaran_bertumbuh 
            WHERE status IN ('draf', 'diajukan', 'revisi')
            AND tanggal_buat < %s
        ''', (jadwal_selesai,))
        
        anggaran_bertumbuh_to_reject = cursor.fetchall()
        log_info(f"Found {len(anggaran_bertumbuh_to_reject)} anggaran bertumbuh to reject")
        
        # Update anggaran awal menjadi ditolak
        for anggaran in anggaran_awal_to_reject:
            try:
                cursor.execute('''
                    UPDATE anggaran_awal 
                    SET status = 'ditolak' 
                    WHERE id = %s
                ''', (anggaran['id'],))
                log_info(f"Auto-rejected anggaran awal {anggaran['id']} ({anggaran['kegiatan']} - {anggaran['nama_barang']})")
            except Exception as e:
                log_error(f"Error rejecting anggaran awal {anggaran['id']}: {str(e)}")
        
        # Update anggaran bertumbuh menjadi ditolak
        for anggaran in anggaran_bertumbuh_to_reject:
            try:
                cursor.execute('''
                    UPDATE anggaran_bertumbuh 
                    SET status = 'ditolak' 
                    WHERE id = %s
                ''', (anggaran['id'],))
                log_info(f"Auto-rejected anggaran bertumbuh {anggaran['id']} ({anggaran['kegiatan']} - {anggaran['nama_barang']})")
            except Exception as e:
                log_error(f"Error rejecting anggaran bertumbuh {anggaran['id']}: {str(e)}")
        
        total_rejected = len(anggaran_awal_to_reject) + len(anggaran_bertumbuh_to_reject)
        if total_rejected > 0:
            log_info(f"Auto-reject anggaran completed: {total_rejected} items rejected")
        else:
            log_info("No anggaran to reject - all items already processed")
        
    except Exception as e:
        log_error(f"Error in auto_reject_anggaran_worker: {str(e)}")

def auto_reject_laporan_kemajuan_worker(cursor, log_info, log_error):
    """
    Otomatis mengubah status laporan kemajuan menjadi 'ditolak' jika melewati kemajuan_selesai
    """
    try:
        # Cek jadwal kemajuan selesai
        cursor.execute('''
            SELECT kemajuan_selesai FROM pengaturan_jadwal 
            ORDER BY id DESC LIMIT 1
        ''')
        jadwal = cursor.fetchone()
        
        if not jadwal or not jadwal.get('kemajuan_selesai'):
            log_info("Tidak ada jadwal kemajuan selesai, skip auto-reject laporan kemajuan")
            return
        
        jadwal_selesai = jadwal['kemajuan_selesai']
        now = datetime.now()
        
        if now <= jadwal_selesai:
            log_info("Jadwal kemajuan belum selesai, skip auto-reject laporan kemajuan")
            return
        
        log_info(f"Jadwal kemajuan sudah selesai ({jadwal_selesai}), cek laporan kemajuan untuk auto-reject")
        
        # Cek laporan kemajuan awal yang belum disetujui
        cursor.execute('''
            SELECT id, id_proposal, kegiatan_utama, kegiatan, nama_barang
            FROM laporan_kemajuan_awal 
            WHERE status IN ('draf', 'diajukan', 'revisi')
            AND tanggal_buat < %s
        ''', (jadwal_selesai,))
        
        laporan_kemajuan_awal_to_reject = cursor.fetchall()
        log_info(f"Found {len(laporan_kemajuan_awal_to_reject)} laporan kemajuan awal to reject")
        
        # Cek laporan kemajuan bertumbuh yang belum disetujui
        cursor.execute('''
            SELECT id, id_proposal, kegiatan_utama, kegiatan, nama_barang
            FROM laporan_kemajuan_bertumbuh 
            WHERE status IN ('draf', 'diajukan', 'revisi')
            AND tanggal_buat < %s
        ''', (jadwal_selesai,))
        
        laporan_kemajuan_bertumbuh_to_reject = cursor.fetchall()
        log_info(f"Found {len(laporan_kemajuan_bertumbuh_to_reject)} laporan kemajuan bertumbuh to reject")
        
        # Update laporan kemajuan awal menjadi ditolak
        for laporan in laporan_kemajuan_awal_to_reject:
            try:
                cursor.execute('''
                    UPDATE laporan_kemajuan_awal 
                    SET status = 'ditolak' 
                    WHERE id = %s
                ''', (laporan['id'],))
                log_info(f"Auto-rejected laporan kemajuan awal {laporan['id']} ({laporan['kegiatan']} - {laporan['nama_barang']})")
            except Exception as e:
                log_error(f"Error rejecting laporan kemajuan awal {laporan['id']}: {str(e)}")
        
        # Update laporan kemajuan bertumbuh menjadi ditolak
        for laporan in laporan_kemajuan_bertumbuh_to_reject:
            try:
                cursor.execute('''
                    UPDATE laporan_kemajuan_bertumbuh 
                    SET status = 'ditolak' 
                    WHERE id = %s
                ''', (laporan['id'],))
                log_info(f"Auto-rejected laporan kemajuan bertumbuh {laporan['id']} ({laporan['kegiatan']} - {laporan['nama_barang']})")
            except Exception as e:
                log_error(f"Error rejecting laporan kemajuan bertumbuh {laporan['id']}: {str(e)}")
        
        total_rejected = len(laporan_kemajuan_awal_to_reject) + len(laporan_kemajuan_bertumbuh_to_reject)
        if total_rejected > 0:
            log_info(f"Auto-reject laporan kemajuan completed: {total_rejected} items rejected")
        else:
            log_info("No laporan kemajuan to reject - all items already processed")
        
    except Exception as e:
        log_error(f"Error in auto_reject_laporan_kemajuan_worker: {str(e)}")

def auto_reject_laporan_akhir_worker(cursor, log_info, log_error):
    """
    Otomatis mengubah status laporan akhir menjadi 'ditolak' jika melewati akhir_selesai
    """
    try:
        # Cek jadwal akhir selesai
        cursor.execute('''
            SELECT akhir_selesai FROM pengaturan_jadwal 
            ORDER BY id DESC LIMIT 1
        ''')
        jadwal = cursor.fetchone()
        
        if not jadwal or not jadwal.get('akhir_selesai'):
            log_info("Tidak ada jadwal akhir selesai, skip auto-reject laporan akhir")
            return
        
        jadwal_selesai = jadwal['akhir_selesai']
        now = datetime.now()
        
        if now <= jadwal_selesai:
            log_info("Jadwal akhir belum selesai, skip auto-reject laporan akhir")
            return
        
        log_info(f"Jadwal akhir sudah selesai ({jadwal_selesai}), cek laporan akhir untuk auto-reject")
        
        # Cek laporan akhir awal yang belum disetujui
        cursor.execute('''
            SELECT id, id_proposal, kegiatan_utama, kegiatan, nama_barang
            FROM laporan_akhir_awal 
            WHERE status IN ('draf', 'diajukan', 'revisi')
            AND tanggal_buat < %s
        ''', (jadwal_selesai,))
        
        laporan_akhir_awal_to_reject = cursor.fetchall()
        log_info(f"Found {len(laporan_akhir_awal_to_reject)} laporan akhir awal to reject")
        
        # Cek laporan akhir bertumbuh yang belum disetujui
        cursor.execute('''
            SELECT id, id_proposal, kegiatan_utama, kegiatan, nama_barang
            FROM laporan_akhir_bertumbuh 
            WHERE status IN ('draf', 'diajukan', 'revisi')
            AND tanggal_buat < %s
        ''', (jadwal_selesai,))
        
        laporan_akhir_bertumbuh_to_reject = cursor.fetchall()
        log_info(f"Found {len(laporan_akhir_bertumbuh_to_reject)} laporan akhir bertumbuh to reject")
        
        # Update laporan akhir awal menjadi ditolak
        for laporan in laporan_akhir_awal_to_reject:
            try:
                cursor.execute('''
                    UPDATE laporan_akhir_awal 
                    SET status = 'ditolak' 
                    WHERE id = %s
                ''', (laporan['id'],))
                log_info(f"Auto-rejected laporan akhir awal {laporan['id']} ({laporan['kegiatan']} - {laporan['nama_barang']})")
            except Exception as e:
                log_error(f"Error rejecting laporan akhir awal {laporan['id']}: {str(e)}")
        
        # Update laporan akhir bertumbuh menjadi ditolak
        for laporan in laporan_akhir_bertumbuh_to_reject:
            try:
                cursor.execute('''
                    UPDATE laporan_akhir_bertumbuh 
                    SET status = 'ditolak' 
                    WHERE id = %s
                ''', (laporan['id'],))
                log_info(f"Auto-rejected laporan akhir bertumbuh {laporan['id']} ({laporan['kegiatan']} - {laporan['nama_barang']})")
            except Exception as e:
                log_error(f"Error rejecting laporan akhir bertumbuh {laporan['id']}: {str(e)}")
        
        total_rejected = len(laporan_akhir_awal_to_reject) + len(laporan_akhir_bertumbuh_to_reject)
        if total_rejected > 0:
            log_info(f"Auto-reject laporan akhir completed: {total_rejected} items rejected")
        else:
            log_info("No laporan akhir to reject - all items already processed")
        
    except Exception as e:
        log_error(f"Error in auto_reject_laporan_akhir_worker: {str(e)}")

def auto_update_status_selesai_worker(cursor, log_info, log_error):
    """
    Otomatis mengubah status mahasiswa dan proposal menjadi 'selesai'
    jika sudah melewati akhir_selesai dan semua penilaian sudah lengkap
    """
    try:
        # Cek jadwal akhir selesai
        cursor.execute('''
            SELECT akhir_selesai FROM pengaturan_jadwal 
            ORDER BY id DESC LIMIT 1
        ''')
        jadwal = cursor.fetchone()
        
        if not jadwal or not jadwal.get('akhir_selesai'):
            log_info("Tidak ada jadwal akhir selesai, skip auto-update status")
            return
        
        jadwal_selesai = jadwal['akhir_selesai']
        now = datetime.now()
        
        if now <= jadwal_selesai:
            log_info("Jadwal akhir belum selesai, skip auto-update status")
            return
        
        log_info(f"Jadwal akhir sudah selesai ({jadwal_selesai}), cek status selesai")
        
        # Cek proposal yang sudah lengkap penilaiannya (mahasiswa, proposal, kemajuan, akhir)
        cursor.execute('''
            SELECT DISTINCT 
                p.id as proposal_id, 
                p.nim, 
                p.status as status_proposal, 
                m.status as status_mahasiswa,
                m.nama_ketua
            FROM proposal p
            JOIN mahasiswa m ON p.nim = m.nim
            WHERE p.status_admin = 'lolos'
            AND EXISTS (SELECT 1 FROM penilaian_mahasiswa pm WHERE pm.id_proposal = p.id)
            AND EXISTS (SELECT 1 FROM penilaian_proposal pp WHERE pp.id_proposal = p.id)
            AND EXISTS (SELECT 1 FROM penilaian_laporan_kemajuan plk WHERE plk.id_proposal = p.id)
            AND EXISTS (SELECT 1 FROM penilaian_laporan_akhir pla WHERE pla.id_proposal = p.id)
            AND (p.status != 'selesai' OR m.status != 'selesai')
        ''')
        
        proposals_to_update = cursor.fetchall()
        log_info(f"Found {len(proposals_to_update)} proposals with complete assessments ready for status update")
        
        updated_proposals = 0
        updated_mahasiswa = 0
        
        for proposal in proposals_to_update:
            try:
                # Update status proposal menjadi 'selesai'
                if proposal['status_proposal'] != 'selesai':
                    cursor.execute('''
                        UPDATE proposal SET status = 'selesai' WHERE id = %s
                    ''', (proposal['proposal_id'],))
                    updated_proposals += 1
                    log_info(f"Updated proposal {proposal['proposal_id']} ({proposal['nama_ketua']}) status to 'selesai'")
                
                # Update status mahasiswa menjadi 'selesai'
                if proposal['status_mahasiswa'] != 'selesai':
                    cursor.execute('''
                        UPDATE mahasiswa SET status = 'selesai' WHERE nim = %s
                    ''', (proposal['nim'],))
                    updated_mahasiswa += 1
                    log_info(f"Updated mahasiswa {proposal['nim']} ({proposal['nama_ketua']}) status to 'selesai'")
                
            except Exception as e:
                log_error(f"Error updating status for proposal {proposal['proposal_id']}: {str(e)}")
        
        if updated_proposals > 0 or updated_mahasiswa > 0:
            log_info(f"Auto-update status selesai completed: {updated_proposals} proposals, {updated_mahasiswa} mahasiswa updated")
        else:
            log_info("No status updates needed - all proposals and mahasiswa already have 'selesai' status")
        
    except Exception as e:
        log_error(f"Error in auto_update_status_selesai_worker: {str(e)}")

# Fungsi legacy sudah dihapus karena tidak digunakan lagi

if __name__ == '__main__':
    # ✅ PERBAIKAN: Mulai worker auto-fill secara otomatis saat aplikasi dimulai
    start_auto_fill_worker_on_startup()
    
    print("Sistem auto-fill berjalan di background...")
    app.run(debug=True, host='0.0.0.0', port=5000)
